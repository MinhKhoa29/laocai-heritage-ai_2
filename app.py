# streamlit run app.py --server.runOnSave true
# git add .
# git commit -m "mô tả thay đổi"
# git push

import os
import re
import unicodedata
import json as pyjson
import streamlit as st
import streamlit.components.v1 as components

import sqlite3
import hashlib
import secrets
from datetime import datetime, timedelta

from textwrap import dedent
from html import escape
from google import genai
from google.genai import types
import base64
from dotenv import load_dotenv
load_dotenv()

DB_PATH = "laocai_heritage_ai.db"
SESSION_HOURS = 24 * 7  # 7 ngày

def get_db_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_auth_db():
    conn = get_db_conn()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS user_sessions (
            sid TEXT PRIMARY KEY,
            user_id INTEGER NOT NULL,
            expires_at TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
                
    """)

    cur.execute("""
        CREATE TABLE IF NOT EXISTS community_posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            username TEXT NOT NULL,
            content TEXT NOT NULL,
            image_url TEXT,
            category TEXT,
            area TEXT,
            created_at TEXT NOT NULL,
            is_seed INTEGER NOT NULL DEFAULT 0,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
    """)

    conn.commit()
    conn.close()

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()

def create_user(username: str, password: str):
    username = str(username or "").strip()
    password = str(password or "").strip()

    if len(username) < 4:
        return False, "Tên đăng nhập phải từ 4 ký tự."
    if len(password) < 6:
        return False, "Mật khẩu phải từ 6 ký tự."

    conn = get_db_conn()
    cur = conn.cursor()

    try:
        cur.execute(
            "INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
            (username, hash_password(password), datetime.utcnow().isoformat())
        )
        conn.commit()
        return True, "Đăng ký thành công."
    except sqlite3.IntegrityError:
        return False, "Tên đăng nhập đã tồn tại."
    finally:
        conn.close()

def verify_user(username: str, password: str):
    username = str(username or "").strip()
    password = str(password or "").strip()

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "SELECT * FROM users WHERE username = ?",
        (username,)
    )
    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    if row["password_hash"] != hash_password(password):
        return None

    return {
        "id": row["id"],
        "username": row["username"]
    }

def get_user_by_id(user_id: int):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT id, username, created_at FROM users WHERE id = ?", (user_id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        return None

    return {
        "id": row["id"],
        "username": row["username"],
        "created_at": row["created_at"]
    }

def update_user_password(user_id: int, old_password: str, new_password: str):
    old_password = str(old_password or "").strip()
    new_password = str(new_password or "").strip()

    if len(new_password) < 6:
        return False, "Mật khẩu mới phải từ 6 ký tự."

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT password_hash FROM users WHERE id = ?", (user_id,))
    row = cur.fetchone()

    if not row:
        conn.close()
        return False, "Không tìm thấy tài khoản."

    if row["password_hash"] != hash_password(old_password):
        conn.close()
        return False, "Mật khẩu hiện tại chưa đúng."

    cur.execute(
        "UPDATE users SET password_hash = ? WHERE id = ?",
        (hash_password(new_password), user_id)
    )
    conn.commit()
    conn.close()

    return True, "Đổi mật khẩu thành công."

def create_session(user_id: int) -> str:
    sid = secrets.token_urlsafe(32)
    now = datetime.utcnow()
    expires = now + timedelta(hours=SESSION_HOURS)

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO user_sessions (sid, user_id, expires_at, created_at) VALUES (?, ?, ?, ?)",
        (sid, user_id, expires.isoformat(), now.isoformat())
    )
    conn.commit()
    conn.close()
    return sid

def get_user_by_sid(sid: str):
    sid = str(sid or "").strip()
    if not sid:
        return None

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("""
        SELECT s.sid, s.user_id, s.expires_at, u.username
        FROM user_sessions s
        JOIN users u ON u.id = s.user_id
        WHERE s.sid = ?
    """, (sid,))
    row = cur.fetchone()

    if not row:
        conn.close()
        return None

    expires_at = row["expires_at"]
    try:
        if datetime.fromisoformat(expires_at) < datetime.utcnow():
            cur.execute("DELETE FROM user_sessions WHERE sid = ?", (sid,))
            conn.commit()
            conn.close()
            return None
    except Exception:
        cur.execute("DELETE FROM user_sessions WHERE sid = ?", (sid,))
        conn.commit()
        conn.close()
        return None

    conn.close()
    return {
        "id": row["user_id"],
        "username": row["username"],
        "sid": row["sid"]
    }

def delete_session(sid: str):
    sid = str(sid or "").strip()
    if not sid:
        return

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM user_sessions WHERE sid = ?", (sid,))
    conn.commit()
    conn.close()

def save_chat_message(user_id: int, role: str, content: str):
    role = str(role or "").strip()
    content = str(content or "").strip()
    if not user_id or not role or not content:
        return

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO chat_messages (user_id, role, content, created_at) VALUES (?, ?, ?, ?)",
        (user_id, role, content, datetime.utcnow().isoformat())
    )
    conn.commit()
    conn.close()

def load_chat_history(user_id: int, limit: int = 100):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("""
        SELECT role, content
        FROM chat_messages
        WHERE user_id = ?
        ORDER BY id ASC
        LIMIT ?
    """, (user_id, limit))
    rows = cur.fetchall()
    conn.close()

    return [{"role": row["role"], "content": row["content"]} for row in rows]

def clear_chat_history(user_id: int):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("DELETE FROM chat_messages WHERE user_id = ?", (user_id,))
    conn.commit()
    conn.close()

def seed_default_community_posts():
    conn = get_db_conn()
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) AS total FROM community_posts")
    total = cur.fetchone()["total"]

    if total > 0:
        conn.close()
        return

    seed_posts = [
        {
            "username": "Khánh Linh",
            "content": "Mình rất ấn tượng với lễ hội Gầu Tào của người Mông ở Bắc Hà. Không khí rất vui, nhiều hoạt động cộng đồng và mang đậm bản sắc vùng cao.",
            "image_url": "https://images.unsplash.com/photo-1516483638261-f4dbaf036963?q=80&w=1400&auto=format&fit=crop",
            "category": "Lễ hội",
            "area": "Bắc Hà"
        },
        {
            "username": "Minh Khoa",
            "content": "Làng nghề thổ cẩm Tả Phìn là nơi rất đáng để tìm hiểu. Các sản phẩm thủ công đẹp, tinh tế và thể hiện rõ nét văn hóa địa phương.",
            "image_url": "https://images.unsplash.com/photo-1512436991641-6745cdb1723f?q=80&w=1400&auto=format&fit=crop",
            "category": "Làng nghề",
            "area": "Sa Pa"
        },
        {
            "username": "Thu Trang",
            "content": "Cuối tuần mình có dịp ghé chợ phiên Bắc Hà và thấy nơi đây rất nhộn nhịp. Ẩm thực, trang phục và không khí vùng cao đều rất đặc sắc.",
            "image_url": "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1400&auto=format&fit=crop",
            "category": "Trải nghiệm",
            "area": "Bắc Hà"
        }
    ]

    now = datetime.utcnow().isoformat()

    for item in seed_posts:
        cur.execute("""
            INSERT INTO community_posts (
                user_id, username, content, image_url, category, area, created_at, is_seed
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            None,
            item["username"],
            item["content"],
            item["image_url"],
            item["category"],
            item["area"],
            now,
            1
        ))

    conn.commit()
    conn.close()


def create_community_post(user_id: int, username: str, content: str, image_url: str = "", category: str = "", area: str = ""):
    content = str(content or "").strip()
    image_url = str(image_url or "").strip()
    category = str(category or "").strip()
    area = str(area or "").strip()
    username = str(username or "").strip()

    if len(content) < 8:
        return False, "Nội dung bài viết phải từ 8 ký tự."

    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO community_posts (
            user_id, username, content, image_url, category, area, created_at, is_seed
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        user_id,
        username,
        content,
        image_url,
        category,
        area,
        datetime.utcnow().isoformat(),
        0
    ))
    conn.commit()
    conn.close()

    return True, "Đăng bài thành công."


def load_community_posts(limit: int = 50):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("""
        SELECT id, user_id, username, content, image_url, category, area, created_at, is_seed
        FROM community_posts
        ORDER BY id DESC
        LIMIT ?
    """, (limit,))
    rows = cur.fetchall()
    conn.close()

    return [dict(row) for row in rows]


def format_post_time(iso_text: str) -> str:
    try:
        dt = datetime.fromisoformat(iso_text)
        return dt.strftime("%d/%m/%Y • %H:%M")
    except Exception:
        return "Vừa xong"


def ensure_default_chat_history(user_id: int, default_messages=None):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) AS total FROM chat_messages WHERE user_id = ?", (user_id,))
    total = cur.fetchone()["total"]
    conn.close()

    if total == 0 and default_messages:
        for msg in default_messages:
            role = str(msg.get("role", "")).strip()
            content = str(msg.get("content", "")).strip()

            if role in ["user", "assistant"] and content:
                save_chat_message(user_id, role, content)

def build_app_url(page_name: str = "home", **params):
    query = {"page": page_name}

    current_sid = st.session_state.get("auth_sid", "")
    if current_sid:
        query["sid"] = current_sid

    for k, v in params.items():
        if v is not None and str(v).strip() != "":
            query[k] = v

    parts = [f"{k}={v}" for k, v in query.items()]
    return "?" + "&".join(parts)

def auth_logout():
    sid = st.session_state.get("auth_sid", "")
    if sid:
        delete_session(sid)

    for key in [
        "auth_user",
        "auth_sid",
        "auth_checked",
        "gemini_messages",
        "gemini_loaded_user_id"
    ]:
        if key in st.session_state:
            del st.session_state[key]

    st.query_params.clear()
    st.query_params["page"] = "home"
    st.rerun()

def restore_auth_from_query():
    if st.session_state.get("auth_checked"):
        return

    st.session_state.auth_checked = True
    sid = str(st.query_params.get("sid", "")).strip()

    if not sid:
        return

    user = get_user_by_sid(sid)
    if user:
        st.session_state.auth_user = {
            "id": user["id"],
            "username": user["username"]
        }
        st.session_state.auth_sid = user["sid"]

def get_gemini_api_key():
    try:
        key = st.secrets.get("GEMINI_API_KEY", "").strip()
        if key:
            return key
    except Exception:
        pass

    return os.environ.get("GEMINI_API_KEY", "").strip()

def image_to_data_uri(candidates):
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }

    for path in candidates:
        if path and os.path.exists(path):
            ext = os.path.splitext(path)[1].lower()
            mime = mime_map.get(ext, "image/png")
            try:
                with open(path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode("utf-8")
                return f"data:{mime};base64,{encoded}"
            except Exception:
                pass

    return ""

st.set_page_config(
    page_title="LAO CAI HERITAGE AI",
    page_icon="assets/anime_teamtrangchu1.png",
    layout="wide"
)

init_auth_db()
seed_default_community_posts()
restore_auth_from_query()

if str(st.query_params.get("logout", "0")) == "1":
    auth_logout()

page = st.query_params.get("page", "home")

page = st.query_params.get("page", "home")


# ChatBot
# =========================
# CHATBOT LOGIC UPGRADE
# Thay toàn bộ đoạn từ:
#   # ChatBot
# đến ngay trước:
#   hero_home_src = image_to_data_uri([
# =========================

# ChatBot
import math
from collections import Counter


def _load_json_safe(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = pyjson.load(f)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def _normalize_free_text(value):
    text = str(value or "").strip().lower()
    text = text.replace("đ", "d").replace("Đ", "D")
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return text


def normalize_query_text(value: str) -> str:
    return _normalize_free_text(value)


def tokenize_query(value: str) -> list[str]:
    text = normalize_query_text(value)
    raw_tokens = re.split(r"[^a-zA-Z0-9]+", text)
    stopwords = {
        "la", "va", "o", "di", "den", "cho", "toi", "minh", "ban", "co", "gi",
        "nao", "khong", "duoc", "mot", "nhung", "cac", "ve", "tai", "ngay", "dem",
        "tham", "khao", "vui", "giup", "em", "anh", "chi", "noi", "nay", "kia"
    }
    return [tok for tok in raw_tokens if len(tok) >= 2 and tok not in stopwords]


def _safe_join_text(parts) -> str:
    return " ".join(str(x or "").strip() for x in parts if str(x or "").strip())


def _load_community_ai_docs(limit: int = 30):
    docs = []
    try:
        posts = load_community_posts(limit=limit)
    except Exception:
        posts = []

    for post in posts:
        content = str(post.get("content", "")).strip()
        if len(content) < 8:
            continue

        area = str(post.get("area", "")).strip()
        category = str(post.get("category", "")).strip()
        username = str(post.get("username", "")).strip()
        created_at = str(post.get("created_at", "")).strip()

        docs.append({
            "id": f"community_{post.get('id', '')}",
            "doc_type": "cong_dong",
            "title": f"Bài đăng cộng đồng {username}".strip(),
            "content": content,
            "metadata": {
                "area": area,
                "category": category,
                "username": username,
                "created_at": created_at,
                "source": "community"
            },
            "source_payload": post
        })

    return docs


def _build_ai_knowledge_base():
    diemden = _load_json_safe("diemden.json")
    lichtrinh = _load_json_safe("lichtrinh.json")
    lehoi_langnghe = _load_json_safe("lehoi_langnghe.json")

    documents = []
    places = []
    routes = []
    festivals = []

    for idx, item in enumerate(diemden[:300], start=1):
        highlights = item.get("highlights", [])
        if not isinstance(highlights, list):
            highlights = []

        place_obj = {
            "name": str(item.get("name", "")).strip(),
            "slug": str(item.get("slug", "")).strip(),
            "area": str(item.get("area", "")).strip(),
            "category": str(item.get("category", "")).strip(),
            "season": str(item.get("season", "")).strip(),
            "best_time": str(item.get("best_time", "")).strip(),
            "short_desc": str(item.get("short_desc", "")).strip(),
            "full_desc": str(item.get("full_desc", "")).strip(),
            "image": str(item.get("image", "")).strip(),
            "highlights": [str(x).strip() for x in highlights[:6] if str(x).strip()]
        }
        places.append(place_obj)

        documents.append({
            "id": f"place_{idx}",
            "doc_type": "diem_den",
            "title": place_obj["name"],
            "content": _safe_join_text([
                place_obj["name"],
                place_obj["area"],
                place_obj["category"],
                place_obj["season"],
                place_obj["best_time"],
                place_obj["short_desc"],
                place_obj["full_desc"],
                " ".join(place_obj["highlights"]),
            ]),
            "metadata": {
                "area": place_obj["area"],
                "category": place_obj["category"],
                "season": place_obj["season"],
                "best_time": place_obj["best_time"],
                "source": "diemden"
            },
            "source_payload": place_obj
        })

    for idx, item in enumerate(lichtrinh[:300], start=1):
        route_obj = {
            "from": str(item.get("from", "")).strip(),
            "to": str(item.get("to", "")).strip(),
            "slug": str(item.get("slug", "")).strip(),
            "category": str(item.get("category", "")).strip(),
            "time_estimate": str(item.get("time_estimate", "")).strip(),
            "distance_km": str(item.get("distance_km", "")).strip(),
            "transport_name": str(item.get("transport_name", "")).strip(),
            "transport_price": str(item.get("transport_price", "")).strip(),
            "ticket_price": str(item.get("ticket_price", "")).strip(),
            "hotel_name": str(item.get("hotel_name", "")).strip(),
            "hotel_price": str(item.get("hotel_price", "")).strip(),
            "total_price": str(item.get("total_price", "")).strip(),
            "note": str(item.get("note", "")).strip(),
            "short_desc": str(item.get("short_desc", "")).strip()
        }
        routes.append(route_obj)

        documents.append({
            "id": f"route_{idx}",
            "doc_type": "lich_trinh",
            "title": f"{route_obj['from']} - {route_obj['to']}".strip(" -"),
            "content": _safe_join_text([
                route_obj["from"],
                route_obj["to"],
                route_obj["category"],
                route_obj["time_estimate"],
                route_obj["distance_km"],
                route_obj["transport_name"],
                route_obj["transport_price"],
                route_obj["ticket_price"],
                route_obj["hotel_name"],
                route_obj["hotel_price"],
                route_obj["total_price"],
                route_obj["note"],
                route_obj["short_desc"],
            ]),
            "metadata": {
                "from": route_obj["from"],
                "to": route_obj["to"],
                "category": route_obj["category"],
                "time_estimate": route_obj["time_estimate"],
                "distance_km": route_obj["distance_km"],
                "total_price": route_obj["total_price"],
                "source": "lichtrinh"
            },
            "source_payload": route_obj
        })

    for idx, item in enumerate(lehoi_langnghe[:200], start=1):
        festival_obj = {
            "name": str(item.get("name", "")).strip(),
            "type": str(item.get("type", "")).strip(),
            "area": str(item.get("area", "")).strip(),
            "season": str(item.get("season", "")).strip(),
            "short_desc": str(item.get("short_desc", "")).strip(),
            "full_desc": str(item.get("full_desc", "")).strip(),
            "image": str(item.get("image", "")).strip(),
            "highlights": [str(x).strip() for x in item.get("highlights", []) if str(x).strip()]
            if isinstance(item.get("highlights", []), list) else []
        }
        festivals.append(festival_obj)

        documents.append({
            "id": f"festival_{idx}",
            "doc_type": festival_obj["type"] or "le_hoi_lang_nghe",
            "title": festival_obj["name"],
            "content": _safe_join_text([
                festival_obj["name"],
                festival_obj["type"],
                festival_obj["area"],
                festival_obj["season"],
                festival_obj["short_desc"],
                festival_obj["full_desc"],
                " ".join(festival_obj["highlights"]),
            ]),
            "metadata": {
                "area": festival_obj["area"],
                "category": festival_obj["type"],
                "season": festival_obj["season"],
                "source": "lehoi_langnghe"
            },
            "source_payload": festival_obj
        })

    documents.extend(_load_community_ai_docs(limit=30))

    return {
        "brand": "Lao Cai Heritage AI",
        "phone": "0346 538 917",
        "places": places,
        "routes": routes,
        "festivals": festivals,
        "documents": documents,
        "quick_questions": [
            "Sa Pa có gì đẹp?",
            "Mùa nào đi Fansipan đẹp?",
            "Gợi ý lịch trình đi Bắc Hà",
            "Chi phí tham khảo thế nào?"
        ]
    }


CHATBOT_PAYLOAD = _build_ai_knowledge_base()
GEMINI_MODEL = "gemini-2.5-flash"


PREFERENCE_PATTERNS = {
    "budget_low": ["re", "thap", "tiet kiem", "gia re", "it tien", "sinh vien"],
    "budget_mid": ["vua phai", "hop ly", "tam trung"],
    "budget_high": ["cao cap", "sang", "xinh", "nghi duong", "cao"],

    "nature": ["thien nhien", "san may", "nui", "rung", "hoang so", "canh dep", "thac", "ruong bac thang"],
    "culture": ["van hoa", "ban sac", "dan toc", "le hoi", "lang nghe", "truyen thong", "cho phien"],
    "history": ["lich su", "di tich", "tam linh", "den", "chua"],
    "relax": ["yen tinh", "thu gian", "nghi duong", "chill", "it dong", "vang ve"],
    "adventure": ["trekking", "leo nui", "kham pha", "mao hiem", "phuot"],
    "food": ["am thuc", "an ngon", "dac san", "mon ngon", "an uong", "quan an"],

    "family": ["gia dinh", "tre em", "bo me", "ca nha"],
    "friends": ["nhom ban", "ban be", "team", "hoi ban"],
    "couple": ["cap doi", "nguoi yeu", "lang man"],
    "solo": ["mot minh", "di mot minh", "solo"],

    "crowd_low": ["it dong", "yen tinh", "vang", "tranh dong"],
    "crowd_high": ["dong vui", "nhon nhip", "soi dong"],

    "weekend": ["cuoi tuan", "thu 7", "chu nhat"],
}


def _extract_duration_preferences(q_norm: str) -> dict:
    prefs = {}
    if "2 ngay 1 dem" in q_norm:
        prefs["duration"] = "2 ngày 1 đêm"
    elif "3 ngay 2 dem" in q_norm:
        prefs["duration"] = "3 ngày 2 đêm"
    elif "1 ngay" in q_norm:
        prefs["duration"] = "1 ngày"
    return prefs

def _extract_companion_preferences(q_norm: str) -> dict:
    prefs = {}

    if any(x in q_norm for x in ["gia dinh", "tre em", "bo me", "ca nha"]):
        prefs["companion_type"] = "gia_dinh"
    elif any(x in q_norm for x in ["nhom ban", "ban be", "team", "hoi ban"]):
        prefs["companion_type"] = "ban_be"
    elif any(x in q_norm for x in ["cap doi", "nguoi yeu", "lang man"]):
        prefs["companion_type"] = "cap_doi"
    elif any(x in q_norm for x in ["mot minh", "di mot minh", "solo"]):
        prefs["companion_type"] = "mot_minh"

    return prefs


def _extract_crowd_preferences(q_norm: str) -> dict:
    prefs = {}

    if any(x in q_norm for x in ["it dong", "yen tinh", "vang", "tranh dong"]):
        prefs["crowd_preference"] = "thap"
    elif any(x in q_norm for x in ["dong vui", "nhon nhip", "soi dong"]):
        prefs["crowd_preference"] = "cao"

    return prefs


def _extract_weather_preferences(q_norm: str) -> dict:
    prefs = {}

    if any(x in q_norm for x in ["mat me", "troi mat", "lanh", "se lanh"]):
        prefs["weather_preference"] = "mat"
    elif any(x in q_norm for x in ["nang dep", "troi dep", "kho rao"]):
        prefs["weather_preference"] = "nang_dep"
    elif any(x in q_norm for x in ["khong mua", "tranh mua", "it mua"]):
        prefs["weather_preference"] = "kho_rao"

    return prefs


def _extract_trip_time_preferences(q_norm: str) -> dict:
    prefs = {}

    if any(x in q_norm for x in ["cuoi tuan", "thu 7", "chu nhat"]):
        prefs["trip_time"] = "cuoi_tuan"
    elif any(x in q_norm for x in ["ngay thuong", "trong tuan"]):
        prefs["trip_time"] = "ngay_thuong"
    elif any(x in q_norm for x in ["dip le", "nghi le"]):
        prefs["trip_time"] = "dip_le"

    return prefs

def analyze_user_query(user_message: str, payload: dict, messages: list[dict] | None = None) -> dict:
    q_norm = normalize_query_text(user_message)
    tokens = tokenize_query(user_message)

    intent = "general"
    if any(k in q_norm for k in ["goi y", "nen di dau", "phu hop", "chon", "de xuat"]):
        intent = "goi_y"
    elif any(k in q_norm for k in ["chi phi", "bao nhieu", "gia", "tong tien"]):
        intent = "chi_phi"
    elif any(k in q_norm for k in ["lich trinh", "hanh trinh", "di nhu nao"]):
        intent = "lich_trinh"
    elif any(k in q_norm for k in ["mua nao", "thoi diem", "mua dep"]):
        intent = "thoi_diem"
    elif any(k in q_norm for k in ["le hoi", "lang nghe"]):
        intent = "van_hoa"
    elif any(k in q_norm for k in ["hom nay", "hien tai", "moi nhat", "cap nhat", "thoi tiet"]):
        intent = "need_web"

    entities = {"places": [], "areas": [], "route_to": []}

    for item in payload.get("places", []):
        name = str(item.get("name", "")).strip()
        area = str(item.get("area", "")).strip()
        if name and normalize_query_text(name) in q_norm:
            entities["places"].append(name)
        if area and normalize_query_text(area) in q_norm and area not in entities["areas"]:
            entities["areas"].append(area)

    for item in payload.get("routes", []):
        to_place = str(item.get("to", "")).strip()
        if to_place and normalize_query_text(to_place) in q_norm:
            entities["route_to"].append(to_place)

    preferences = {}
    preferences.update(_extract_duration_preferences(q_norm))
    preferences.update(_extract_companion_preferences(q_norm))
    preferences.update(_extract_crowd_preferences(q_norm))
    preferences.update(_extract_weather_preferences(q_norm))
    preferences.update(_extract_trip_time_preferences(q_norm))

    for pref_key, keywords in PREFERENCE_PATTERNS.items():
        if any(k in q_norm for k in keywords):
            preferences[pref_key] = True

        if any(x in q_norm for x in ["gan", "xung quanh", "gan day", "khong qua xa"]):
            preferences["nearby_preference"] = True

        if any(x in q_norm for x in ["an ngon", "dac san", "am thuc", "an uong"]):
            preferences["food_focus"] = True

    if "mua xuan" in q_norm:
        preferences["season"] = "mùa xuân"
    elif "mua he" in q_norm:
        preferences["season"] = "mùa hè"
    elif "mua thu" in q_norm:
        preferences["season"] = "mùa thu"
    elif "mua dong" in q_norm:
        preferences["season"] = "mùa đông"

    return {
        "original_question": user_message,
        "normalized_question": q_norm,
        "tokens": tokens,
        "intent": intent,
        "entities": entities,
        "preferences": preferences,
    }



def _merge_preferences(old_prefs: dict | None, new_prefs: dict | None) -> dict:
    merged = dict(old_prefs or {})
    for k, v in (new_prefs or {}).items():
        if v not in [None, "", False, [], {}]:
            merged[k] = v
    return merged


def get_user_memory_state(user_id: int):
    key = f"ai_user_memory_{user_id}"
    if key not in st.session_state:
        st.session_state[key] = {
            "preferences": {},
            "last_entities": {},
            "last_intent": "general"
        }
    return st.session_state[key]


def update_user_memory_state(user_id: int, analysis: dict):
    memory = get_user_memory_state(user_id)
    memory["preferences"] = _merge_preferences(memory.get("preferences", {}), analysis.get("preferences", {}))
    memory["last_entities"] = analysis.get("entities", {})
    memory["last_intent"] = analysis.get("intent", "general")
    st.session_state[f"ai_user_memory_{user_id}"] = memory
    return memory


def _build_token_weights(tokens: list[str]) -> Counter:
    counter = Counter(tokens)
    for tok in list(counter.keys()):
        if len(tok) >= 5:
            counter[tok] += 1
    return counter


def _score_document_semantic(doc: dict, analysis: dict, memory: dict | None = None) -> float:
    content_norm = normalize_query_text(doc.get("content", ""))
    title_norm = normalize_query_text(doc.get("title", ""))
    metadata = doc.get("metadata", {}) or {}

    token_weights = _build_token_weights(analysis.get("tokens", []))
    score = 0.0

    for tok, weight in token_weights.items():
        if tok in title_norm:
            score += 4.5 * weight
        elif tok in content_norm:
            score += 2.2 * weight

    entities = analysis.get("entities", {})
    for place in entities.get("places", []):
        place_norm = normalize_query_text(place)
        if place_norm and (place_norm in title_norm or place_norm in content_norm):
            score += 9

    for area in entities.get("areas", []):
        area_norm = normalize_query_text(area)
        if area_norm and area_norm in normalize_query_text(metadata.get("area", "")):
            score += 6

    for route_to in entities.get("route_to", []):
        route_norm = normalize_query_text(route_to)
        if route_norm and route_norm in normalize_query_text(metadata.get("to", "")):
            score += 8

    prefs = {}
    prefs.update(memory.get("preferences", {}) if isinstance(memory, dict) else {})
    prefs.update(analysis.get("preferences", {}))

    if prefs.get("culture") and doc.get("doc_type") in ["le_hoi", "lang_nghe", "cong_dong"]:
        score += 4.5
    if prefs.get("history") and any(k in content_norm for k in ["di tich", "lich su", "den", "chua", "tam linh"]):
        score += 4
    if prefs.get("nature") and any(k in content_norm for k in ["thien nhien", "san may", "nui", "rung", "thac", "canh quan"]):
        score += 4
    if prefs.get("relax") and any(k in content_norm for k in ["yen tinh", "thu gian", "nghi duong"]):
        score += 3
    if prefs.get("adventure") and any(k in content_norm for k in ["trekking", "leo nui", "kham pha"]):
        score += 3

    if prefs.get("budget_low") and any(k in content_norm for k in ["re", "tiet kiem", "hop ly", "0d"]):
        score += 2

    intent = analysis.get("intent", "general")
    if intent == "chi_phi" and doc.get("doc_type") == "lich_trinh":
        score += 5
    elif intent == "lich_trinh" and doc.get("doc_type") == "lich_trinh":
        score += 5
    elif intent == "van_hoa" and doc.get("doc_type") in ["le_hoi", "lang_nghe", "diem_den", "cong_dong"]:
        score += 4
    elif intent == "thoi_diem" and any(k in content_norm for k in ["mua", "thoi diem", "best time", "season"]):
        score += 4
    elif intent == "goi_y":
        score += 2

    return score


def retrieve_relevant_documents(question: str, payload: dict, analysis: dict, memory: dict | None = None, top_k: int = 8):
    ranked = []
    for doc in payload.get("documents", []):
        score = _score_document_semantic(doc, analysis=analysis, memory=memory)
        if score > 0:
            ranked.append((score, doc))

    ranked.sort(key=lambda x: x[0], reverse=True)
    return [doc for score, doc in ranked[:top_k]]

def _infer_place_features(place: dict) -> dict:
    text = normalize_query_text(_safe_join_text([
        place.get("name", ""),
        place.get("area", ""),
        place.get("category", ""),
        place.get("season", ""),
        place.get("best_time", ""),
        place.get("short_desc", ""),
        place.get("full_desc", ""),
        " ".join(place.get("highlights", [])),
    ]))

    features = {
        "nature": any(x in text for x in ["thien nhien", "san may", "nui", "rung", "thac", "ruong bac thang", "canh dep"]),
        "culture": any(x in text for x in ["van hoa", "dan toc", "ban sac", "cho phien", "le hoi", "lang nghe"]),
        "history": any(x in text for x in ["di tich", "lich su", "den", "chua", "tam linh"]),
        "relax": any(x in text for x in ["yen tinh", "thu gian", "nghi duong"]),
        "adventure": any(x in text for x in ["trekking", "leo nui", "kham pha", "mao hiem"]),
        "food": any(x in text for x in ["am thuc", "dac san", "mon ngon", "cho"]),
        "budget_low": any(x in text for x in ["re", "tiet kiem", "hop ly"]),
        "weekend_fit": True,
    }

    return features


def _score_place_for_recommendation(place: dict, analysis: dict, memory: dict | None = None) -> tuple[float, list[str]]:
    prefs = {}
    if isinstance(memory, dict):
        prefs.update(memory.get("preferences", {}) or {})
    prefs.update(analysis.get("preferences", {}) or {})

    features = _infer_place_features(place)
    score = 0.0
    reasons = []

    if prefs.get("nature") and features["nature"]:
        score += 3
        reasons.append("hợp sở thích thiên nhiên")

    if prefs.get("culture") and features["culture"]:
        score += 3
        reasons.append("có giá trị văn hóa bản địa")

    if prefs.get("history") and features["history"]:
        score += 3
        reasons.append("phù hợp nhu cầu tìm hiểu lịch sử - tâm linh")

    if prefs.get("relax") and features["relax"]:
        score += 2.5
        reasons.append("phù hợp nghỉ ngơi, thư giãn")

    if prefs.get("adventure") and features["adventure"]:
        score += 2.5
        reasons.append("phù hợp trải nghiệm khám phá")

    if prefs.get("food_focus") and features["food"]:
        score += 2
        reasons.append("có thể kết hợp trải nghiệm ẩm thực")

    if prefs.get("budget_low") and features["budget_low"]:
        score += 2
        reasons.append("phù hợp ngân sách tiết kiệm")

    if prefs.get("trip_time") == "cuoi_tuan":
        score += 1.5
        reasons.append("phù hợp chuyến đi cuối tuần")

    if prefs.get("duration") in ["1 ngày", "2 ngày 1 đêm", "3 ngày 2 đêm"]:
        score += 1.5
        reasons.append(f"hợp thời lượng {prefs.get('duration')}")

    if prefs.get("companion_type") == "gia_dinh":
        score += 1.2
        reasons.append("dễ đi cùng gia đình")
    elif prefs.get("companion_type") == "ban_be":
        score += 1.2
        reasons.append("phù hợp đi cùng nhóm bạn")
    elif prefs.get("companion_type") == "cap_doi":
        score += 1.2
        reasons.append("phù hợp cho cặp đôi")

    if prefs.get("crowd_preference") == "thap" and (features["relax"] or features["nature"]):
        score += 1.5
        reasons.append("phù hợp nhu cầu tránh nơi đông đúc")

    if prefs.get("season"):
        place_season = normalize_query_text(place.get("season", "") + " " + place.get("best_time", ""))
        if normalize_query_text(prefs["season"]) in place_season:
            score += 2
            reasons.append(f"phù hợp {prefs['season']}")

    if analysis.get("entities", {}).get("areas"):
        place_area_norm = normalize_query_text(place.get("area", ""))
        for area in analysis["entities"]["areas"]:
            if normalize_query_text(area) in place_area_norm:
                score += 2
                reasons.append(f"đúng khu vực {area}")
                break

    return score, reasons


def recommend_places_for_user(payload: dict, analysis: dict, memory: dict | None = None, top_k: int = 3):
    ranked = []

    for place in payload.get("places", []):
        score, reasons = _score_place_for_recommendation(place, analysis, memory)
        if score > 0:
            ranked.append({
                "score": score,
                "place": place,
                "reasons": reasons[:4]
            })

    ranked.sort(key=lambda x: x["score"], reverse=True)
    return ranked[:top_k]

def build_personalized_recommendation_text(payload: dict, analysis: dict, memory: dict | None = None) -> str:
    recommendations = recommend_places_for_user(payload, analysis, memory, top_k=3)

    if not recommendations:
        return ""

    lines = []
    lines.append("Gợi ý cá nhân hóa phù hợp với nhu cầu của bạn:")
    lines.append("")

    for idx, item in enumerate(recommendations, 1):
        place = item["place"]
        reasons = item["reasons"]
        short_desc = str(place.get("short_desc", "")).strip()

        lines.append(f"{idx}. {place.get('name', 'Địa điểm')}")
        if place.get("area"):
            lines.append(f"   - Khu vực: {place['area']}")
        if short_desc:
            lines.append(f"   - Mô tả: {short_desc}")
        if reasons:
            lines.append(f"   - Lý do phù hợp: {', '.join(reasons)}")
        lines.append("")

    lines.append("Bạn có thể hỏi tiếp như: 'gợi ý lịch trình chi tiết cho lựa chọn số 1' hoặc 'địa điểm nào phù hợp nhất cho cuối tuần này?'")
    return "\n".join(lines).strip()

def _format_preferences_for_prompt(memory: dict | None, analysis: dict | None) -> str:
    prefs = {}
    if isinstance(memory, dict):
        prefs.update(memory.get("preferences", {}) or {})
    if isinstance(analysis, dict):
        prefs.update(analysis.get("preferences", {}) or {})

    if not prefs:
        return "Chưa có sở thích/ngữ cảnh nổi bật được ghi nhớ."

    lines = []
    mapping = {
        "duration": "Thời lượng mong muốn",
        "season": "Mùa quan tâm",
        "budget_low": "Ngân sách thấp",
        "budget_mid": "Ngân sách vừa phải",
        "budget_high": "Ngân sách cao",
        "nature": "Ưu tiên thiên nhiên",
        "culture": "Ưu tiên văn hóa bản địa",
        "history": "Ưu tiên lịch sử/tâm linh",
        "relax": "Ưu tiên yên tĩnh/thư giãn",
        "adventure": "Ưu tiên khám phá/mạo hiểm",
        "companion_type": "Đối tượng đi cùng",
        "crowd_preference": "Mức độ đông đúc mong muốn",
        "weather_preference": "Thời tiết mong muốn",
        "trip_time": "Thời điểm chuyến đi",
        "nearby_preference": "Ưu tiên địa điểm gần",
        "food_focus": "Ưu tiên ẩm thực",
    }

    pretty_value_map = {
        "gia_dinh": "Gia đình",
        "ban_be": "Bạn bè / nhóm",
        "cap_doi": "Cặp đôi",
        "mot_minh": "Một mình",
        "thap": "Ít đông / yên tĩnh",
        "cao": "Đông vui / sôi động",
        "cuoi_tuan": "Cuối tuần",
        "ngay_thuong": "Ngày thường",
        "dip_le": "Dịp lễ",
        "mat": "Mát mẻ",
        "nang_dep": "Nắng đẹp",
        "kho_rao": "Khô ráo",
    }

    for key, value in prefs.items():
        label = mapping.get(key, key)
        if value is True:
            lines.append(f"- {label}")
        else:
            pretty_value = pretty_value_map.get(str(value), value)
            lines.append(f"- {label}: {pretty_value}")

    return "\n".join(lines)


def build_context_from_payload(question: str, payload: dict, current_page: str, analysis: dict | None = None, memory: dict | None = None) -> str:
    analysis = analysis or analyze_user_query(question, payload)
    relevant_docs = retrieve_relevant_documents(
        question=question,
        payload=payload,
        analysis=analysis,
        memory=memory,
        top_k=8,
    )

    lines = []
    lines.append(f"Trang hiện tại: {current_page}")
    lines.append(f"Tổng số địa điểm trong hệ thống: {len(payload.get('places', []))}")
    lines.append(f"Tổng số lịch trình trong hệ thống: {len(payload.get('routes', []))}")
    lines.append(f"Tổng số tư liệu AI: {len(payload.get('documents', []))}")
    lines.append("")

    lines.append("PHÂN TÍCH CÂU HỎI:")
    lines.append(f"- Intent: {analysis.get('intent', 'general')}")
    lines.append(f"- Địa điểm nhận diện: {', '.join(analysis.get('entities', {}).get('places', [])) or 'Không rõ'}")
    lines.append(f"- Khu vực nhận diện: {', '.join(analysis.get('entities', {}).get('areas', [])) or 'Không rõ'}")
    lines.append(f"- Điểm đến lịch trình: {', '.join(analysis.get('entities', {}).get('route_to', [])) or 'Không rõ'}")
    lines.append("")

    lines.append("NGỮ CẢNH NGƯỜI DÙNG ĐANG ĐƯỢC GHI NHỚ:")
    lines.append(_format_preferences_for_prompt(memory, analysis))
    lines.append("")

    if relevant_docs:
        lines.append("TÀI LIỆU LIÊN QUAN THEO NGỮ NGHĨA:")
        for i, doc in enumerate(relevant_docs, 1):
            metadata = doc.get("metadata", {}) or {}
            content = str(doc.get("content", "")).strip()
            short_content = content[:320].rsplit(" ", 1)[0] + "..." if len(content) > 320 else content
            lines.append(
                f"{i}. Loại: {doc.get('doc_type', 'khac')} | "
                f"Tiêu đề: {doc.get('title', 'Đang cập nhật')} | "
                f"Khu vực: {metadata.get('area', metadata.get('to', 'Đang cập nhật'))} | "
                f"Danh mục: {metadata.get('category', 'Đang cập nhật')} | "
                f"Nội dung: {short_content}"
            )
    else:
        lines.append("Không tìm thấy tài liệu nội bộ đủ gần theo ngữ nghĩa.")
        lines.append("Khi trả lời, hãy nói rõ dữ liệu hệ thống chưa đủ cụ thể.")

    return "\n".join(lines)


def build_history_text(messages: list[dict], limit: int = 8) -> str:
    recent = messages[-limit:] if messages else []
    lines = []

    for msg in recent:
        role = "Người dùng" if msg.get("role") == "user" else "Trợ lý"
        content = str(msg.get("content", "")).strip()
        if content:
            lines.append(f"{role}: {content}")

    return "\n".join(lines) if lines else "Chưa có lịch sử hội thoại."


def route_question(user_message: str, analysis: dict | None = None) -> str:
    analysis = analysis or analyze_user_query(user_message, CHATBOT_PAYLOAD)
    q = analysis.get("normalized_question", normalize_query_text(user_message))
    intent = analysis.get("intent", "general")

    search_keywords = [
        "hom nay", "hien tai", "moi nhat", "gan day", "tin tuc", "thoi tiet",
        "su kien", "gia ve", "gio mo cua", "lich mo cua", "cap nhat",
        "nam nay", "thang nay", "tuan nay"
    ]

    external_keywords = [
        "o dau", "dia chi", "di chuyen", "duong di", "review", "danh gia",
        "an gi", "khach san nao", "quan an", "cafe", "gan", "xung quanh"
    ]

    if intent == "need_web":
        return "search"

    if any(k in q for k in search_keywords):
        return "search"

    if any(k in q for k in external_keywords):
        return "search"

    if analysis.get("entities", {}).get("places") or analysis.get("entities", {}).get("route_to"):
        return "internal"

    if intent in ["goi_y", "chi_phi", "lich_trinh", "thoi_diem", "van_hoa"]:
        return "internal"

    return "hybrid"


def _build_gemini_prompt(user_message: str, current_page: str, context_text: str, history_text: str) -> str:
    return f"""
DỮ LIỆU NỘI BỘ:
{context_text}

LỊCH SỬ HỘI THOẠI GẦN ĐÂY:
{history_text}

CÂU HỎI MỚI CỦA NGƯỜI DÙNG:
{user_message}

YÊU CẦU TRẢ LỜI:
- Ưu tiên dùng dữ liệu nội bộ và ngữ cảnh người dùng đã được ghi nhớ.
- Không trả lời kiểu khớp từ khóa máy móc; hãy hiểu theo ý định và nhu cầu thực tế.
- Nếu dữ liệu hệ thống chưa đủ, nói rõ là chưa có đủ trong hệ thống.
- Trả lời bằng tiếng Việt, tự nhiên, rõ ràng, hữu ích.
""".strip()


def ask_gemini(user_message: str, payload: dict, current_page: str, messages: list[dict]) -> str:
    api_key = get_gemini_api_key()
    if not api_key:
        return "Bạn chưa cấu hình GEMINI_API_KEY."

    current_user = st.session_state.get("auth_user", {}) or {}
    current_user_id = current_user.get("id")

    analysis = analyze_user_query(user_message, payload, messages)
    memory_before = get_user_memory_state(current_user_id) if current_user_id else {"preferences": {}}
    route_mode = route_question(user_message, analysis=analysis)

    personalized_hint = ""
    if analysis.get("intent") == "goi_y":
        personalized_hint = build_personalized_recommendation_text(
            payload=payload,
            analysis=analysis,
            memory=memory_before
    )

    try:
        client = genai.Client(api_key=api_key)

        context_text = build_context_from_payload(
            question=user_message,
            payload=payload,
            current_page=current_page,
            analysis=analysis,
            memory=memory_before,
        )
        history_text = build_history_text(messages, limit=8)
        prompt = _build_gemini_prompt(user_message, current_page, context_text, history_text)

        if personalized_hint:
            prompt += "\n\nGỢI Ý CÁ NHÂN HÓA NỘI BỘ:\n" + personalized_hint

        config_kwargs = {
            "system_instruction": (
                "Bạn là Lao Cai Heritage AI, trợ lý du lịch thông minh cho website giới thiệu "
                "địa điểm, lịch trình và văn hóa Lào Cai. "
                "Bạn phải ưu tiên hiểu ý định câu hỏi, sở thích người dùng và dữ liệu nội bộ theo ngữ nghĩa. "
                "Không bịa thông tin. Nếu dữ liệu hệ thống chưa có, hãy nói rõ là chưa có trong hệ thống. "
                "Trả lời bằng tiếng Việt, ngắn gọn, đúng trọng tâm, dễ hiểu. "
                "Ưu tiên trả lời theo dạng tư vấn thực tế cho khách du lịch. "
                "Khi phù hợp, gợi ý thêm 1 đến 2 câu hỏi tiếp theo. "
                "Không nói rằng bạn là Gemini hay Google."
            ),
            "temperature": 0.35,
            "max_output_tokens": 650,
        }

        if route_mode in ["search", "hybrid"]:
            grounding_tool = types.Tool(google_search=types.GoogleSearch())
            config_kwargs["tools"] = [grounding_tool]

        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(**config_kwargs),
        )

        answer = (response.text or "").strip()
        if answer:
            if current_user_id:
                update_user_memory_state(current_user_id, analysis)
            return answer

        return "Mình chưa tạo được câu trả lời phù hợp. Bạn hãy hỏi lại ngắn gọn hơn."

    except Exception as e:
        error_text = str(e)

        if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text.upper():
            return "Gemini đang báo hết quota hoặc vượt giới hạn tốc độ ở API key hiện tại."

        if "API key" in error_text or "api_key" in error_text.lower():
            return "API key Gemini chưa đúng hoặc chưa được cấp quyền."

        return f"Lỗi Gemini: {error_text}"


hero_home_src = image_to_data_uri([
    "kho_anh/chinh.png",
    "kho_anh/chinh.jpg",
    "kho_anh/chinh.jpeg",
    "kho_anh/chinh.webp"
])

# CSS trang chủ
st.markdown(f"""
<style>
/* Ẩn menu mặc định streamlit cho giống web thật hơn */
#MainMenu {{visibility: hidden;}}
footer {{visibility: hidden;}}
header {{visibility: hidden;}}

/* Toàn bộ nền */
.stApp {{
    background-color: #f6f6f6;
}}

/* Giảm padding mặc định */
.block-container {{
    padding-top: 45px;
    padding-left: 0rem;
    padding-right: 0rem;
    max-width: 100%;
}}

/* Navbar */
.navbar {{
    background: rgba(255,255,255,0.98);
    height: 72px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 0 60px;
    box-shadow: 0 2px 12px rgba(0,0,0,0.08);
    position: fixed;
    top: 0;
    left: 0;
    right: 0;
    width: 100%;
    z-index: 99999;
    box-sizing: border-box;
    backdrop-filter: blur(8px);
    -webkit-backdrop-filter: blur(8px);
}}

.logo {{
    display: flex;
    align-items: center;
    gap: 14px;
    text-decoration: none;
    white-space: nowrap;
}}

.logo img {{
    height: 71px;
    width: auto;
    display: block;
    object-fit: contain;
}}

.logo-text {{
    font-size: 34px;
    font-weight: 800;
    color: #1565c0;
    line-height: 1;
}}

.logo-text span {{
    color: #2e7d32;
}}

.nav-links {{
    display: flex;
    gap: 34px;
    align-items: center;
    flex-wrap: wrap;
}}

.nav-links a {{
    font-size: 18px;
    font-weight: 600;
    color: #222;
    text-decoration: none;
    position: relative;
    padding: 6px 0;
    transition: color 0.25s ease;
}}

.nav-links a:hover {{
    color: #1565c0;
}}

.nav-links a.active {{
    color: #1565c0;
}}

.nav-links a.active::after {{
    content: "";
    position: absolute;
    left: 0;
    bottom: -6px;
    width: 100%;
    height: 3px;
    border-radius: 999px;
    background: #1565c0;
}}
            
.nav-user{{
    display:inline-flex;
    align-items:center;
    gap:8px;
    padding:8px 14px !important;
    border-radius:999px;
    background:#f1f5f9;
    color:#0f172a !important;
    border:1px solid #dbe4ef;
    font-weight:800 !important;
}}

.nav-user:hover{{
    background:#e2e8f0;
    color:#0f172a !important;
}}

.nav-user::after{{
    display:none !important;
}}

/* Hero */
.hero {{
    position: relative;
    min-height: 100vh;
    background-image: linear-gradient(rgba(0,0,0,0.28), rgba(0,0,0,0.28)),
                      url('{hero_home_src}');
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
    display: flex;
    align-items: center;
    padding: 0 80px;
    box-sizing: border-box;
    overflow: hidden;
}}

.hero-content {{
    color: white;
    max-width: 760px;
    width: 100%;
    position: relative;
    z-index: 2;
}}

.hero-title {{
    font-size: 54px;
    font-weight: 800;
    line-height: 1.15;
    margin-bottom: 16px;
}}

.hero-subtitle {{
    font-size: 22px;
    line-height: 1.6;
    opacity: 0.96;
}}

@media (max-width: 768px) {{
    .hero {{
        min-height: 420px !important;
        padding: 88px 16px 28px 16px !important;
        align-items: center !important;
        background-position: center center !important;
    }}

    .hero-content {{
        max-width: 100% !important;
        text-align: center !important;
    }}

    .hero-title {{
        font-size: 28px !important;
        line-height: 1.28 !important;
        margin-bottom: 12px !important;
    }}

    .hero-subtitle {{
        font-size: 14px !important;
        line-height: 1.7 !important;
        opacity: 0.96 !important;
    }}
}}

.hero-content {{
    color: white;
    max-width: 700px;
    animation: heroFadeUp 1s ease-out forwards;
}}

.hero-title {{
    font-size: 54px;
    font-weight: 800;
    line-height: 1.15;
    margin-bottom: 16px;
    opacity: 0;
    transform: translateY(28px);
    animation: heroTextUp 1s ease-out forwards;
}}

.hero-subtitle {{
    font-size: 22px;
    line-height: 1.6;
    opacity: 0;
    transform: translateY(30px);
    animation: heroTextUp 1s ease-out 0.35s forwards;
}}

@keyframes heroZoom {{
    from {{
        background-size: 100%;
    }}
    to {{
        background-size: 108%;
    }}
}}

@keyframes heroTextUp {{
    from {{
        opacity: 0;
        transform: translateY(28px);
    }}
    to {{
        opacity: 1;
        transform: translateY(0);
    }}
}}

@keyframes heroFadeUp {{
    from {{
        opacity: 0.2;
    }}
    to {{
        opacity: 1;
    }}
}}

.hero-content {{
    color: white;
    max-width: 700px;
}}

.hero-title {{
    font-size: 54px;
    font-weight: 800;
    line-height: 1.15;
    margin-bottom: 16px;
}}

.hero-subtitle {{
    font-size: 22px;
    line-height: 1.6;
    opacity: 0.96;
}}       

@media (max-width: 992px) {{
    .search-wrapper {{
        margin-top: -52px;
    }}

    .search-box {{
        width: 92%;
    }}

    .search-row {{
        grid-template-columns: 1fr 1fr;
    }}
}}

@media (max-width: 640px) {{
    .search-wrapper {{
        margin-top: -36px;
    }}

    .search-box {{
        width: 94%;
        padding: 14px;
    }}

    .search-row {{
        grid-template-columns: 1fr;
    }}
}}            

.search-float {{
    position: relative;
    margin-top: -95px;  
    margin-bottom: -20px;
    z-index: 50;
}}         

/* Danh mục nhanh */
.section {{
    padding: 46px 70px 10px 70px;
}}

.quick-grid {{
    display: grid;
    grid-template-columns: repeat(5, 1fr);
    gap: 24px;
    margin-top: 20px;
}}

.quick-card {{
    background: white;
    border-radius: 20px;
    padding: 26px 16px;
    text-align: center;
    box-shadow: 0 6px 18px rgba(0,0,0,0.06);
    transition: 0.2s;
}}

.quick-card:hover {{
    transform: translateY(-4px);
}}

.quick-icon {{
    font-size: 34px;
    margin-bottom: 10px;
}}

.quick-title {{
    font-size: 18px;
    font-weight: 700;
    color: #1f2937;
}}

/* Tiêu đề section */
.section-title {{
    font-size: 34px;
    font-weight: 800;
    color: #111827;
    margin-bottom: 8px;
}}

.section-subtitle {{
    font-size: 18px;
    color: #6b7280;
}}

/* Card địa điểm */
.place-grid {{
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 26px;
    margin-top: 28px;
    margin-bottom: 60px;
}}

.place-card {{
    background: white;
    border-radius: 22px;
    overflow: hidden;
    box-shadow: 0 8px 20px rgba(0,0,0,0.08);
}}

.place-image {{
    height: 220px;
    background-size: cover;
    background-position: center;
}}

.place-body {{
    padding: 18px;
}}

.place-name {{
    font-size: 22px;
    font-weight: 800;
    color: #111827;
    margin-bottom: 10px;
}}

.place-desc {{
    color: #4b5563;
    line-height: 1.6;
    font-size: 16px;
}}

/* Footer */
.footer {{
    background: #dfeefc;
    padding: 40px 70px;
    margin-top: 30px;
}}

.footer-title {{
    font-size: 24px;
    font-weight: 800;
    margin-bottom: 14px;
    color: #0d2b4d;
}}

.footer-text {{
    font-size: 17px;
    color: #334155;
    line-height: 1.8;
}}

@media (max-width: 768px) {{
    html, body {{
        overflow-x: hidden !important;
    }}

    .stApp {{
        overflow-x: hidden !important;
    }}

    .block-container {{
        padding-top: 78px !important;
        padding-left: 0 !important;
        padding-right: 0 !important;
        max-width: 100% !important;
        overflow-x: hidden !important;
    }}

    .navbar {{
        height: auto !important;
        min-height: 64px !important;
        padding: 10px 14px !important;
        display: flex !important;
        flex-direction: column !important;
        align-items: flex-start !important;
        justify-content: center !important;
        gap: 8px !important;
    }}

    .logo {{
        width: 100% !important;
        min-width: 0 !important;
        gap: 10px !important;
        align-items: center !important;
    }}

    .logo img {{
        height: 40px !important;
        width: auto !important;
        flex-shrink: 0 !important;
    }}

    .logo-text {{
        font-size: 18px !important;
        line-height: 1.2 !important;
        white-space: normal !important;
        word-break: break-word !important;
    }}

    .nav-links {{
        width: 100% !important;
        display: flex !important;
        flex-wrap: wrap !important;
        gap: 8px 14px !important;
        align-items: center !important;
    }}

    .nav-links a {{
        font-size: 14px !important;
        padding: 2px 0 !important;
        line-height: 1.3 !important;
        white-space: nowrap !important;
    }}

    .nav-links a.active::after {{
        bottom: -3px !important;
        height: 2px !important;
    }}
}}
</style>
""", unsafe_allow_html=True)

def render_auth_page():
    auth_bg_src = image_to_data_uri([
        "kho_anh/trang_chu/chinh.png",
        "kho_anh/trang_chu/chinh.jpg",
        "kho_anh/trang_chu/chinh.jpeg",
        "kho_anh/trang_chu/chinh.webp",
        "kho_anh/chinh.png",
        "kho_anh/chinh.jpg",
        "kho_anh/chinh.jpeg",
        "kho_anh/chinh.webp"
    ])

    bg_css = f"background-image: url('{auth_bg_src}');" if auth_bg_src else """
    background:
        radial-gradient(circle at 20% 20%, rgba(59,255,136,0.18), transparent 26%),
        radial-gradient(circle at 80% 30%, rgba(0,255,170,0.14), transparent 25%),
        radial-gradient(circle at 50% 78%, rgba(0,200,255,0.12), transparent 28%),
        linear-gradient(135deg, #03140f 0%, #0a251d 55%, #071814 100%);
    """

    login_url = build_app_url("login")
    signup_url = build_app_url("signup")

    st.markdown(f"""
    <style>
    .block-container {{
        padding-top: 0 !important;
        padding-left: 0 !important;
        padding-right: 0 !important;
        max-width: 100% !important;
    }}

    .auth-screen {{
        min-height: 100vh;
        width: 100%;
        {bg_css}
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        position: relative;
        overflow: hidden;
    }}

    .auth-screen::before {{
        content: "";
        position: absolute;
        inset: 0;
        background: linear-gradient(180deg, rgba(0,0,0,0.18) 0%, rgba(0,0,0,0.38) 100%);
        backdrop-filter: blur(5px);
        -webkit-backdrop-filter: blur(5px);
    }}

    .auth-center {{
        position: relative;
        z-index: 2;
        min-height: 100vh;
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 20px;
        box-sizing: border-box;
    }}

    .auth-box {{
        width: 100%;
        max-width: 320px;
        padding: 28px 20px 20px 20px;
        border-radius: 22px;
        background: rgba(8, 36, 28, 0.38);
        border: 1px solid rgba(255,255,255,0.16);
        box-shadow: 0 18px 40px rgba(0,0,0,0.30);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
    }}

    .auth-title {{
        font-size: 18px;
        font-weight: 800;
        color: #ffffff;
        margin-bottom: 10px;
        line-height: 1.2;
    }}

    .auth-sub {{
        font-size: 12px;
        line-height: 1.6;
        color: rgba(255,255,255,0.78);
        margin-bottom: 18px;
    }}

    .auth-btn {{
        display: flex;
        align-items: center;
        justify-content: center;
        width: 100%;
        height: 44px;
        border-radius: 999px;
        text-decoration: none !important;
        font-size: 13px;
        font-weight: 800;
        letter-spacing: 0.03em;
        box-sizing: border-box;
        transition: all 0.18s ease;
    }}

    .auth-btn-signin {{
        background: linear-gradient(180deg, #6ad638 0%, #56c92a 100%);
        color: #ffffff !important;
        margin-bottom: 12px;
        box-shadow: 0 10px 22px rgba(95, 214, 42, 0.28);
    }}

    .auth-btn-signin:hover {{
        transform: translateY(-1px);
    }}

    .auth-btn-signup {{
        background: transparent;
        color: #ffffff !important;
        border: 1px solid rgba(255,255,255,0.45);
    }}

    .auth-btn-signup:hover {{
        background: rgba(255,255,255,0.05);
        transform: translateY(-1px);
    }}
    </style>

    <div class="auth-screen">
        <div class="auth-center">
            <div class="auth-box">
                <div class="auth-title">Hello!</div>
                <div class="auth-sub">
                    Đăng nhập hoặc tạo tài khoản để lưu lịch sử sử dụng theo từng người dùng.
                </div>
                <a class="auth-btn auth-btn-signin" href="{login_url}" target="_self">SIGN IN</a>
                <a class="auth-btn auth-btn-signup" href="{signup_url}" target="_self">SIGN UP</a>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_login_form_page():
    auth_bg_src = image_to_data_uri([
        "kho_anh/trang_chu/chinh.png",
        "kho_anh/trang_chu/chinh.jpg",
        "kho_anh/trang_chu/chinh.jpeg",
        "kho_anh/trang_chu/chinh.webp",
        "kho_anh/chinh.png",
        "kho_anh/chinh.jpg",
        "kho_anh/chinh.jpeg",
        "kho_anh/chinh.webp"
    ])

    bg_css = f"background-image: url('{auth_bg_src}');" if auth_bg_src else """
    background:
        radial-gradient(circle at 20% 20%, rgba(59,255,136,0.18), transparent 26%),
        radial-gradient(circle at 80% 30%, rgba(0,255,170,0.14), transparent 25%),
        radial-gradient(circle at 50% 78%, rgba(0,200,255,0.12), transparent 28%),
        linear-gradient(135deg, #03140f 0%, #0a251d 55%, #071814 100%);
    """

    st.markdown(f"""
    <style>
    .stApp {{
        {bg_css}
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        min-height: 100vh;
    }}

    .stApp::before {{
        content: "";
        position: fixed;
        inset: 0;
        background: rgba(0,0,0,0.28);
        backdrop-filter: blur(5px);
        -webkit-backdrop-filter: blur(5px);
        z-index: 0;
        pointer-events: none;
    }}

    .block-container {{
        position: relative;
        z-index: 1;
        padding-top: 40px !important;
        padding-bottom: 40px !important;
        max-width: 100% !important;
    }}

    .auth-card-wrap {{
        padding: 10px;
        border-radius: 24px;
        background: rgba(8, 36, 28, 0.38);
        border: 1px solid rgba(255,255,255,0.16);
        box-shadow: 0 18px 40px rgba(0,0,0,0.30);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
    }}

    .auth-back {{
        display: inline-block;
        margin-bottom: 12px;
        color: rgba(255,255,255,0.82) !important;
        text-decoration: none !important;
        font-size: 14px;
        font-weight: 700;
    }}

    .auth-title {{
        font-size: 22px;
        font-weight: 800;
        color: #ffffff;
        margin-bottom: 8px;
    }}

    .auth-sub {{
        font-size: 15px;
        line-height: 1.6;
        color: rgba(255,255,255,0.82);
        margin-bottom: 16px;
    }}

    div[data-testid="stForm"] {{
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
    }}

    div[data-testid="stTextInput"] label {{
        display: none !important;
    }}

    div[data-testid="stTextInput"] input {{
        height: 44px !important;
        min-height: 44px !important;
        border-radius: 999px !important;
        border: 1.2px solid rgba(255,255,255,0.52) !important;
        background: rgba(255,255,255,0.90) !important;
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        caret-color: #000000 !important;
        font-size: 14px !important;
        padding: 0 16px !important;
        box-shadow: none !important;
    }}

    div[data-testid="stTextInput"] input::placeholder {{
        color: rgba(0,0,0,0.55) !important;
        -webkit-text-fill-color: rgba(0,0,0,0.55) !important;
    }}

    div[data-testid="stFormSubmitButton"] button {{
        width: 100% !important;
        height: 44px !important;
        min-height: 44px !important;
        border-radius: 999px !important;
        border: none !important;
        background: linear-gradient(180deg, #6ad638 0%, #56c92a 100%) !important;
        color: #ffffff !important;
        font-size: 14px !important;
        font-weight: 800 !important;
        letter-spacing: 0.03em !important;
    }}

    div[data-testid="stAlert"] {{
        border-radius: 14px !important;
    }}
    </style>
    """, unsafe_allow_html=True)

    left, center, right = st.columns([1.2, 1.8, 1.2])

    with center:
        st.markdown('<div class="auth-card-wrap">', unsafe_allow_html=True)
        st.markdown(
            f'<a class="auth-back" href="{build_app_url("home")}" target="_self">← Quay lại</a>',
            unsafe_allow_html=True
        )
        st.markdown('<div class="auth-title">Sign in</div>', unsafe_allow_html=True)
        st.markdown('<div class="auth-sub">Đăng nhập để tiếp tục sử dụng hệ thống.</div>', unsafe_allow_html=True)

        with st.form("login_form", clear_on_submit=False):
            login_username = st.text_input(
                "Tên đăng nhập",
                placeholder="Tên đăng nhập",
                key="login_username"
            )
            login_password = st.text_input(
                "Mật khẩu",
                type="password",
                placeholder="Mật khẩu",
                key="login_password"
            )
            login_submit = st.form_submit_button("SIGN IN", use_container_width=True)

        if login_submit:
            user = verify_user(login_username, login_password)
            if not user:
                st.error("Sai tên đăng nhập hoặc mật khẩu.")
            else:
                sid = create_session(user["id"])
                st.session_state.auth_user = user
                st.session_state.auth_sid = sid
                st.query_params.clear()
                st.query_params["page"] = "home"
                st.query_params["sid"] = sid
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)


def render_signup_form_page():
    auth_bg_src = image_to_data_uri([
        "kho_anh/trang_chu/chinh.png",
        "kho_anh/trang_chu/chinh.jpg",
        "kho_anh/trang_chu/chinh.jpeg",
        "kho_anh/trang_chu/chinh.webp",
        "kho_anh/chinh.png",
        "kho_anh/chinh.jpg",
        "kho_anh/chinh.jpeg",
        "kho_anh/chinh.webp"
    ])

    bg_css = f"background-image: url('{auth_bg_src}');" if auth_bg_src else """
    background:
        radial-gradient(circle at 20% 20%, rgba(59,255,136,0.18), transparent 26%),
        radial-gradient(circle at 80% 30%, rgba(0,255,170,0.14), transparent 25%),
        radial-gradient(circle at 50% 78%, rgba(0,200,255,0.12), transparent 28%),
        linear-gradient(135deg, #03140f 0%, #0a251d 55%, #071814 100%);
    """

    st.markdown(f"""
    <style>
    .stApp {{
        {bg_css}
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        min-height: 100vh;
    }}

    .stApp::before {{
        content: "";
        position: fixed;
        inset: 0;
        background: rgba(0,0,0,0.28);
        backdrop-filter: blur(5px);
        -webkit-backdrop-filter: blur(5px);
        z-index: 0;
        pointer-events: none;
    }}

    .block-container {{
        position: relative;
        z-index: 1;
        padding-top: 40px !important;
        padding-bottom: 40px !important;
        max-width: 100% !important;
    }}

    .auth-card-wrap {{
        padding: 10px;
        border-radius: 24px;
        background: rgba(8, 36, 28, 0.38);
        border: 1px solid rgba(255,255,255,0.16);
        box-shadow: 0 18px 40px rgba(0,0,0,0.30);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
    }}

    .auth-back {{
        display: inline-block;
        margin-bottom: 12px;
        color: rgba(255,255,255,0.82) !important;
        text-decoration: none !important;
        font-size: 14px;
        font-weight: 700;
    }}

    .auth-title {{
        font-size: 22px;
        font-weight: 800;
        color: #ffffff;
        margin-bottom: 8px;
    }}

    .auth-sub {{
        font-size: 15px;
        line-height: 1.6;
        color: rgba(255,255,255,0.82);
        margin-bottom: 16px;
    }}

    div[data-testid="stForm"] {{
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
    }}

    div[data-testid="stTextInput"] label {{
        display: none !important;
    }}

    div[data-testid="stTextInput"] input {{
        height: 44px !important;
        min-height: 44px !important;
        border-radius: 999px !important;
        border: 1.2px solid rgba(255,255,255,0.52) !important;
        background: rgba(255,255,255,0.90) !important;
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        caret-color: #000000 !important;
        font-size: 14px !important;
        padding: 0 16px !important;
        box-shadow: none !important;
    }}

    div[data-testid="stTextInput"] input::placeholder {{
        color: rgba(0,0,0,0.55) !important;
        -webkit-text-fill-color: rgba(0,0,0,0.55) !important;
    }}

    div[data-testid="stFormSubmitButton"] button {{
        width: 100% !important;
        height: 44px !important;
        min-height: 44px !important;
        border-radius: 999px !important;
        border: 1px solid rgba(255,255,255,0.45) !important;
        background: transparent !important;
        color: #ffffff !important;
        font-size: 14px !important;
        font-weight: 800 !important;
        letter-spacing: 0.03em !important;
    }}

    div[data-testid="stAlert"] {{
        border-radius: 14px !important;
    }}
    </style>
    """, unsafe_allow_html=True)

    left, center, right = st.columns([1.2, 1.8, 1.2])

    with center:
        st.markdown('<div class="auth-card-wrap">', unsafe_allow_html=True)
        st.markdown(
            f'<a class="auth-back" href="{build_app_url("home")}" target="_self">← Quay lại</a>',
            unsafe_allow_html=True
        )
        st.markdown('<div class="auth-title">Sign up</div>', unsafe_allow_html=True)
        st.markdown('<div class="auth-sub">Tạo tài khoản mới để lưu lịch sử riêng.</div>', unsafe_allow_html=True)

        with st.form("register_form", clear_on_submit=False):
            reg_username = st.text_input(
                "Tạo tên đăng nhập",
                placeholder="Tạo tên đăng nhập",
                key="reg_username"
            )
            reg_password = st.text_input(
                "Tạo mật khẩu",
                type="password",
                placeholder="Tạo mật khẩu",
                key="reg_password"
            )
            reg_password_2 = st.text_input(
                "Nhập lại mật khẩu",
                type="password",
                placeholder="Nhập lại mật khẩu",
                key="reg_password_2"
            )
            reg_submit = st.form_submit_button("SIGN UP", use_container_width=True)

        if reg_submit:
            if reg_password != reg_password_2:
                st.error("Mật khẩu nhập lại chưa khớp.")
            else:
                ok, msg = create_user(reg_username, reg_password)
                if ok:
                    st.success("Đăng ký thành công. Bạn hãy quay lại để đăng nhập.")
                else:
                    st.error(msg)

        st.markdown("</div>", unsafe_allow_html=True)

auth_pages = {"home", "login", "signup"}

if "auth_user" not in st.session_state:
    if page == "login":
        render_login_form_page()
    elif page == "signup":
        render_signup_form_page()
    else:
        render_auth_page()
    st.stop()

navbar_logo_src = image_to_data_uri([
    "assets/anime_teamtrangchu1.png",
    "anime_teamtrangchu1.png",
    "assets/anime_teamtrangchu1.jpg",
    "anime_teamtrangchu1.jpg",
    "assets/anime_teamtrangchu1.jpeg",
    "anime_teamtrangchu1.jpeg",
    "assets/anime_teamtrangchu1.webp",
    "anime_teamtrangchu1.webp"
])

# Topbar
current_username = escape(st.session_state["auth_user"]["username"])

navbar_html = dedent(f"""
<div class="navbar">
    <a class="logo" href="{build_app_url('home')}" target="_self">
        <img src="{navbar_logo_src}" alt="Logo">
        <span class="logo-text">Lao Cai <span>Heritage AI</span></span>
    </a>
    <div class="nav-links">
        <a href="{build_app_url('home')}" target="_self" class="{'active' if page == 'home' else ''}">Trang chủ</a>
        <a href="{build_app_url('diemden')}" target="_self" class="{'active' if page in ['diemden', 'diemden_detail'] else ''}">Điểm đến</a>
        <a href="{build_app_url('lichtrinh')}" target="_self" class="{'active' if page in ['lichtrinh', 'lichtrinh_detail'] else ''}">Lịch trình</a>
        <a href="{build_app_url('chatbot')}" target="_self" class="{'active' if page == 'chatbot' else ''}">Chatbot</a>
        <a href="{build_app_url('gioithieu')}" target="_self" class="{'active' if page == 'gioithieu' else ''}">Về dự án</a>
        <a href="{build_app_url('lehoi_langnghe')}" target="_self" class="{'active' if page == 'lehoi_langnghe' else ''}">Lễ hội - Làng nghề</a>
        <a href="{build_app_url('congdong')}" target="_self" class="{'active' if page == 'congdong' else ''}">Cộng đồng</a>
        <a href="{build_app_url('caidat')}" target="_self" class="nav-user {'active' if page == 'caidat' else ''}">👤 {current_username}</a>
    </div>
</div>
""")
st.markdown(navbar_html, unsafe_allow_html=True)

hero_home_src = image_to_data_uri([
    "kho_anh/trang_chu/chinh.png",
    "kho_anh/trang_chu/chinh.jpg",
    "kho_anh/trang_chu/chinh.jpeg",
    "kho_anh/trang_chu/chinh.webp"
])

if page == "home":
    # Hero
    st.markdown("""
    <div class="hero">
        <div class="hero-content">
            <div class="hero-title">Khám phá vẻ đẹp văn hóa và lịch sử Lào Cai</div>
            <div class="hero-subtitle">
                Một nền tảng du lịch thông minh giúp bạn tìm hiểu di tích,
                danh lam thắng cảnh và gợi ý hành trình phù hợp bằng AI.
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    kham_pha_1 = image_to_data_uri([
        "kho_anh/trang_chu/kham_pha/1.png",
        "kho_anh/trang_chu/kham_pha/1.jpg",
        "kho_anh/trang_chu/kham_pha/1.jpeg",
        "kho_anh/trang_chu/kham_pha/1.webp"
    ])

    kham_pha_2 = image_to_data_uri([
        "kho_anh/trang_chu/kham_pha/2.png",
        "kho_anh/trang_chu/kham_pha/2.jpg",
        "kho_anh/trang_chu/kham_pha/2.jpeg",
        "kho_anh/trang_chu/kham_pha/2.webp"
    ])

    kham_pha_3 = image_to_data_uri([
        "kho_anh/trang_chu/kham_pha/3.png",
        "kho_anh/trang_chu/kham_pha/3.jpg",
        "kho_anh/trang_chu/kham_pha/3.jpeg",
        "kho_anh/trang_chu/kham_pha/3.webp"
    ])

    kham_pha_4 = image_to_data_uri([
        "kho_anh/trang_chu/kham_pha/4.png",
        "kho_anh/trang_chu/kham_pha/4.jpg",
        "kho_anh/trang_chu/kham_pha/4.jpeg",
        "kho_anh/trang_chu/kham_pha/4.webp"
    ])

    kham_pha_5 = image_to_data_uri([
        "kho_anh/trang_chu/kham_pha/5.png",
        "kho_anh/trang_chu/kham_pha/5.jpg",
        "kho_anh/trang_chu/kham_pha/5.jpeg",
        "kho_anh/trang_chu/kham_pha/5.webp"
    ])

    # Khám phá nhanh
    components.html(f"""
    <style>
    body {{
        margin: 0;
        font-family: Arial, sans-serif;
    }}

    .product-section {{
        padding: 30px 0 20px 0;
        width: 100%;
        box-sizing: border-box;
        overflow: hidden;
    }}

    .product-container {{
        width: 100%;
        max-width: 1280px;
        margin: 0 auto;
        padding: 0 24px;
        box-sizing: border-box;
    }}

    .product-header {{
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        gap: 24px;
        margin-bottom: 28px;
    }}

    .product-header-left {{
        flex: 1;
        min-width: 0;
    }}

    .product-title {{
        font-size: 35px;
        font-weight: 800;
        color: #1565c0;
        text-transform: uppercase;
        line-height: 1.3;
        margin-bottom: 8px;
    }}

    .product-line {{
        width: 145px;
        height: 4px;
        background: #1565c0;
        border-radius: 999px;
        margin-bottom: 18px;
    }}

    .product-desc {{
        max-width: 940px;
        font-size: 18px;
        line-height: 1.8;
        color: #1f2937;
    }}

    .arrow-box {{
        display: flex;
        gap: 14px;
        align-items: center;
        flex-shrink: 0;
        margin-top: 10px;
    }}

    .arrow-btn {{
        width: 50px;
        height: 50px;
        border: none;
        border-radius: 50%;
        background: #ffffff;
        color: #6b7280;
        font-size: 24px;
        cursor: pointer;
        box-shadow: 0 6px 18px rgba(0,0,0,0.10);
        transition: all 0.25s ease;
    }}

    .arrow-btn:hover {{
        transform: translateY(-2px);
        color: #1565c0;
        box-shadow: 0 10px 22px rgba(0,0,0,0.14);
    }}

    .arrow-btn:disabled {{
        opacity: 0.45;
        cursor: not-allowed;
        transform: none;
    }}

    /* khung ngoài */
    .carousel-wrap {{
        width: 100%;
        overflow: hidden;
        position: relative;
        padding-right: 90px; /* chừa chỗ để lộ card sau */
        box-sizing: border-box;
    }}

    /* thanh trượt */
    .carousel-track {{
        display: flex;
        gap: 18px;
        transition: transform 0.6s ease;
        will-change: transform;
    }}

    /* mỗi card */
    .product-card {{
        position: relative;
        flex: 0 0 calc((100% - 36px) / 3); /* 3 card chính */
        height: 300px;
        border-radius: 22px;
        overflow: hidden;
        background-size: cover;
        background-position: center;
        box-shadow: 0 8px 20px rgba(0,0,0,0.08);
    }}

    .product-card::after {{
        content: "";
        position: absolute;
        inset: 0;
        background: linear-gradient(to top, rgba(0,0,0,0.62), rgba(0,0,0,0.08));
    }}

    .product-text {{
        position: absolute;
        left: 18px;
        right: 18px;
        bottom: 16px;
        z-index: 2;
        color: #fff;
        font-size: 15px;
        font-weight: 800;
        text-transform: uppercase;
        line-height: 1.5;
        text-shadow: 0 2px 8px rgba(0,0,0,0.35);
        text-align: center;
    }}

    @media (max-width: 992px) {{
        .product-header {{
            flex-direction: column;
            align-items: flex-start;
        }}

        .carousel-wrap {{
            padding-right: 40px;
        }}

        .product-card {{
            flex: 0 0 calc((100% - 18px) / 2);
            height: 270px;
        }}
    }}

    @media (max-width: 768px) {{
        .product-section {{
            padding: 18px 0 10px 0;
        }}

        .product-container {{
            padding: 0 12px;
        }}

        .product-header {{
            flex-direction: column;
            align-items: flex-start;
            gap: 14px;
            margin-bottom: 18px;
        }}

        .product-title {{
            font-size: 20px;
            line-height: 1.35;
        }}

        .product-line {{
            width: 90px;
            margin-bottom: 12px;
        }}

        .product-desc {{
            font-size: 14px;
            line-height: 1.7;
        }}

        .arrow-box {{
            width: 100%;
            justify-content: flex-start;
            gap: 10px;
            margin-top: 0;
        }}

        .carousel-wrap {{
            padding-right: 0;
        }}

        .product-card {{
            flex: 0 0 100%;
            height: 220px;
            border-radius: 18px;
        }}

        .product-text {{
            left: 12px;
            right: 12px;
            bottom: 12px;
            font-size: 13px;
            line-height: 1.45;
        }}

        .arrow-btn {{
            width: 40px;
            height: 40px;
            font-size: 18px;
        }}
    }}
                    
    </style>

    <div class="product-section">
        <div class="product-container">
            <div class="product-header">
                <div class="product-header-left">
                    <div class="product-title">KHÁM PHÁ LAO CAI HERITAGE AI</div>
                    <div class="product-line"></div>
                    <div class="product-desc">
                        Nền tảng ứng dụng trí tuệ nhân tạo giúp giới thiệu các địa danh, văn hóa, lịch sử và du lịch Lào Cai một cách trực quan, hiện đại và dễ tiếp cận.
                        Mang đến trải nghiệm khám phá thông minh, hỗ trợ tra cứu thông tin, gợi ý hành trình và kết nối người dùng với vẻ đẹp của quê hương Lào Cai.
                    </div>
                </div>

                <div class="arrow-box">
                    <button class="arrow-btn" id="prevBtn">←</button>
                    <button class="arrow-btn" id="nextBtn">→</button>
                </div>
            </div>

            <div class="carousel-wrap">
                <div class="carousel-track" id="carouselTrack"></div>
            </div>
        </div>
    </div>

    <script>
    const productData = [
        {{
            image: "{kham_pha_1}",
            text: "TIỀM NĂNG, ĐỊNH HƯỚNG PHÁT TRIỂN"
        }},
        {{
            image: "{kham_pha_2}",
            text: "Sun World Fansipan Legend"
        }},
        {{
            image: "{kham_pha_3}",
            text: "đua ngựa bắc hà"
        }},
        {{
            image: "{kham_pha_4}",
            text: "nhà thờ đá sapa"
        }},
        {{
            image: "{kham_pha_5}",
            text: "tâm linh - đền chiềng ken"
        }}
    ];

    const track = document.getElementById("carouselTrack");
    const prevBtn = document.getElementById("prevBtn");
    const nextBtn = document.getElementById("nextBtn");

    let currentIndex = 0;

    function buildCards() {{
        track.innerHTML = "";
        productData.forEach(item => {{
            const card = document.createElement("div");
            card.className = "product-card";
            card.style.backgroundImage = `url('${{item.image}}')`;
            card.innerHTML = `<div class="product-text">${{item.text}}</div>`;
            track.appendChild(card);
        }});
    }}

    function getVisibleCount() {{
        if (window.innerWidth <= 768) return 1;
        if (window.innerWidth <= 992) return 2;
        return 3;
    }}

    function updateCarousel() {{
        const cards = track.querySelectorAll(".product-card");
        if (!cards.length) return;

        const gap = 18;
        const cardWidth = cards[0].offsetWidth + gap;
        track.style.transform = `translateX(-${{currentIndex * cardWidth}}px)`;

        prevBtn.disabled = currentIndex === 0;
        nextBtn.disabled = currentIndex >= productData.length - getVisibleCount();
    }}

    prevBtn.addEventListener("click", () => {{
        if (currentIndex > 0) {{
            currentIndex--;
            updateCarousel();
        }}
    }});

    nextBtn.addEventListener("click", () => {{
        if (currentIndex < productData.length - getVisibleCount()) {{
            currentIndex++;
            updateCarousel();
        }}
    }});

    window.addEventListener("resize", updateCarousel);

    buildCards();
    setTimeout(updateCarousel, 100);
    </script>
    """, height=540)

    # Miền Bắc
    nb_quangninh = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/quangninh.png",
        "kho_anh/trang_chu/noi_bat/bac/quangninh.jpg",
        "kho_anh/trang_chu/noi_bat/bac/quangninh.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/quangninh.webp"
    ])

    nb_hagiang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/hagiang.png",
        "kho_anh/trang_chu/noi_bat/bac/hagiang.jpg",
        "kho_anh/trang_chu/noi_bat/bac/hagiang.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/hagiang.webp"
    ])

    nb_laocai = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/laocai.png",
        "kho_anh/trang_chu/noi_bat/bac/laocai.jpg",
        "kho_anh/trang_chu/noi_bat/bac/laocai.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/laocai.webp"
    ])

    nb_ninhbinh = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/ninhbinh.png",
        "kho_anh/trang_chu/noi_bat/bac/ninhbinh.jpg",
        "kho_anh/trang_chu/noi_bat/bac/ninhbinh.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/ninhbinh.webp"
    ])

    nb_yenbai = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/yenbai.png",
        "kho_anh/trang_chu/noi_bat/bac/yenbai.jpg",
        "kho_anh/trang_chu/noi_bat/bac/yenbai.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/yenbai.webp"
    ])

    nb_sonla = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/sonla.png",
        "kho_anh/trang_chu/noi_bat/bac/sonla.jpg",
        "kho_anh/trang_chu/noi_bat/bac/sonla.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/sonla.webp"
    ])

    nb_caobang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/caobang.png",
        "kho_anh/trang_chu/noi_bat/bac/caobang.jpg",
        "kho_anh/trang_chu/noi_bat/bac/caobang.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/caobang.webp"
    ])

    nb_haiphong = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/haiphong.png",
        "kho_anh/trang_chu/noi_bat/bac/haiphong.jpg",
        "kho_anh/trang_chu/noi_bat/bac/haiphong.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/haiphong.webp"
    ])

    nb_hanoi = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/bac/hanoi.png",
        "kho_anh/trang_chu/noi_bat/bac/hanoi.jpg",
        "kho_anh/trang_chu/noi_bat/bac/hanoi.jpeg",
        "kho_anh/trang_chu/noi_bat/bac/hanoi.webp"
    ])

    # Miền Trung
    nt_hue = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/hue.png",
        "kho_anh/trang_chu/noi_bat/trung/hue.jpg",
        "kho_anh/trang_chu/noi_bat/trung/hue.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/hue.webp"
    ])

    nt_danang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/danang.png",
        "kho_anh/trang_chu/noi_bat/trung/danang.jpg",
        "kho_anh/trang_chu/noi_bat/trung/danang.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/danang.webp"
    ])

    nt_quangnam = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/quangnam.png",
        "kho_anh/trang_chu/noi_bat/trung/quangnam.jpg",
        "kho_anh/trang_chu/noi_bat/trung/quangnam.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/quangnam.webp"
    ])

    nt_nhatrang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/nhatrang.png",
        "kho_anh/trang_chu/noi_bat/trung/nhatrang.jpg",
        "kho_anh/trang_chu/noi_bat/trung/nhatrang.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/nhatrang.webp"
    ])

    nt_phuyen = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/phuyen.png",
        "kho_anh/trang_chu/noi_bat/trung/phuyen.jpg",
        "kho_anh/trang_chu/noi_bat/trung/phuyen.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/phuyen.webp"
    ])

    nt_binhdinh = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/binhdinh.png",
        "kho_anh/trang_chu/noi_bat/trung/binhdinh.jpg",
        "kho_anh/trang_chu/noi_bat/trung/binhdinh.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/binhdinh.webp"
    ])

    nt_quangbinh = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/quangbinh.png",
        "kho_anh/trang_chu/noi_bat/trung/quangbinh.jpg",
        "kho_anh/trang_chu/noi_bat/trung/quangbinh.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/quangbinh.webp"
    ])

    nt_thanhhoa = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/thanhhoa.png",
        "kho_anh/trang_chu/noi_bat/trung/thanhhoa.jpg",
        "kho_anh/trang_chu/noi_bat/trung/thanhhoa.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/thanhhoa.webp"
    ])

    nt_nghean = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/trung/nghean.png",
        "kho_anh/trang_chu/noi_bat/trung/nghean.jpg",
        "kho_anh/trang_chu/noi_bat/trung/nghean.jpeg",
        "kho_anh/trang_chu/noi_bat/trung/nghean.webp"
    ])

    # Miền Nam
    nn_hochiminh = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/hochiminh.png",
        "kho_anh/trang_chu/noi_bat/nam/hochiminh.jpg",
        "kho_anh/trang_chu/noi_bat/nam/hochiminh.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/hochiminh.webp"
    ])

    nn_kiengiang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/kiengiang.png",
        "kho_anh/trang_chu/noi_bat/nam/kiengiang.jpg",
        "kho_anh/trang_chu/noi_bat/nam/kiengiang.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/kiengiang.webp"
    ])

    nn_lamdong = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/lamdong.png",
        "kho_anh/trang_chu/noi_bat/nam/lamdong.jpg",
        "kho_anh/trang_chu/noi_bat/nam/lamdong.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/lamdong.webp"
    ])

    nn_vungtau = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/vungtau.png",
        "kho_anh/trang_chu/noi_bat/nam/vungtau.jpg",
        "kho_anh/trang_chu/noi_bat/nam/vungtau.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/vungtau.webp"
    ])

    nn_cantho = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/cantho.png",
        "kho_anh/trang_chu/noi_bat/nam/cantho.jpg",
        "kho_anh/trang_chu/noi_bat/nam/cantho.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/cantho.webp"
    ])

    nn_angiang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/angiang.png",
        "kho_anh/trang_chu/noi_bat/nam/angiang.jpg",
        "kho_anh/trang_chu/noi_bat/nam/angiang.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/angiang.webp"
    ])

    nn_tiengiang = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/tiengiang.png",
        "kho_anh/trang_chu/noi_bat/nam/tiengiang.jpg",
        "kho_anh/trang_chu/noi_bat/nam/tiengiang.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/tiengiang.webp"
    ])

    nn_bentre = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/bentre.png",
        "kho_anh/trang_chu/noi_bat/nam/bentre.jpg",
        "kho_anh/trang_chu/noi_bat/nam/bentre.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/bentre.webp"
    ])

    nn_dongthap = image_to_data_uri([
        "kho_anh/trang_chu/noi_bat/nam/dongthap.png",
        "kho_anh/trang_chu/noi_bat/nam/dongthap.jpg",
        "kho_anh/trang_chu/noi_bat/nam/dongthap.jpeg",
        "kho_anh/trang_chu/noi_bat/nam/dongthap.webp"
    ])

    # Điểm đến yêu thích
    components.html(dedent(f"""
    <style>
    body {{
        margin: 0;
        font-family: Arial, sans-serif;
    }}

    .favorite-section {{
        width: 100%;
        padding: 24px 0;
        box-sizing: border-box;
    }}

    .favorite-container {{
        max-width: 1350px;
        margin: 0 auto;
        padding: 0 18px;
        box-sizing: border-box;
    }}

    .favorite-title {{
        text-align: center;
        font-size: 35px;
        font-weight: 900;
        color: #1565c0;
        text-transform: uppercase;
        margin-bottom: 8px;
        line-height: 1.3;
    }}

    .favorite-line {{
        width: 150px;
        height: 4px;
        background: #1565c0;
        border-radius: 999px;
        margin: 0 auto 18px auto;
    }}

    .favorite-desc {{
        text-align: center;
        max-width: 820px;
        margin: 0 auto 24px auto;
        font-size: 18px;
        line-height: 1.8;
        color: #374151;
    }}

    .favorite-tabs {{
        display: flex;
        justify-content: center;
        flex-wrap: wrap;
        gap: 36px;
        margin-bottom: 28px;
    }}

    .favorite-tab {{
        position: relative;
        font-size: 17px;
        font-weight: 700;
        color: #374151;
        padding-bottom: 8px;
        cursor: pointer;
        transition: 0.25s ease;
    }}

    .favorite-tab:hover {{
        color: #1565c0;
    }}

    .favorite-tab.active {{
        color: #1565c0;
    }}

    .favorite-tab.active::after {{
        content: "";
        position: absolute;
        left: 0;
        bottom: 0;
        width: 100%;
        height: 3px;
        border-radius: 999px;
        background: #1565c0;
    }}

    .favorite-grid {{
        display: grid;
        grid-template-columns: 1.35fr 1.05fr 1.05fr 1.05fr;
        grid-template-rows: 180px 180px 220px;
        gap: 10px;
    }}

    /* CARD */
    .favorite-card {{
        position: relative;
        overflow: hidden;
        border-radius: 16px;
        min-height: 180px;
        background-size: cover;
        background-position: center;
        box-shadow: 0 6px 18px rgba(0,0,0,0.08);
        isolation: isolate;
        opacity: 1;
        transform: translateY(0);
        transition: opacity 0.35s ease, transform 0.35s ease;
    }}

    .favorite-card.fade-out {{
        opacity: 0;
        transform: translateY(8px);
    }}

    .favorite-card::before {{
        content: "";
        position: absolute;
        inset: 0;
        background-image: inherit;
        background-size: cover;
        background-position: center;
        transform: scale(1);
        transition: transform 0.5s ease;
        z-index: 0;
    }}

    .favorite-overlay {{
        position: absolute;
        inset: 0;
        background: linear-gradient(to top, rgba(0,0,0,0.55), rgba(0,0,0,0.12));
        transition: 0.35s ease;
        z-index: 1;
    }}

    .favorite-content {{
        position: absolute;
        inset: 0;
        z-index: 2;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 18px;
        text-align: center;
        color: #fff;
    }}

    .favorite-name {{
        margin: 0;
        font-size: 18px;
        font-weight: 800;
        text-transform: uppercase;
        line-height: 1.3;
        transition: 0.35s ease;
    }}

    .favorite-content {{
        position: absolute;
        inset: 0;
        z-index: 2;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 18px;
        text-align: center;
        color: #fff;
    }}

    .favorite-name {{
        margin: 0;
        font-size: 18px;
        font-weight: 800;
        text-transform: uppercase;
        line-height: 1.3;
        transition: 0.35s ease;
    }}

    .favorite-info {{
        margin-top: 12px;
        max-width: 88%;
        font-size: 14px;
        line-height: 1.6;
        color: #ffffff;
        opacity: 0;
        transform: translateY(12px);
        transition: 0.35s ease;
        display: -webkit-box;
        -webkit-line-clamp: 4;
        -webkit-box-orient: vertical;
        overflow: hidden;
        text-shadow: 0 2px 8px rgba(0,0,0,0.35);
    }}

.favorite-card:hover::before {{
    transform: scale(1.08);
}}

.favorite-card:hover .favorite-overlay {{
    background: linear-gradient(to top, rgba(0,0,0,0.78), rgba(0,0,0,0.28));
}}

.favorite-card:hover .favorite-name {{
    transform: translateY(-8px);
}}

.favorite-card:hover .favorite-info {{
    opacity: 1;
    transform: translateY(0);
}}

    /* vị trí từng card */
    .card-1 {{
        grid-column: 1 / 2;
        grid-row: 1 / 3;
    }}

    .card-2 {{
        grid-column: 2 / 3;
        grid-row: 1 / 2;
    }}

    .card-3 {{
        grid-column: 3 / 5;
        grid-row: 1 / 2;
    }}

    .card-4 {{
        grid-column: 2 / 3;
        grid-row: 2 / 3;
    }}

    .card-5 {{
        grid-column: 3 / 4;
        grid-row: 2 / 3;
    }}

    .card-6 {{
        grid-column: 4 / 5;
        grid-row: 2 / 4;
    }}

    .card-7 {{
        grid-column: 1 / 2;
        grid-row: 3 / 4;
    }}

    .card-8 {{
        grid-column: 2 / 3;
        grid-row: 3 / 4;
    }}

    .card-9 {{
        grid-column: 3 / 4;
        grid-row: 3 / 4;
    }}

    @media (max-width: 1100px) {{
        .favorite-grid {{
            grid-template-columns: 1fr 1fr 1fr;
            grid-template-rows: 220px 220px 220px;
        }}

        .card-1 {{ grid-column: 1 / 2; grid-row: 1 / 3; }}
        .card-2 {{ grid-column: 2 / 3; grid-row: 1 / 2; }}
        .card-3 {{ grid-column: 3 / 4; grid-row: 1 / 2; }}
        .card-4 {{ grid-column: 2 / 3; grid-row: 2 / 3; }}
        .card-5 {{ grid-column: 3 / 4; grid-row: 2 / 3; }}
        .card-6 {{ grid-column: 1 / 2; grid-row: 3 / 4; }}
        .card-7 {{ grid-column: 2 / 3; grid-row: 3 / 4; }}
        .card-8 {{ grid-column: 3 / 4; grid-row: 3 / 4; }}
        .card-9 {{ display: none; }}
    }}

    @media (max-width: 768px) {{
        .favorite-section {{
            padding: 12px 0 4px 0;
        }}

        .favorite-container {{
            padding: 0 12px;
        }}

        .favorite-title {{
            font-size: 22px;
            line-height: 1.35;
        }}

        .favorite-line {{
            width: 90px;
            margin-bottom: 12px;
        }}

        .favorite-desc {{
            font-size: 14px;
            line-height: 1.7;
            margin-bottom: 18px;
        }}

        .favorite-tabs {{
            gap: 10px 16px;
            margin-bottom: 18px;
        }}

        .favorite-tab {{
            font-size: 14px;
            padding-bottom: 6px;
        }}

        .favorite-grid {{
            grid-template-columns: 1.15fr 1fr;
            grid-template-rows: 145px 145px 95px 95px 95px;
            gap: 10px;
        }}

        .favorite-card {{
            border-radius: 12px;
        }}

        .favorite-content {{
            padding: 10px;
        }}

        .favorite-name {{
            font-size: 14px;
            line-height: 1.25;
        }}

        .favorite-info {{
            margin-top: 6px;
            max-width: 94%;
            font-size: 11px;
            line-height: 1.45;
            opacity: 1;
            transform: none;
            -webkit-line-clamp: 3;
        }}

        .card-1 {{
            grid-column: 1 / 2;
            grid-row: 1 / 3;
            min-height: 300px;
        }}

        .card-2 {{
            grid-column: 2 / 3;
            grid-row: 1 / 2;
            min-height: 145px;
        }}

        .card-3 {{
            grid-column: 2 / 3;
            grid-row: 2 / 3;
            min-height: 145px;
        }}

        .card-4 {{
            grid-column: 1 / 2;
            grid-row: 3 / 4;
            min-height: 95px;
        }}

        .card-5 {{
            grid-column: 2 / 3;
            grid-row: 3 / 4;
            min-height: 95px;
        }}

        .card-6 {{
            grid-column: 1 / 3;
            grid-row: 4 / 5;
            min-height: 95px;
        }}

        .card-7 {{
            grid-column: 1 / 2;
            grid-row: 5 / 6;
            min-height: 95px;
        }}

        .card-8 {{
            grid-column: 2 / 3;
            grid-row: 5 / 6;
            min-height: 95px;
        }}

        .card-9 {{
            grid-column: 1 / 3;
            grid-row: auto;
            min-height: 95px;
        }}
    }}
                           
    </style>

    <div class="favorite-section">
        <div class="favorite-container">

            <div class="favorite-title">Điểm đến nổi bật</div>
            <div class="favorite-line"></div>

            <div class="favorite-desc">
                Các điểm đến tiêu biểu cùng những thông tin hữu ích về thiên nhiên, văn hóa, con người và các hành trình trải nghiệm nổi bật tại mỗi vùng miền.
            </div>

            <div class="favorite-tabs">
                <div class="favorite-tab active" data-region="north">Miền Bắc</div>
                <div class="favorite-tab" data-region="central">Miền Trung</div>
                <div class="favorite-tab" data-region="south">Miền Nam</div>
            </div>

            <div class="favorite-grid" id="favoriteGrid">
                <div class="favorite-card card-1">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-2">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-3">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-4">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-5">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-6">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-7">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-8">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>

                <div class="favorite-card card-9">
                    <div class="favorite-overlay"></div>
                    <div class="favorite-content">
                        <div class="favorite-name"></div>
                        <div class="favorite-info"></div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script>
    const destinationData = {{
        north: [
            {{
                name: "Quảng Ninh",
                image: "{nb_quangninh}",
                info: "Nổi tiếng với Vịnh Hạ Long, Quảng Ninh là điểm đến hấp dẫn với cảnh biển đảo kỳ vĩ, thiên nhiên đặc sắc và nhiều trải nghiệm du lịch đáng nhớ."
            }},
            {{
                name: "Hà Giang",
                image: "{nb_hagiang}",
                info: "Hà Giang gây ấn tượng bởi núi non hùng vĩ, những cung đèo nổi tiếng, bản làng vùng cao và vẻ đẹp hoang sơ đậm chất miền đá."
            }},
            {{
                name: "Lào Cai",
                image: "{nb_laocai}",
                info: "Lào Cai là vùng đất hội tụ vẻ đẹp của núi rừng Tây Bắc, nổi bật với Sa Pa, Fansipan, ruộng bậc thang và bản sắc văn hóa dân tộc đặc sắc."
            }},
            {{
                name: "Ninh Bình",
                image: "{nb_ninhbinh}",
                info: "Ninh Bình thu hút du khách bởi quần thể danh thắng Tràng An, Tam Cốc, Hang Múa và khung cảnh non nước hữu tình như tranh vẽ."
            }},
            {{
                name: "Yên Bái",
                image: "{nb_yenbai}",
                info: "Yên Bái được biết đến với Mù Cang Chải, những thửa ruộng bậc thang tuyệt đẹp cùng vẻ đẹp bình dị, thơ mộng của núi rừng Tây Bắc."
            }},
            {{
                name: "Sơn La",
                image: "{nb_sonla}",
                info: "Sơn La nổi bật với Mộc Châu xanh mát, đồi chè trải dài, mùa hoa rực rỡ và không khí trong lành, dễ chịu quanh năm."
            }},
            {{
                name: "Cao Bằng",
                image: "{nb_caobang}",
                info: "Cao Bằng là vùng đất biên giới nổi tiếng với thác Bản Giốc, động Ngườm Ngao và cảnh quan thiên nhiên hùng vĩ, nguyên sơ."
            }},
            {{
                name: "Hải Phòng",
                image: "{nb_haiphong}",
                info: "Hải Phòng là thành phố cảng sôi động, hấp dẫn du khách với Cát Bà, Đồ Sơn, ẩm thực hải sản phong phú và nhịp sống trẻ trung."
            }},
            {{
                name: "Hà Nội",
                image: "{nb_hanoi}",
                info: "Hà Nội là thủ đô ngàn năm văn hiến, nổi bật với Hồ Gươm, phố cổ, Văn Miếu và sự giao thoa hài hòa giữa truyền thống với hiện đại."
            }}
        ],

        central: [
            {{
                name: "Huế",
                image: "{nt_hue}",
                info: "Huế mang vẻ đẹp trầm mặc và sâu lắng, nổi tiếng với Đại Nội, lăng tẩm triều Nguyễn, chùa Thiên Mụ và nền văn hóa giàu bản sắc."
            }},
            {{
                name: "Đà Nẵng",
                image: "{nt_danang}",
                info: "Đà Nẵng là thành phố biển hiện đại, nổi bật với cầu Rồng, biển Mỹ Khê, Bà Nà Hills và môi trường du lịch năng động, văn minh."
            }},
            {{
                name: "Quảng Nam",
                image: "{nt_quangnam}",
                info: "Quảng Nam hấp dẫn bởi phố cổ Hội An, thánh địa Mỹ Sơn, các làng nghề truyền thống và vẻ đẹp văn hóa đậm chất miền Trung."
            }},
            {{
                name: "Nha Trang",
                image: "{nt_nhatrang}",
                info: "Nha Trang là điểm đến lý tưởng với biển xanh, cát trắng, nhiều hòn đảo đẹp, hải sản phong phú và không khí nghỉ dưỡng sôi động."
            }},
            {{
                name: "Phú Yên",
                image: "{nt_phuyen}",
                info: "Phú Yên gây ấn tượng bởi vẻ đẹp hoang sơ, yên bình với Gành Đá Đĩa, Bãi Xép và khung cảnh biển trời trong trẻo, thơ mộng."
            }},
            {{
                name: "Bình Định",
                image: "{nt_binhdinh}",
                info: "Bình Định nổi tiếng với Quy Nhơn, Eo Gió, Kỳ Co cùng sự kết hợp hài hòa giữa cảnh đẹp biển trời và truyền thống võ cổ truyền."
            }},
            {{
                name: "Quảng Bình",
                image: "{nt_quangbinh}",
                info: "Quảng Bình được mệnh danh là vương quốc hang động với Phong Nha - Kẻ Bàng, thiên nhiên hùng vĩ và nhiều kỳ quan độc đáo."
            }},
            {{
                name: "Thanh Hóa",
                image: "{nt_thanhhoa}",
                info: "Thanh Hóa sở hữu nhiều điểm đến nổi bật như biển Sầm Sơn, thành nhà Hồ, suối cá Cẩm Lương và dấu ấn lịch sử lâu đời."
            }},
            {{
                name: "Nghệ An",
                image: "{nt_nghean}",
                info: "Nghệ An là vùng đất giàu truyền thống văn hóa và lịch sử, nổi bật với quê Bác, biển Cửa Lò và vẻ đẹp mộc mạc của xứ Nghệ."
            }}
        ],

        south: [
            {{
                name: "TP. Hồ Chí Minh",
                image: "{nn_hochiminh}",
                info: "TP. Hồ Chí Minh là trung tâm kinh tế sôi động bậc nhất cả nước, nổi bật với nhịp sống hiện đại, các công trình biểu tượng và sức trẻ năng động."
            }},
            {{
                name: "Kiên Giang",
                image: "{nn_kiengiang}",
                info: "Kiên Giang, đặc biệt là Phú Quốc, nổi tiếng với biển xanh, cát trắng, hoàng hôn đẹp, làng chài truyền thống và nhiều khu nghỉ dưỡng hấp dẫn."
            }},
            {{
                name: "Lâm Đồng",
                image: "{nn_lamdong}",
                info: "Lâm Đồng với thành phố Đà Lạt mang vẻ đẹp dịu dàng, thơ mộng, nổi bật bởi khí hậu mát mẻ, đồi thông, hồ nước và muôn vàn loài hoa."
            }},
            {{
                name: "Bà Rịa - Vũng Tàu",
                image: "{nn_vungtau}",
                info: "Bà Rịa - Vũng Tàu là điểm đến quen thuộc với du khách yêu biển, nổi bật với bãi tắm đẹp, hải sản ngon và không khí nghỉ dưỡng thoải mái."
            }},
            {{
                name: "Cần Thơ",
                image: "{nn_cantho}",
                info: "Cần Thơ là đô thị trung tâm của miền Tây Nam Bộ, nổi tiếng với chợ nổi Cái Răng, bến Ninh Kiều và nét văn hóa sông nước đặc trưng."
            }},
            {{
                name: "An Giang",
                image: "{nn_angiang}",
                info: "An Giang thu hút du khách với rừng tràm Trà Sư, núi Cấm, miếu Bà Chúa Xứ và vẻ đẹp giao hòa giữa thiên nhiên với văn hóa tâm linh."
            }},
            {{
                name: "Tiền Giang",
                image: "{nn_tiengiang}",
                info: "Tiền Giang hấp dẫn bởi những cù lao xanh mát, vườn trái cây trù phú, chợ nổi và trải nghiệm đậm chất miền sông nước Nam Bộ."
            }},
            {{
                name: "Bến Tre",
                image: "{nn_bentre}",
                info: "Bến Tre nổi tiếng là xứ dừa với hệ thống kênh rạch xanh mát, làng nghề truyền thống và nhịp sống miệt vườn thanh bình, gần gũi."
            }},
            {{
                name: "Đồng Tháp",
                image: "{nn_dongthap}",
                info: "Đồng Tháp ghi dấu ấn với đất sen hồng, làng hoa Sa Đéc, khu sinh thái đặc sắc và vẻ đẹp mộc mạc, yên bình của miền Tây."
            }}
        ]
    }};

    const tabs = document.querySelectorAll(".favorite-tab");
    const cards = document.querySelectorAll(".favorite-card");

    function renderRegion(regionKey) {{
        const data = destinationData[regionKey];

        cards.forEach(card => card.classList.add("fade-out"));

        setTimeout(() => {{
            cards.forEach((card, index) => {{
                card.style.backgroundImage = `url('${{data[index].image}}')`;
                card.querySelector(".favorite-name").textContent = data[index].name;
                card.querySelector(".favorite-info").textContent = data[index].info;
                card.classList.remove("fade-out");
            }});
        }}, 220);
    }}

    tabs.forEach(tab => {{
        tab.addEventListener("click", () => {{
            tabs.forEach(item => item.classList.remove("active"));
            tab.classList.add("active");
            renderRegion(tab.dataset.region);
        }});
    }});

    renderRegion("north");
    </script>
    """), height=850, scrolling=False)

    # Footer
    home_logo_src = image_to_data_uri([
        "assets/anime_teamtrangchu.png",
        "anime_teamtrangchu.png",
        "assets/anime_teamtrangchu.jpg",
        "anime_teamtrangchu.jpg",
        "assets/anime_teamtrangchu.jpeg",
        "anime_teamtrangchu.jpeg",
        "assets/anime_teamtrangchu.webp",
        "anime_teamtrangchu.webp"
    ]) or "https://dummyimage.com/400x400/f8fafc/94a3b8.png&text=LOGO"

    vitri_icon_src = image_to_data_uri([
        "assets/vitri.png", "vitri.png",
        "assets/vitri.jpg", "vitri.jpg",
        "assets/vitri.jpeg", "vitri.jpeg",
        "assets/vitri.webp", "vitri.webp"
    ])

    email_icon_src = image_to_data_uri([
        "assets/email.png", "email.png",
        "assets/email.jpg", "email.jpg",
        "assets/email.jpeg", "email.jpeg",
        "assets/email.webp", "email.webp"
    ])

    dienthoai_icon_src = image_to_data_uri([
        "assets/dienthoai.png", "dienthoai.png",
        "assets/dienthoai.jpg", "dienthoai.jpg",
        "assets/dienthoai.jpeg", "dienthoai.jpeg",
        "assets/dienthoai.webp", "dienthoai.webp"
    ])

    facebook_icon_src = image_to_data_uri([
        "assets/facebook.png", "facebook.png",
        "assets/facebook.jpg", "facebook.jpg",
        "assets/facebook.jpeg", "facebook.jpeg",
        "assets/facebook.webp", "facebook.webp"
    ])

    tiktok_icon_src = image_to_data_uri([
        "assets/tiktok.png", "tiktok.png",
        "assets/tiktok.jpg", "tiktok.jpg",
        "assets/tiktok.jpeg", "tiktok.jpeg",
        "assets/tiktok.webp", "tiktok.webp"
    ])

    youtube_icon_src = image_to_data_uri([
        "assets/youtube.png", "youtube.png",
        "assets/youtube.jpg", "youtube.jpg",
        "assets/youtube.jpeg", "youtube.jpeg",
        "assets/youtube.webp", "youtube.webp"
    ])

    zalo_icon_src = image_to_data_uri([
        "assets/zalo.png", "zalo.png",
        "assets/zalo.jpg", "zalo.jpg",
        "assets/zalo.jpeg", "zalo.jpeg",
        "assets/zalo.webp", "zalo.webp"
    ])

    facebook_icon_html = f'<img class="social-icon" src="{facebook_icon_src}" alt="Facebook">' if facebook_icon_src else "f"
    tiktok_icon_html = f'<img class="social-icon" src="{tiktok_icon_src}" alt="TikTok">' if tiktok_icon_src else "♪"
    youtube_icon_html = f'<img class="social-icon" src="{youtube_icon_src}" alt="YouTube">' if youtube_icon_src else "▶"
    zalo_icon_html = f'<img class="social-icon" src="{zalo_icon_src}" alt="Zalo">' if zalo_icon_src else "z"

    facebook_link = "https://www.facebook.com/share/1CubUJMiYU/"
    tiktok_link = "https://www.tiktok.com/@minhkhoadayyy?is_from_webapp=1&sender_device=pc"
    youtube_link = "http://www.youtube.com/@MINHKHOALT"
    zalo_link = "https://zalo.me/0346538917"

    facebook_href = escape(facebook_link, quote=True)
    tiktok_href = escape(tiktok_link, quote=True)
    youtube_href = escape(youtube_link, quote=True)
    zalo_href = escape(zalo_link, quote=True)

    vitri_icon_html = f'<img class="info-icon" src="{vitri_icon_src}" alt="Vị trí">' if vitri_icon_src else "📍"
    email_icon_html = f'<img class="info-icon" src="{email_icon_src}" alt="Email">' if email_icon_src else "✉️"
    dienthoai_icon_html = f'<img class="info-icon" src="{dienthoai_icon_src}" alt="Điện thoại">' if dienthoai_icon_src else "📞"

    project_title = "LAO CAI HERITAGE AI"
    project_subtitle = (
        "Là sản phẩm số giúp giới thiệu điểm đến, lịch trình và giá trị văn hóa "
        "Lào Cai bằng giao diện trực quan kết hợp AI."
    )

    contact_address = "Khánh Yên, Lào Cai, Việt Nam"
    contact_email = "khoagaming999@gmail.com"
    contact_phone = "0346 538 917"

    components.html(f"""
    <style>
    * {{
        box-sizing: border-box;
    }}

    html, body {{
        margin: 0;
        padding: 0;
        background: transparent;
        font-family: Inter, Arial, sans-serif;
    }}

    .about-contact-wrap {{
        max-width: 1600px;
        margin: 56px auto 28px auto;
        padding: 0 16px;
    }}

    .about-contact {{
        position: relative;
        overflow: hidden;
        border-radius: 34px;
        padding: 34px;
        background:
            radial-gradient(circle at top left, rgba(255,255,255,0.18), transparent 28%),
            linear-gradient(135deg, #6f829e 0%, #8799b3 35%, #b9c6d6 100%);
        box-shadow: 0 20px 52px rgba(15,23,42,0.16);
    }}

    .about-contact::before {{
        content: "";
        position: absolute;
        width: 320px;
        height: 320px;
        border-radius: 50%;
        background: rgba(255,255,255,0.08);
        top: -140px;
        right: -80px;
    }}

    .about-contact::after {{
        content: "";
        position: absolute;
        width: 240px;
        height: 240px;
        border-radius: 50%;
        background: rgba(255,255,255,0.06);
        bottom: -110px;
        left: -70px;
    }}

    .about-contact-grid {{
        position: relative;
        z-index: 2;
        display: grid;
        grid-template-columns: 1.08fr 0.92fr;
        gap: 34px;
        align-items: stretch;
    }}

    .about-left {{
        color: #ffffff;
        display: flex;
        flex-direction: column;
        justify-content: center;
    }}

    .brand-row {{
        display: flex;
        gap: 22px;
        align-items: center;
        margin-bottom: 22px;
    }}

    .brand-logo {{
        width: 176px;
        height: 176px;
        min-width: 176px;
        border-radius: 50%;
        overflow: hidden;
        border: 5px solid rgba(255,255,255,0.68);
        box-shadow: 0 12px 28px rgba(0,0,0,0.16);
        background: rgba(255,255,255,0.16);
    }}

    .brand-logo img {{
        width: 100%;
        height: 100%;
        object-fit: cover;
        display: block;
    }}

    .brand-title {{
        font-size: 48px;
        line-height: 1.02;
        font-weight: 900;
        margin: 0 0 12px 0;
        letter-spacing: 0.01em;
        color: #ffffff;
    }}

    .brand-subtitle {{
        font-size: 18px;
        line-height: 1.82;
        color: rgba(255,255,255,0.94);
        max-width: 540px;
    }}

    .brand-line {{
        width: 190px;
        height: 4px;
        margin: 18px 0 0 0;
        border-radius: 999px;
        background: linear-gradient(90deg, rgba(255,255,255,0.92), rgba(255,255,255,0.18));
    }}

    .about-bottom {{
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 24px;
        margin-top: 12px;
    }}

    .info-box {{
        background: rgba(255,255,255,0.10);
        border: 1px solid rgba(255,255,255,0.18);
        border-radius: 24px;
        padding: 20px 22px;
        backdrop-filter: blur(8px);
        -webkit-backdrop-filter: blur(8px);
    }}
                    
    .social-row {{
        display: flex;
        flex-wrap: nowrap;
        align-items: center;
        justify-content: center;
        gap: 16px;
        margin-top: 10px;
    }}

    .info-title {{
        font-size: 27px;
        font-weight: 900;
        margin-bottom: 14px;
        color: #ffffff;            
        text-align: center;
        width: 100%;
    }}

    .info-item{{
        display:flex;
        align-items:center;
        gap:10px;
        font-size:20px;
        line-height:1.9;
        color:rgba(255,255,255,0.93);
    }}

    .info-icon{{
        width:18px;
        height:18px;
        object-fit:contain;
        display:block;
        flex-shrink:0;
    }}

    .social-row {{
        display: flex;
        flex-wrap: wrap;
        gap: 12px;
        margin-top: 8px;
    }}

    .social-btn{{
        width:54px;
        height:54px;
        border-radius:50%;
        display:flex;
        align-items:center;
        justify-content:center;
        background:#ffffff;
        color:#253247;
        font-size:16px;
        font-weight:900;
        box-shadow:0 8px 18px rgba(0,0,0,0.12);
        overflow:hidden;
        flex-shrink:0;
        text-decoration:none;
        cursor:pointer;
        transition: transform 0.18s ease, box-shadow 0.18s ease;
    }}

    .social-btn:hover{{
        transform: translateY(-2px);
        box-shadow:0 12px 24px rgba(0,0,0,0.16);
    }}

    .social-icon{{
        width:24px;
        height:24px;
        object-fit:contain;
        display:block;
    }}

    .about-right {{
        display: flex;
        align-items: center;
    }}

    .form-card {{
        width: 100%;
        background: rgba(255,255,255,0.28);
        border: 1px solid rgba(255,255,255,0.28);
        border-radius: 30px;
        padding: 16px;
        backdrop-filter: blur(14px);
        -webkit-backdrop-filter: blur(14px);
        box-shadow: inset 0 1px 0 rgba(255,255,255,0.34);
    }}

    .form-inner {{
        background: rgba(255,255,255,0.76);
        border-radius: 24px;
        padding: 18px;
    }}

    .form-row {{
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 14px;
        margin-bottom: 14px;
    }}

    .input,
    .textarea {{
        width: 100%;
        border: 1px solid rgba(148,163,184,0.28);
        background: rgba(255,255,255,0.94);
        border-radius: 16px;
        padding: 14px 16px;
        font-size: 15px;
        color: #0f172a;
        outline: none;
    }}

    .input::placeholder,
    .textarea::placeholder {{
        color: #94a3b8;
    }}

    .input:focus,
    .textarea:focus {{
        border-color: #8ca8c7;
        box-shadow: 0 0 0 3px rgba(140,168,199,0.14);
    }}

    .textarea {{
        min-height: 210px;
        resize: none;
        margin-bottom: 16px;
    }}

    .submit-wrap {{
        display: flex;
        justify-content: center;
    }}

    .submit-btn {{
        border: none;
        min-width: 180px;
        height: 50px;
        border-radius: 999px;
        background: linear-gradient(135deg, #d8dde6 0%, #bcc6d4 100%);
        color: #3f4b5d;
        font-size: 22px;
        font-weight: 800;
        cursor: pointer;
        box-shadow: 0 10px 22px rgba(15,23,42,0.10);
        transition: transform 0.18s ease, box-shadow 0.18s ease;
    }}

    .submit-btn:hover {{
        transform: translateY(-1px);
        box-shadow: 0 14px 26px rgba(15,23,42,0.14);
    }}
                    
    .feedback-status {{
        margin-top: 12px;
        text-align: center;
        font-size: 14px;
        font-weight: 700;
        color: #475569;
        min-height: 20px;
    }}

    @media (max-width: 1100px) {{
        .about-contact-grid {{
            grid-template-columns: 1fr;
        }}

        .brand-title {{
            font-size: 42px;
        }}
    }}

    @media (max-width: 720px) {{
        .about-contact-wrap {{
            margin: 20px auto 18px auto;
            padding: 0 10px;
        }}

        .about-contact {{
            padding: 16px 12px;
            border-radius: 22px;
        }}

        .about-contact-grid {{
            grid-template-columns: 1fr;
            gap: 14px;
        }}

        .about-left {{
            gap: 0;
        }}

        .brand-row {{
            flex-direction: column;
            align-items: flex-start;
            gap: 12px;
            margin-bottom: 14px;
        }}

        .brand-logo {{
            width: 92px;
            height: 92px;
            min-width: 92px;
            border-width: 3px;
        }}

        .brand-title {{
            font-size: 25px;
            line-height: 1.05;
            margin: 0 0 8px 0;
        }}

        .brand-subtitle {{
            font-size: 13px;
            line-height: 1.7;
            max-width: 100%;
        }}

        .brand-line {{
            width: 120px;
            margin-top: 12px;
        }}

        .about-bottom {{
            grid-template-columns: 1fr;
            gap: 12px;
            margin-top: 6px;
        }}

        .info-box {{
            border-radius: 18px;
            padding: 14px 14px;
        }}

        .info-title {{
            font-size: 18px;
            margin-bottom: 10px;
        }}

        .info-item {{
            font-size: 14px;
            line-height: 1.8;
            gap: 8px;
            align-items: flex-start;
            word-break: break-word;
        }}

        .info-icon {{
            width: 15px;
            height: 15px;
            margin-top: 4px;
        }}

        .social-row {{
            justify-content: center;
            gap: 10px;
            margin-top: 4px;
        }}

        .social-btn {{
            width: 42px;
            height: 42px;
        }}

        .social-icon {{
            width: 19px;
            height: 19px;
        }}

        .about-right {{
            display: block;
        }}

        .form-card {{
            border-radius: 20px;
            padding: 10px;
        }}

        .form-inner {{
            border-radius: 16px;
            padding: 12px;
        }}

        .form-row {{
            grid-template-columns: 1fr;
            gap: 10px;
            margin-bottom: 10px;
        }}

        .input,
        .textarea {{
            border-radius: 12px;
            padding: 12px 13px;
            font-size: 14px;
        }}

        .textarea {{
            min-height: 130px;
            margin-bottom: 12px;
        }}

        .submit-btn {{
            width: 100%;
            min-width: 100%;
            height: 44px;
            font-size: 17px;
        }}

        .feedback-status {{
            font-size: 13px;
            margin-top: 10px;
        }}
    }}
                    
    </style>

    <div class="about-contact-wrap">
        <div class="about-contact">
            <div class="about-contact-grid">

                <div class="about-left">
                    <div class="brand-row">
                        <div class="brand-logo">
                            <img src="{home_logo_src}" alt="Lao Cai Heritage AI">
                        </div>

                        <div>
                            <h2 class="brand-title">{escape(project_title)}</h2>
                            <div class="brand-subtitle">{escape(project_subtitle)}</div>
                            <div class="brand-line"></div>
                        </div>
                    </div>

                    <div class="about-bottom">
                        <div class="info-box">
                            <div class="info-title">Liên hệ</div>
                            <div class="info-item">{vitri_icon_html} Khánh Yên, Lào Cai, Việt Nam</div>
                            <div class="info-item">{email_icon_html} khoagaming999@gmail.com</div>
                            <div class="info-item">{dienthoai_icon_html} 0346 538 917</div>
                        </div>
                        <div class="info-box">
                            <div class="info-title">Mạng xã hội</div>
                            <div class="social-row">
                                <a class="social-btn" href="{facebook_href}" target="_blank" rel="noopener noreferrer">{facebook_icon_html}</a>
                                <a class="social-btn" href="{tiktok_href}" target="_blank" rel="noopener noreferrer">{tiktok_icon_html}</a>
                                <a class="social-btn" href="{youtube_href}" target="_blank" rel="noopener noreferrer">{youtube_icon_html}</a>
                                <a class="social-btn" href="{zalo_href}" target="_blank" rel="noopener noreferrer">{zalo_icon_html}</a>
                            </div>
                        </div>
                    </div>
                </div>

                <div class="about-right">
                    <div class="form-card">
                        <div class="form-inner">

                            <iframe
                                id="feedback_hidden_frame"
                                name="feedback_hidden_frame"
                                style="display:none;"
                            ></iframe>

                            <form
                                id="feedback-form"
                                action="https://script.google.com/macros/s/AKfycbyqh-jHTAxvtDB4aTgA_vJfsHq9PLS_QI-AH-zH0LWVbjY0XRRsMXN7XXTJiUzdRgzbLw/exec"
                                method="POST"
                                target="feedback_hidden_frame"
                                onsubmit="return submitFeedbackForm(this);"
                            >
                                <div class="form-row">
                                    <input class="input" name="name" type="text" placeholder="Họ tên" required>
                                    <input class="input" name="email" type="email" placeholder="Gmail.com" required>
                                </div>

                                <textarea class="textarea" name="message" placeholder="Nội dung..." required></textarea>

                                <input type="hidden" name="to_email" value="khoagaming999@gmail.com">

                                <div class="submit-wrap">
                                    <button id="feedback-submit-btn" class="submit-btn" type="submit">Gửi</button>
                                </div>

                                <div id="feedback-status" class="feedback-status"></div>
                            </form>

                        </div>
                    </div>
                </div>

                <script>
                let feedbackSubmitting = false;
                let feedbackFormRef = null;

                const feedbackFrame = document.getElementById("feedback_hidden_frame");

                if (feedbackFrame) {{
                    feedbackFrame.addEventListener("load", function () {{
                        if (!feedbackSubmitting || !feedbackFormRef) {{
                            return;
                        }}

                        const status = document.getElementById("feedback-status");
                        const btn = document.getElementById("feedback-submit-btn");

                        status.textContent = "Đã gửi góp ý thành công.";
                        feedbackFormRef.reset();

                        btn.disabled = false;
                        btn.style.opacity = "1";
                        btn.style.cursor = "pointer";

                        feedbackSubmitting = false;
                        feedbackFormRef = null;
                    }});
                }}

                function submitFeedbackForm(form) {{
                    const name = (form.name.value || "").trim();
                    const email = (form.email.value || "").trim();
                    const message = (form.message.value || "").trim();

                    const status = document.getElementById("feedback-status");
                    const btn = document.getElementById("feedback-submit-btn");

                    if (!name || !email || !message) {{
                        status.textContent = "Vui lòng nhập đầy đủ thông tin.";
                        return false;
                    }}

                    status.textContent = "Đang gửi...";
                    btn.disabled = true;
                    btn.style.opacity = "0.7";
                    btn.style.cursor = "not-allowed";

                    feedbackSubmitting = true;
                    feedbackFormRef = form;
                    return true;
                }}
                </script>

            </div>
        </div>
    </div>
    """, height=950, scrolling=False)

elif page == "lichtrinh":
    import json
    import math
    import re
    import unicodedata
    from html import escape
    import streamlit.components.v1 as components

    try:
        with open("lichtrinh.json", "r", encoding="utf-8") as f:
            lichtrinh_data = json.load(f)
    except Exception as e:
        st.error(f"Lỗi đọc file lichtrinh.json: {e}")
        st.stop()

    def safe_text(value, fallback="Đang cập nhật"):
        text = str(value).strip() if value is not None else ""
        return escape(text) if text else escape(fallback)

    def normalize_text(value):
        return str(value or "").strip().lower()

    def extract_price(value):
        if value is None:
            return 0
        digits = re.findall(r"\d+", str(value).replace(".", "").replace(",", ""))
        return int("".join(digits)) if digits else 0
    
    def safe_url(value, fallback="https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"):
        text = str(value).strip() if value is not None else ""
        return escape(text if text else fallback, quote=True)

    def normalize_image_name(value):
        text = str(value or "").strip().lower()
        text = text.replace("đ", "d").replace("Đ", "D")
        text = unicodedata.normalize("NFD", text)
        text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
        text = re.sub(r"[^a-z0-9]+", "", text)
        return text

    def itinerary_image(item):
        to_place = str(item.get("to", "")).strip()
        normalized_to = normalize_image_name(to_place)

        local_candidates = [
            f"kho_anh/diem_den/danh_sach/{normalized_to}.jpg",
            f"kho_anh/diem_den/danh_sach/{normalized_to}.png",
            f"kho_anh/diem_den/danh_sach/{normalized_to}.jpeg",
            f"kho_anh/diem_den/danh_sach/{normalized_to}.webp",
        ]

        local_image = image_to_data_uri(local_candidates)
        if local_image:
            return local_image

        for key in ["image", "cover", "thumbnail", "hero_image", "banner_image"]:
            value = str(item.get(key, "")).strip()
            if value:
                return value

        return "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"

    def make_slug(item):
        raw = str(item.get("slug", "")).strip()
        if raw:
            return raw

        text = f"{item.get('from', '')} {item.get('to', '')}"
        text = text.replace("Đ", "D").replace("đ", "d")
        text = unicodedata.normalize("NFD", text)
        text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
        text = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
        return text or "lich-trinh"

    destination_options = ["Tất cả"] + sorted({item.get("to", "") for item in lichtrinh_data if item.get("to")})
    category_options = ["Tất cả"] + sorted({item.get("category", "") for item in lichtrinh_data if item.get("category")})

    total_routes = len(lichtrinh_data)
    lowest_price_all = min((extract_price(item.get("total_price")) for item in lichtrinh_data), default=0)
    hero_title = "Lịch trình du lịch Lào Cai"
    hero_subtitle = "Plan your trip smarter — chọn điểm đến, phong cách du lịch & ngân sách để nhận gợi ý tối ưu."

    st.markdown("""
    <style>
    .lt-hero-section{
        position: relative;
        height: 500px;
        background:
            linear-gradient(180deg, rgba(5, 18, 38, 0.20) 0%, rgba(6, 22, 43, 0.55) 100%),
            url('https://booking.muongthanh.com/upload_images/images/H%60/dinh-nui-fansipan.jpg');
        background-size: cover;
        background-position: center;
        display: flex;
        align-items: center;
        justify-content: center;
        text-align: center;
        padding: 0 24px 140px 24px;
        overflow: hidden;
    }

    .lt-hero-section::after{
        content:"";
        position:absolute;
        inset:0;
        background: radial-gradient(circle at top, rgba(255,255,255,0.14), transparent 38%);
        pointer-events:none;
    }

    .lt-hero-inner{
        position: relative;
        z-index: 2;
        max-width: 980px;
        color: #ffffff;
    }

    .lt-hero-kicker{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        padding: 10px 18px;
        border-radius: 999px;
        background: rgba(255,255,255,0.16);
        border: 1px solid rgba(255,255,255,0.28);
        backdrop-filter: blur(8px);
        -webkit-backdrop-filter: blur(8px);
        font-size: 14px;
        font-weight: 800;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        margin-bottom: 18px;
    }

    .lt-hero-title{
        font-size: 56px;
        font-weight: 900;
        line-height: 1.1;
        margin: 0;
        text-shadow: 0 6px 22px rgba(0,0,0,0.28);
    }

    .lt-hero-desc{
        max-width: 760px;
        margin: 16px auto 0 auto;
        font-size: 18px;
        line-height: 1.75;
        color: rgba(255,255,255,0.96);
        text-shadow: 0 4px 16px rgba(0,0,0,0.24);
    }

    .lt-hero-stats {
        display: flex;
        justify-content: center;
        gap: 14px;
        flex-wrap: wrap;
        margin-top: 24px;
    }

    .lt-hero-stat {
        min-width: 168px;
        padding: 14px 18px;
        border-radius: 18px;
        background: rgba(255,255,255,0.14);
        border: 1px solid rgba(255,255,255,0.24);
        backdrop-filter: blur(8px);
        -webkit-backdrop-filter: blur(8px);
        box-shadow: 0 14px 30px rgba(0,0,0,0.10);
        position: relative;
        overflow: hidden;
        opacity: 0;
        transform: translateY(18px);
        animation: ltStatUp 0.75s ease-out forwards;
        transition: transform 0.28s ease, box-shadow 0.28s ease, border-color 0.28s ease;
    }

    .lt-hero-stat::before {
        content: "";
        position: absolute;
        inset: 0;
        background: linear-gradient(
            135deg,
            rgba(255,255,255,0.22) 0%,
            rgba(255,255,255,0.06) 45%,
            rgba(255,255,255,0.02) 100%
        );
        pointer-events: none;
    }

    .lt-hero-stat::after {
        content: "";
        position: absolute;
        top: -28px;
        left: -36px;
        width: 90px;
        height: 90px;
        border-radius: 50%;
        background: rgba(255,255,255,0.10);
        filter: blur(2px);
        pointer-events: none;
    }

    .lt-hero-stat:nth-child(1) { animation-delay: 0.10s; }
    .lt-hero-stat:nth-child(2) { animation-delay: 0.22s; }
    .lt-hero-stat:nth-child(3) { animation-delay: 0.34s; }

    .lt-hero-stat:hover {
        transform: translateY(-8px) scale(1.03);
        box-shadow: 0 22px 42px rgba(0,0,0,0.16);
        border-color: rgba(255,255,255,0.42);
    }

    .lt-hero-stat-value {
        position: relative;
        z-index: 2;
        font-size: 24px;
        font-weight: 900;
        color: #ffffff;
        line-height: 1.15;
        text-shadow: 0 4px 12px rgba(0,0,0,0.20);
        transition: transform 0.28s ease, text-shadow 0.28s ease;
    }

    .lt-hero-stat-label {
        position: relative;
        z-index: 2;
        font-size: 12px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.08em;
        color: rgba(255,255,255,0.82);
        margin-top: 6px;
        transition: color 0.28s ease;
    }

    .lt-hero-stat:hover .lt-hero-stat-value {
        transform: translateY(-1px);
        text-shadow: 0 8px 18px rgba(0,0,0,0.28);
    }

    .lt-hero-stat:hover .lt-hero-stat-label {
        color: rgba(255,255,255,0.98);
    }

    @keyframes ltStatUp {
        from {
            opacity: 0;
            transform: translateY(18px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }

    .lt-form-intro{
        text-align: center;
        margin: -190px auto 0 auto;
        position: relative;
        z-index: 10;
    }

    .lt-form-title{
        font-size: 30px;
        font-weight: 900;
        color: #0f172a;
        margin: 14px 0 8px 0;
    }

    .lt-form-subtitle{
        font-size: 16px;
        line-height: 1.7;
        color: #475569;
        margin: 0 auto;
        max-width: 760px;
    }

    div[data-testid="stForm"]{
        position: relative;
        z-index: 20;
        margin-top: -160px;
        background: #ffffff;
        border: 1px solid rgba(15, 23, 42, 0.07);
        border-radius: 15px;
        padding: 20px 24px 22px 24px;
        box-shadow: 0 24px 54px rgba(15,23,42,0.12);
    }

    div[data-testid="stForm"] [data-testid="stHorizontalBlock"]{
        align-items: flex-end !important;
        gap: 14px !important;
    }

    div[data-testid="stColumn"]{
        display: flex;
        align-items: stretch;
    }

    div[data-testid="stColumn"] > div{
        width: 100%;
    }

    div[data-testid="stSelectbox"]{
        width: 100%;
        margin-bottom: 0 !important;
    }

    div[data-testid="stSelectbox"] label{
        min-height: 22px !important;
        display: flex !important;
        align-items: center !important;
        margin-bottom: 8px !important;
        line-height: 1.2 !important;
    }

    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
    div[data-testid="stTextInput"] input{
        min-height: 52px !important;
        border-radius: 16px !important;
        border: 1px solid #dbe4ef !important;
        background: #f8fafc !important;
        box-shadow: none !important;
        font-size: 15px !important;
        display: flex !important;
        align-items: center !important;
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }

    div[data-testid="stFormSubmitButton"]{
        width: 100%;
        display: flex;
        align-items: flex-end;
        height: 100%;
    }

    div[data-testid="stFormSubmitButton"] button{
        width: 100%;
        min-height: 46px !important;
        margin-top: 0 !important;
        border-radius: 14px !important;
    }

    div[data-testid="stFormSubmitButton"] button{
        width: 100%;
        min-height: 52px;
        border: none !important;
        border-radius: 16px !important;
        background: linear-gradient(135deg, #1565c0 0%, #2b7fd3 100%) !important;
        color: #ffffff !important;
        font-size: 16px !important;
        font-weight: 800 !important;
        letter-spacing: 0.02em;
        box-shadow: 0 16px 28px rgba(21,101,192,0.24) !important;
    }

    div[data-testid="stFormSubmitButton"] button:hover{
        background: linear-gradient(135deg, #0f4f9a 0%, #1565c0 100%) !important;
    }
                
    .lt-form-overlap{
        height: 0;
        margin-top: -145px;
    }

    @media (max-width: 991px){
        .lt-hero-section{
            height: 390px;
            padding-bottom: 82px;
        }

        .lt-hero-title{
            font-size: 44px;
        }

        .lt-form-title{
            font-size: 26px;
        }
    }

    @media (max-width: 767px){
        .lt-hero-section{
            height: 360px;
            padding: 0 16px 74px 16px;
        }

        .lt-hero-title{
            font-size: 34px;
        }

        .lt-hero-desc{
            font-size: 15px;
        }

        .lt-hero-stat{
            min-width: calc(50% - 8px);
            padding: 12px 14px;
        }

        .lt-form-intro{
            margin-top: -72px;
        }

        .lt-form-title{
            font-size: 23px;
        }

        .lt-form-subtitle,
        .lt-summary-text{
            font-size: 15px;
        }

        div[data-testid="stForm"]{
            padding: 18px 16px 16px 16px;
            border-radius: 20px;
        }
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <section class="lt-hero-section">
        <div class="lt-hero-inner">
            <h1 class="lt-hero-title">{escape(hero_title)}</h1>
            <p class="lt-hero-desc">{escape(hero_subtitle)}</p>
            <div class="lt-hero-stats">
                <div class="lt-hero-stat">
                    <div class="lt-hero-stat-value">{total_routes}</div>
                    <div class="lt-hero-stat-label">Hành trình đề xuất</div>
                </div>
                <div class="lt-hero-stat">
                    <div class="lt-hero-stat-value">{len(destination_options) - 1}</div>
                    <div class="lt-hero-stat-label">Điểm đến tiêu biểu</div>
                </div>
                <div class="lt-hero-stat">
                    <div class="lt-hero-stat-value">{lowest_price_all:,}đ</div>
                    <div class="lt-hero-stat-label">Giá tham khảo</div>
                </div>
            </div>
        </div>
    </section>
    """, unsafe_allow_html=True)

    _, center_col, _ = st.columns([0.12, 0.76, 0.12])
    with center_col:
        st.markdown('<div class="lt-form-overlap"></div>', unsafe_allow_html=True)

        with st.form("lichtrinh_filter_form", border=False):
            c1, c2, c3, c4, c5 = st.columns([1.55, 1.15, 1.15, 1.15, 0.85])

            with c1:
                destination_query = st.selectbox("Bạn muốn đến đâu?", destination_options, index=0, key="lichtrinh_destination")

            with c2:
                departure_place = st.selectbox("Nơi khởi hành", ["TP. Lào Cai"], key="lichtrinh_departure")

            with c3:
                selected_category = st.selectbox("Loại hình", category_options, index=0, key="lichtrinh_category")

            with c4:
                sort_option = st.selectbox("Sắp xếp chi phí", ["Mặc định", "Chi phí tăng dần", "Chi phí giảm dần"], index=0, key="lichtrinh_sort")

            with c5:
                st.markdown("<div style='height: 28px;'></div>", unsafe_allow_html=True)
                st.form_submit_button("Tìm tour", use_container_width=True)

        filtered_data = [
            item for item in lichtrinh_data
            if normalize_text(item.get("from")) == normalize_text(departure_place)
        ]

        if selected_category != "Tất cả":
            filtered_data = [
                item for item in filtered_data
                if normalize_text(item.get("category")) == normalize_text(selected_category)
            ]

        if destination_query != "Tất cả":
            keyword = normalize_text(destination_query)
            filtered_data = [
                item for item in filtered_data
                if keyword in normalize_text(item.get("to"))
            ]

    if sort_option == "Chi phí tăng dần":
        filtered_data = sorted(filtered_data, key=lambda x: extract_price(x.get("total_price")))
    elif sort_option == "Chi phí giảm dần":
        filtered_data = sorted(filtered_data, key=lambda x: extract_price(x.get("total_price")), reverse=True)

    result_count = len(filtered_data)
    cheapest_filtered = min((extract_price(item.get("total_price")) for item in filtered_data), default=0)

    selected_destination_label = destination_query if destination_query != "Tất cả" else "tất cả điểm đến"
    selected_category_label = selected_category if selected_category != "Tất cả" else "mọi loại hình"


    card_css = """
    <style>
    body{
        margin:0;
        background:transparent;
        font-family: Inter, Arial, sans-serif;
    }

    .lt-results-wrap{
        max-width: 1320px;
        margin: 18px auto 42px auto;
        padding: 0 4px;
        box-sizing: border-box;
    }

    .lt-grid{
        display:grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap:24px;
        box-sizing:border-box;
    }

    .lt-card{
        position:relative;
        min-height:430px;
        perspective:1600px;
    }

    .lt-card-shell{
        position:relative;
        width:100%;
        height:100%;
        min-height:430px;
        transform-style:preserve-3d;
        transition:transform 0.82s cubic-bezier(0.22, 0.7, 0.18, 1);
    }

    .lt-card:hover .lt-card-shell{
        transform:rotateY(-180deg);
    }

    .lt-card-face{
        position:absolute;
        inset:0;
        border-radius:26px;
        overflow:hidden;
        backface-visibility:hidden;
        -webkit-backface-visibility:hidden;
    }

    .lt-card-front{
        position:relative;
        background:linear-gradient(180deg, #ffffff 0%, #fbfdff 100%);
        border:1px solid #e5edf6;
        padding:24px 22px 22px 22px;
        display:flex;
        flex-direction:column;
        box-shadow:0 18px 38px rgba(15,23,42,0.08);
    }

    .lt-card-front::before{
        content:"";
        position:absolute;
        top:0;
        left:0;
        right:0;
        height:5px;
        background:linear-gradient(90deg, #1565c0 0%, #2b7fd3 52%, #43b0f1 100%);
    }

    .lt-card-back{
        transform:rotateY(180deg);
        background-size:cover;
        background-position:center;
        background-repeat:no-repeat;
        box-shadow:0 18px 38px rgba(15,23,42,0.12);
    }

    .lt-card-back::before{
        content:"";
        position:absolute;
        inset:0;
        background:linear-gradient(180deg, rgba(5,18,38,0.10) 0%, rgba(5,18,38,0.78) 100%);
    }

    .lt-back-inner{
        position:relative;
        z-index:2;
        height:100%;
        display:flex;
        flex-direction:column;
        justify-content:center;
        align-items:center;        
        padding:4px 18px 20px 18px;
        color:#ffffff;
    }

    .lt-back-badge{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        align-self:flex-start;
        margin-bottom:12px;
        padding:8px 13px;
        border-radius:999px;
        background:rgba(255,255,255,0.16);
        border:1px solid rgba(255,255,255,0.28);
        font-size:13px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.06em;
    }

    .lt-back-route{
        font-size:26px;
        font-weight:900;
        line-height:1.25;
        margin-bottom:10px;
        text-transform:uppercase;
        text-shadow:0 4px 14px rgba(0,0,0,0.28);
    }

    .lt-back-desc{
        font-size:15px;
        line-height:1.75;
        color:rgba(255,255,255,0.92);
        margin-bottom:18px;
        max-width:92%;
    }

    .lt-view-btn{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        align-self:center;
        min-width:150px;
        height:48px;
        border:none;
        border-radius:14px;
        background:#ffffff;
        color:#0f172a;
        font-size:15px;
        font-weight:900;
        cursor:pointer;
        box-shadow:0 14px 26px rgba(0,0,0,0.18);
        transition:transform 0.2s ease, box-shadow 0.2s ease;
    }

    .lt-view-btn:hover{
        transform:translateY(-2px);
        box-shadow:0 18px 32px rgba(0,0,0,0.24);
    }

    .lt-card-top{
        display:flex;
        align-items:flex-start;
        justify-content:space-between;
        gap:14px;
        margin-bottom:18px;
    }

    .lt-route{
        font-size:20px;
        font-weight:900;
        text-transform:uppercase;
        letter-spacing:0.01em;
        line-height:1.3;
        color:#0f172a;
    }

    .lt-badge{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        padding:8px 12px;
        border-radius:999px;
        background:#ecf5ff;
        color:#1565c0;
        font-size:15px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.06em;
        white-space:nowrap;
        border:1px solid #cfe5ff;
    }

    .lt-info-grid{
        display:grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap:12px;
    }

    .lt-info-item{
        background:#f1f5f9;
        border:1px solid #dde6ef;
        border-radius:18px;
        padding:14px;
        min-height:88px;
        box-sizing:border-box;
        display:flex;
        flex-direction:column;
        align-items:center;
        justify-content:center;
        text-align:center;
    }

    .lt-info-label{
        font-size:17px;
        font-weight:800;
        text-transform:uppercase;
        color:#64748b;
        margin-bottom:8px;
        text-align:center;
    }

    .lt-info-value{
        font-size:16px;
        line-height:1.6;
        font-weight:700;
        color:#0f172a;
        word-break:break-word;
        text-align:center;
    }

    .lt-hotel-box{
        margin-top:14px;
        background:linear-gradient(180deg, #fff7ed 0%, #fffaf5 100%);
        border:1px solid #fed7aa;
        border-radius:18px;
        padding:15px 16px;
    }

    .lt-hotel-title{
        font-size:12px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.08em;
        color:#c2410c;
        margin-bottom:8px;
    }

    .lt-hotel-value{
        font-size:14px;
        line-height:1.7;
        font-weight:700;
        color:#7c2d12;
        word-break:break-word;
    }

    .lt-divider{
        height:1px;
        background:#e7eef6;
        margin:18px 0 16px 0;
    }

    .lt-bottom{
        display:flex;
        justify-content:space-between;
        align-items:flex-end;
        gap:12px;
        margin-top:auto;
    }

    .lt-bottom-note{
        font-size:13px;
        line-height:1.65;
        color:#64748b;
        max-width:55%;
    }

    .lt-price-block{
        text-align:right;
        flex-shrink:0;
    }

    .lt-price-caption{
        font-size:12px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.08em;
        color:#64748b;
        margin-bottom:6px;
    }

    .lt-price-value{
        font-size:17px;
        font-weight:900;
        line-height:1.2;
        color:#ef4444;
        word-break:break-word;
        overflow-wrap:anywhere;
    }

    .lt-price-sub{
        font-size:13px;
        line-height:1.55;
        color:#94a3b8;
        margin-top:4px;
    }

    .lt-empty{
        background:linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
        border:1px dashed #cbd5e1;
        border-radius:26px;
        padding:42px 26px;
        text-align:center;
        color:#334155;
        box-shadow:0 14px 36px rgba(15,23,42,0.05);
    }

    .lt-empty-title{
        font-size:26px;
        font-weight:900;
        color:#0f172a;
        margin-bottom:10px;
    }

    .lt-empty-text{
        max-width:620px;
        margin:0 auto;
        font-size:16px;
        line-height:1.8;
        color:#475569;
    }

    @media (max-width: 1180px){
        .lt-grid{
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }
    }

    @media (max-width: 760px){
        .lt-results-wrap{
            padding:0;
        }

        .lt-grid{
            grid-template-columns:1fr;
            gap:18px;
        }

        .lt-card,
        .lt-card-shell{
            min-height:420px;
        }

        .lt-card-top,
        .lt-bottom{
            flex-direction:column;
            align-items:flex-start;
        }

        .lt-route{
            font-size:17px;
        }

        .lt-info-grid{
            grid-template-columns:1fr;
        }

        .lt-bottom-note,
        .lt-price-block{
            max-width:100%;
            width:100%;
            text-align:left;
        }

        .lt-price-value{
            font-size:26px;
        }

        .lt-back-route{
            font-size:22px;
        }

        .lt-empty-title{
            font-size:22px;
        }

        .lt-empty-text{
            font-size:15px;
        }
    }
    </style>
    """

    cards_html = '<div class="lt-results-wrap">'

    if filtered_data:
        cards_html += '<div class="lt-grid">'
        for item in filtered_data:
            total_price = safe_text(item.get("total_price"), "Liên hệ")
            transport_price = safe_text(item.get("transport_price"), "Đang cập nhật")
            ticket_price = safe_text(item.get("ticket_price"), "0đ")
            hotel_name = safe_text(item.get("hotel_name"), "Đang cập nhật")
            hotel_price = safe_text(item.get("hotel_price"), "Liên hệ")
            category = safe_text(item.get("category"), "Lịch trình")
            route_text = f"{safe_text(item.get('from'))} → {safe_text(item.get('to'))}"
            cover_image = safe_url(itinerary_image(item))
            slug_value = make_slug(item)

            hover_desc = safe_text(
                item.get("note") or item.get("short_desc"),
                "Nhấn để xem chi tiết lịch trình, chi phí tham khảo và thông tin lưu trú."
            )

            cards_html += f"""
            <div class="lt-card">
                <div class="lt-card-shell">

                    <div class="lt-card-face lt-card-front">
                        <div class="lt-card-top">
                            <div class="lt-route">{route_text}</div>
                            <div class="lt-badge">{category}</div>
                        </div>

                        <div class="lt-info-grid">
                            <div class="lt-info-item">
                                <div class="lt-info-label">Quãng đường</div>
                                <div class="lt-info-value">{safe_text(item.get('distance_km'))}</div>
                            </div>
                            <div class="lt-info-item">
                                <div class="lt-info-label">Thời gian</div>
                                <div class="lt-info-value">{safe_text(item.get('time_estimate'))}</div>
                            </div>
                            <div class="lt-info-item">
                                <div class="lt-info-label">Phương tiện</div>
                                <div class="lt-info-value">{safe_text(item.get('transport_name'))}</div>
                            </div>
                            <div class="lt-info-item">
                                <div class="lt-info-label">Giá vé tham quan</div>
                                <div class="lt-info-value">{ticket_price}</div>
                            </div>
                        </div>

                        <div class="lt-hotel-box">
                            <div class="lt-hotel-title">Khách sạn / Nhà nghỉ gợi ý</div>
                            <div class="lt-hotel-value">{hotel_name} • {hotel_price}</div>
                        </div>

                        <div class="lt-divider"></div>

                        <div class="lt-bottom">
                            <div class="lt-bottom-note">
                                Giá xe tham khảo: <strong>{transport_price}</strong><br>
                                Hành trình hiển thị theo bộ lọc bạn đã chọn.
                            </div>
                            <div class="lt-price-block">
                                <div class="lt-price-caption">Giá từ / khách</div>
                                <div class="lt-price-value">{total_price}</div>
                                <div class="lt-price-sub">Đã gồm chi phí tổng quan</div>
                            </div>
                        </div>
                    </div>

                    <div class="lt-card-face lt-card-back" style="background-image:url('{cover_image}');">
                        <div class="lt-back-inner">
                            <div class="lt-back-badge">{category}</div>
                            <div class="lt-back-route">{route_text}</div>
                            <div class="lt-back-desc">{hover_desc}</div>
                            <a class="lt-view-btn"
                                href="{build_app_url('lichtrinh_detail', slug=slug_value)}"
                                target="_blank"
                                rel="noopener noreferrer">
                                    Xem chi tiết
                            </a>
                        </div>
                    </div>

                </div>
            </div>
            """

        cards_html += "</div>"

        cards_html += """
        """

        rows = math.ceil(result_count / 3) if result_count else 1
        height = max(520, rows * 580)
    else:
        cards_html += """
        <div class="lt-empty">
            <div class="lt-empty-title">Chưa tìm thấy lịch trình phù hợp</div>
            <div class="lt-empty-text">
                Bộ lọc hiện tại đang khá hẹp nên chưa có dữ liệu khớp. Bạn hãy chuyển điểm đến hoặc loại hình về
                <strong>Tất cả</strong> để xem lại toàn bộ gợi ý đang có trong hệ thống.
            </div>
        </div>
        """
        height = 320

    cards_html += "</div>"
    components.html(card_css + cards_html, height=height, scrolling=False)

    # Lưu ý
    components.html("""
    <div style="
        max-width: 1320px;
        margin: 18px auto 60px auto;
        padding: 0 4px;
        box-sizing: border-box;
        font-family: Inter, Arial, sans-serif;
    ">
        <div style="
            display: grid;
            grid-template-columns: 280px 1fr;
            align-items: stretch;
            background: #ffffff;
            border: 1px solid #e5e7eb;
            box-shadow: 0 10px 28px rgba(15,23,42,0.08);
            overflow: hidden;
        ">
            <!-- Khối trái -->
            <div style="
                position: relative;
                min-height: 180px;
                background: linear-gradient(135deg, #8b0000 0%, #c1121f 100%);
                display: flex;
                align-items: center;
                justify-content: center;
                padding: 20px;
                box-sizing: border-box;
            ">
                <div style="
                    position: relative;
                    width: 100%;
                    height: 100%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                ">
                    <div style="
                        width: 135px;
                        height: 68px;
                        border: 2px solid rgba(255,255,255,0.95);
                        color: #ffffff;
                        font-size: 25px;
                        font-weight: 800;
                        display: flex;
                        align-items: center;
                        justify-content: center;
                        transform: rotate(-6deg);
                        box-shadow: 0 10px 20px rgba(0,0,0,0.18);
                    ">
                        LƯU Ý
                    </div>

                    <div style="
                        position: absolute;
                        right: 8px;
                        bottom: 18px;
                        font-size: 54px;
                        line-height: 1;
                        filter: drop-shadow(0 6px 10px rgba(0,0,0,0.22));
                    ">
                        📢
                    </div>
                </div>
            </div>

            <!-- Khối phải -->
            <div style="
                padding: 18px 30px;
                box-sizing: border-box;
                background: #ffffff;
                min-height: 180px;
            ">
                <div style="
                    font-size: 25px;
                    font-weight: 800;
                    color: #111827;
                    line-height: 1.45;
                    margin-bottom: 10px;
                ">
                    LƯU Ý KHI THAM KHẢO LỊCH TRÌNH DU LỊCH LÀO CAI
                </div>

                <div style="
                    font-size: 14px;
                    color: #64748b;
                    margin-bottom: 14px;
                ">
                    Cập nhật hệ thống • Tham khảo trước khi khởi hành
                </div>

                <div style="
                    font-size: 12px;
                    color: #374151;
                    line-height: 1.9;
                    margin-bottom: 14px;
                ">
                    Lịch trình được xây dựng nhằm hỗ trợ người dùng tham khảo trong quá trình tìm hiểu và lên kế hoạch chuyến đi. Thời gian di chuyển, chi phí dự kiến và một số thông tin liên quan có thể thay đổi tùy theo điều kiện thời tiết, tình trạng vận hành thực tế, thời điểm lưu trú và phương tiện được lựa chọn.

Người dùng nên chủ động kiểm tra lại thông tin khách sạn, vé tham quan, phương tiện di chuyển và chuẩn bị đầy đủ giấy tờ cá nhân trước khi bắt đầu hành trình để bảo đảm chuyến đi diễn ra thuận tiện và phù hợp với thực tế.
                </div>

                <div style="
                    font-size: 15px;
                    color: #64748b;
                    line-height: 1.8;
                ">
                    Chuyên mục:
                    <span style="font-weight: 800; color: #111827;">Lưu ý lịch trình</span>
                </div>
            </div>
        </div>
    </div>
    """, height=230, scrolling=False)

elif page == "lichtrinh_detail":
    import os
    import json
    import re
    import base64
    import unicodedata
    from html import escape
    from string import Template
    from urllib.parse import quote_plus
    from pathlib import Path
    import fitz
    import streamlit.components.v1 as components

    try:
        with open("lichtrinh.json", "r", encoding="utf-8") as f:
            lichtrinh_data = json.load(f)
    except Exception as e:
        st.error(f"Lỗi đọc file lichtrinh.json: {e}")
        st.stop()

    def clean_text(value):
        if value is None:
            return ""
        return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()

    def normalize_text(value):
        return clean_text(value).lower()

    def safe_text(value, fallback="Đang cập nhật"):
        text = clean_text(value)
        return escape(text if text else fallback)

    def safe_url(value, fallback="https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"):
        text = clean_text(value)
        return escape(text if text else fallback, quote=True)

    def make_slug(item):
        raw = clean_text(item.get("slug"))
        if raw:
            return raw

        text = f"{item.get('from', '')} {item.get('to', '')}"
        text = text.replace("Đ", "D").replace("đ", "d")
        text = unicodedata.normalize("NFD", text)
        text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
        text = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
        return text or "lich-trinh"

    def read_text_file_any(path_str):
        path = Path(path_str)
        if not path.exists():
            return ""

        suffix = path.suffix.lower()

        try:
            if suffix == ".txt":
                return path.read_text(encoding="utf-8").strip()

            elif suffix == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(path))
                    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    return "\n".join(parts).strip()
                except Exception:
                    return ""

            elif suffix == ".pdf":
                try:
                    pdf = fitz.open(str(path))
                    texts = []
                    for page in pdf:
                        txt = page.get_text("text").strip()
                        if txt:
                            texts.append(txt)
                    pdf.close()
                    return "\n".join(texts).strip()
                except Exception:
                    return ""
        except Exception:
            return ""

        return ""

    def pick_intro_text_file(route):
        folder = Path("text_lichtrinh_gioithieu")

        def normalize_name(value):
            text = clean_text(value).lower()
            text = text.replace("đ", "d").replace("Đ", "D")
            text = unicodedata.normalize("NFD", text)
            text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
            text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
            return text

        to_name = clean_text(route.get("to"))
        if not to_name:
            return ""

        base_name = normalize_name(to_name)

        candidates = [
            folder / f"{base_name}.txt",
            folder / f"{base_name}.docx",
            folder / f"{base_name}.pdf",
        ]

        for fp in candidates:
            if fp.exists():
                return read_text_file_any(fp)

        return ""
    
    

    def get_images(route):
        import random

        fallback = "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"

        def normalize_folder_name(value):
            text = clean_text(value).lower()
            text = text.replace("đ", "d").replace("Đ", "D")
            text = unicodedata.normalize("NFD", text)
            text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
            text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
            return text

        def list_local_images(folder):
            if not folder or not os.path.isdir(folder):
                return []

            exts = (".png", ".jpg", ".jpeg", ".webp")
            files = []

            for name in os.listdir(folder):
                if name.lower().endswith(exts):
                    full_path = os.path.join(folder, name)
                    src = image_to_data_uri([full_path])
                    if src:
                        files.append(src)

            return files

        # Ảnh to: lấy từ link trong JSON nếu có
        main_image = ""
        for key in ["image", "cover", "thumbnail", "hero_image", "banner_image"]:
            val = clean_text(route.get(key))
            if val:
                main_image = safe_url(val, fallback)
                break

        # Nếu JSON chưa có ảnh to thì lấy luôn ảnh đầu trong thư mục local
        folder_name = normalize_folder_name(route.get("to", ""))
        small_folder = f"kho_anh/diem_den/tri_tiet/{folder_name}/thuyet_minh"
        local_images = list_local_images(small_folder)

        if not main_image:
            main_image = local_images[0] if local_images else fallback

        # Bỏ ảnh trùng với ảnh to
        small_images = [img for img in local_images if img != main_image]

        # Lấy 4 ảnh nhỏ ngẫu nhiên
        if len(small_images) >= 4:
            small_images = random.sample(small_images, 4)
        elif len(small_images) > 0:
            while len(small_images) < 4:
                small_images.append(random.choice(small_images))
        else:
            small_images = [main_image, main_image, main_image, main_image]

        return [main_image] + small_images[:4]

    def split_paragraphs(*values):
        results = []
        seen = set()

        for value in values:
            text = clean_text(value)
            if not text:
                continue

            parts = [p.strip() for p in text.split("\n") if p.strip()]
            if not parts:
                parts = [text]

            for part in parts:
                normalized = " ".join(part.split()).strip()
                if normalized and normalized.lower() not in seen:
                    seen.add(normalized.lower())
                    results.append(normalized)
        return results

    def build_highlights(route):
        highlights = []

        if clean_text(route.get("distance_km")):
            highlights.append(f"Quãng đường dự kiến: {clean_text(route.get('distance_km'))}.")
        if clean_text(route.get("time_estimate")):
            highlights.append(f"Thời gian di chuyển ước tính: {clean_text(route.get('time_estimate'))}.")
        if clean_text(route.get("transport_name")):
            highlights.append(f"Phương tiện gợi ý: {clean_text(route.get('transport_name'))}.")
        if clean_text(route.get("ticket_price")):
            highlights.append(f"Vé tham quan tham khảo: {clean_text(route.get('ticket_price'))}.")
        if clean_text(route.get("hotel_name")):
            highlights.append(f"Lưu trú gợi ý: {clean_text(route.get('hotel_name'))}.")
        if clean_text(route.get("note")):
            note_text = clean_text(route.get("note"))
            if len(note_text) > 140:
                note_text = note_text[:140].rsplit(" ", 1)[0] + "..."
            highlights.append(note_text)

        if not highlights:
            highlights = [
                "Lịch trình đang được cập nhật thêm điểm nổi bật.",
                "Bạn có thể xem chi phí, phương tiện và lưu trú tham khảo bên dưới."
            ]

        return highlights[:6]

    def build_map_url(route):
        for key in ["map_embed", "iframe_map", "google_map", "map_url", "location_map"]:
            raw = clean_text(route.get(key))
            if not raw:
                continue

            if raw.startswith("<iframe"):
                match = re.search(r'''src=["']([^"']+)["']''', raw)
                if match:
                    return escape(match.group(1), quote=True)

            if raw.startswith("http://") or raw.startswith("https://"):
                return escape(raw, quote=True)

        lat = clean_text(route.get("lat"))
        lng = clean_text(route.get("lng"))
        if lat and lng:
            return escape(f"https://www.google.com/maps?q={lat},{lng}&output=embed", quote=True)

        hotel_name = clean_text(route.get("hotel_name"))
        destination = clean_text(route.get("to"))
        area_query = " ".join([x for x in [hotel_name, destination, "Lào Cai"] if x]).strip()

        if not area_query:
            area_query = "Lào Cai"

        return escape(f"https://www.google.com/maps?q={quote_plus(area_query)}&output=embed", quote=True)

    slug = clean_text(st.query_params.get("slug", ""))

    selected_route = None
    for item in lichtrinh_data:
        if normalize_text(make_slug(item)) == normalize_text(slug):
            selected_route = item
            break

    if not selected_route:
        st.error("Không tìm thấy lịch trình.")
        st.stop()

    from_place = clean_text(selected_route.get("from")) or "Điểm khởi hành"
    to_place = clean_text(selected_route.get("to")) or "Điểm đến"
    route_title = f"{from_place} - {to_place}"
    category = clean_text(selected_route.get("category")) or "Lịch trình"
    route_code = clean_text(selected_route.get("code")) or clean_text(selected_route.get("tour_code")) or clean_text(selected_route.get("id")) or slug.upper()
    note_text = clean_text(selected_route.get("note")) or clean_text(selected_route.get("short_desc")) or "Lịch trình này đang được cập nhật thêm phần mô tả chi tiết."
    total_price = clean_text(selected_route.get("total_price")) or "Liên hệ"
    transport_price = clean_text(selected_route.get("transport_price")) or "Đang cập nhật"
    ticket_price = clean_text(selected_route.get("ticket_price")) or "0đ"
    transport_name = clean_text(selected_route.get("transport_name")) or "Đang cập nhật"
    distance_km = clean_text(selected_route.get("distance_km")) or "Đang cập nhật"
    time_estimate = clean_text(selected_route.get("time_estimate")) or "Đang cập nhật"
    hotel_name = clean_text(selected_route.get("hotel_name")) or "Đang cập nhật"
    hotel_price = clean_text(selected_route.get("hotel_price")) or "Liên hệ"

    intro_file_text = pick_intro_text_file(selected_route)

    if intro_file_text:
        paragraphs = split_paragraphs(intro_file_text)
    else:
        paragraphs = split_paragraphs(selected_route.get("note"), selected_route.get("short_desc"))

    if not paragraphs:
        paragraphs = [
            f"Hành trình từ {from_place} đến {to_place} phù hợp cho du khách muốn khám phá vẻ đẹp địa phương một cách thuận tiện.",
            "Lịch trình hiện đang được hệ thống tổng hợp với các thông tin tham khảo về phương tiện, thời gian, chi phí và lưu trú."
        ]

    highlights = build_highlights(selected_route)
    images = get_images(selected_route)
    map_url = build_map_url(selected_route)

    highlight_html = "".join(f"<li>{escape(item)}</li>" for item in highlights)

    intro_html = "".join(
        f"<p>{escape(paragraph)}</p>"
        for paragraph in paragraphs
        if str(paragraph).strip()
    )

    html_template = Template("""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1.0" />
        <style>
            * {
                box-sizing: border-box;
            }

            html, body {
                margin: 0;
                padding: 0;
                background: #eef2e6;
                font-family: Arial, sans-serif;
                color: #0f172a;
            }

            .ltd2-page {
                width: 100%;
                padding: 18px 16px 40px 16px;
                background: #eef2e6;
            }

            .ltd2-wrap {
                max-width: 1320px;
                margin: 0 auto;
            }

            .ltd2-title {
                font-size: 26px;
                line-height: 1.25;
                font-weight: 900;
                color: #0b4f8a;
                margin: 0 0 8px 0;
            }

            .ltd2-meta {
                display: flex;
                flex-wrap: wrap;
                gap: 16px;
                align-items: center;
                margin-bottom: 16px;
                color: #475569;
                font-size: 14px;
                line-height: 1.6;
            }

            .ltd2-meta strong {
                color: #0f172a;
            }

            .ltd2-gallery {
                display: grid;
                grid-template-columns: 1.55fr 1fr;
                gap: 10px;
                margin-bottom: 16px;
            }

            .ltd2-main-photo {
                min-height: 300px;
                border-radius: 10px;
                background-size: cover;
                background-position: center;
                position: relative;
                overflow: hidden;
                box-shadow: 0 12px 24px rgba(15, 23, 42, 0.10);
            }

            .ltd2-main-photo::before {
                content: "";
                position: absolute;
                inset: 0;
                background: linear-gradient(180deg, rgba(0,0,0,0.05) 0%, rgba(0,0,0,0.25) 100%);
            }

            .ltd2-main-badge {
                position: absolute;
                top: 10px;
                left: 10px;
                z-index: 2;
                background: #e53935;
                color: #ffffff;
                font-size: 12px;
                font-weight: 800;
                padding: 6px 10px;
                border-radius: 6px;
            }

            .ltd2-side-grid {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 10px;
            }

            .ltd2-side-photo {
                min-height: 145px;
                border-radius: 10px;
                background-size: cover;
                background-position: center;
                box-shadow: 0 10px 20px rgba(15, 23, 42, 0.08);
            }

            .ltd2-top-grid {
                display: grid;
                grid-template-columns: 1.7fr 0.8fr;
                gap: 16px;
                margin-bottom: 16px;
            }

            .ltd2-bottom-grid {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 16px;
            }

            .ltd2-card {
                background: #ffffff;
                border: 1px solid #dfe6d6;
                border-radius: 16px;
                padding: 18px;
                box-shadow: 0 10px 22px rgba(15, 23, 42, 0.06);
            }

            .ltd2-card-title {
                font-size: 18px;
                font-weight: 900;
                color: #1f2937;
                margin-bottom: 12px;
            }

            .ltd2-list {
                margin: 0;
                padding-left: 22px;
            }

            .ltd2-list li {
                margin-bottom: 9px;
                color: #475569;
                font-size: 15px;
                line-height: 1.7;
            }

            .ltd2-price-card {
                min-height: 100%;
                background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
                border: 1px solid #dbe7f3;
                border-radius: 18px;
                display: flex;
                flex-direction: column;
                justify-content: center;
                align-items: center;
                text-align: center;
                padding: 26px 22px;
                box-shadow: 0 12px 28px rgba(15, 23, 42, 0.06);
            }

            .ltd2-price-label {
                font-size: 20px;
                font-weight: 800;
                letter-spacing: 0.08em;
                text-transform: uppercase;
                color: #64748b;
                margin-bottom: 12px;
            }

            .ltd2-price-value {
                font-size: 42px;
                line-height: 1.15;
                font-weight: 900;
                color: #1565c0;
                word-break: break-word;
            }

            .ltd2-price-sub {
                margin-top: 12px;
                font-size: 14px;
                line-height: 1.75;
                color: #475569;
                max-width: 250px;
            }

            .ltd2-price-sub {
                margin-top: 10px;
                font-size: 13px;
                line-height: 1.7;
                color: #475569;
                max-width: 230px;
            }

            .ltd2-text p {
                margin: 0 0 12px 0;
                color: #374151;
                font-size: 16px;
                line-height: 1.85;
            }

            .ltd2-info-grid {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 10px;
                margin-bottom: 14px;
            }

            .ltd2-info-item {
                background: #f8fafc;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 12px 14px;
            }

            .ltd2-info-item span {
                display: block;
                font-size: 12px;
                font-weight: 800;
                text-transform: uppercase;
                letter-spacing: 0.06em;
                color: #64748b;
                margin-bottom: 6px;
            }

            .ltd2-info-item strong {
                display: block;
                font-size: 14px;
                line-height: 1.7;
                color: #0f172a;
                font-weight: 800;
                word-break: break-word;
            }

            .ltd2-map {
                width: 100%;
                height: 300px;
                border: none;
                border-radius: 12px;
                overflow: hidden;
            }

            @media (max-width: 980px) {
                .ltd2-gallery,
                .ltd2-top-grid,
                .ltd2-bottom-grid {
                    grid-template-columns: 1fr;
                }

                .ltd2-side-grid {
                    grid-template-columns: 1fr 1fr;
                }
            }

            @media (max-width: 680px) {
                .ltd2-page {
                    padding: 14px 10px 28px 10px;
                }

                .ltd2-title {
                    font-size: 22px;
                }

                .ltd2-meta {
                    gap: 10px;
                    font-size: 13px;
                }

                .ltd2-side-grid,
                .ltd2-info-grid {
                    grid-template-columns: 1fr;
                }

                .ltd2-main-photo,
                .ltd2-side-photo {
                    min-height: 220px;
                }

                .ltd2-price-value {
                    font-size: 28px;
                }
            }
        </style>
    </head>
    <body>
        <div class="ltd2-page">
            <div class="ltd2-wrap">
                <h1 class="ltd2-title">$route_title</h1>

                <div class="ltd2-meta">
                    <div><strong>Loại hình:</strong> $category</div>
                    <div><strong>Khởi hành:</strong> $from_place</div>
                    <div><strong>Mã tour:</strong> $route_code</div>
                </div>

                <div class="ltd2-gallery">
                    <div class="ltd2-main-photo" style="background-image:url('$image_1');">
                        <div class="ltd2-main-badge">Lịch trình nổi bật</div>
                    </div>

                    <div class="ltd2-side-grid">
                        <div class="ltd2-side-photo" style="background-image:url('$image_2');"></div>
                        <div class="ltd2-side-photo" style="background-image:url('$image_3');"></div>
                        <div class="ltd2-side-photo" style="background-image:url('$image_4');"></div>
                        <div class="ltd2-side-photo" style="background-image:url('$image_5');"></div>
                    </div>
                </div>

                <div class="ltd2-top-grid">
                    <div class="ltd2-card">
                        <div class="ltd2-card-title">Thông tin tour</div>
                        <ul class="ltd2-list">
                            $highlight_html
                        </ul>
                    </div>

                    <div class="ltd2-card ltd2-price-card">
                        <div class="ltd2-price-label">Chi phí tham khảo</div>
                        <div class="ltd2-price-value">$total_price</div>
                        <div class="ltd2-price-sub">
                            Mức giá có thể thay đổi tùy thời điểm, phương tiện và dịch vụ lưu trú thực tế.
                        </div>
                    </div>
                </div>

                <div class="ltd2-bottom-grid">
                    <div class="ltd2-card">
                        <div class="ltd2-card-title">Giới thiệu</div>
                        <div class="ltd2-text">
                            $intro_html
                        </div>
                    </div>

                    <div class="ltd2-card">
                        <div class="ltd2-card-title">Thông tin hành trình</div>

                        <div class="ltd2-info-grid">
                            <div class="ltd2-info-item">
                                <span>Khách sạn / nhà nghỉ</span>
                                <strong>$hotel_name</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Giá lưu trú</span>
                                <strong>$hotel_price</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Phương tiện</span>
                                <strong>$transport_name</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Thời gian</span>
                                <strong>$time_estimate</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Quãng đường</span>
                                <strong>$distance_km</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Vé tham quan</span>
                                <strong>$ticket_price</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Giá xe</span>
                                <strong>$transport_price</strong>
                            </div>

                            <div class="ltd2-info-item">
                                <span>Điểm đến</span>
                                <strong>$to_place</strong>
                            </div>
                        </div>

                        <iframe
                            class="ltd2-map"
                            src="$map_url"
                            loading="lazy"
                            referrerpolicy="no-referrer-when-downgrade">
                        </iframe>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """)

    html_output = html_template.substitute(
        route_title=escape(route_title),
        category=escape(category),
        from_place=escape(from_place),
        to_place=escape(to_place),
        route_code=escape(route_code if route_code else "Đang cập nhật"),
        image_1=images[0],
        image_2=images[1],
        image_3=images[2],
        image_4=images[3],
        image_5=images[4],
        highlight_html=highlight_html,
        total_price=escape(total_price),
        intro_html=intro_html,
        hotel_name=escape(hotel_name),
        hotel_price=escape(hotel_price),
        transport_name=escape(transport_name),
        time_estimate=escape(time_estimate),
        distance_km=escape(distance_km),
        ticket_price=escape(ticket_price),
        transport_price=escape(transport_price),
        map_url=map_url,
    )

    page_height = 1650 + max(0, len(" ".join(paragraphs)) // 400) * 220
    components.html(html_output, height=page_height, scrolling=False)

elif page == "diemden":
    import json
    import math
    from html import escape
    from urllib.parse import quote

    try:
        with open("diemden.json", "r", encoding="utf-8") as f:
            diemden_data = json.load(f)
    except Exception as e:
        st.error(f"Lỗi đọc file diemden.json: {e}")
        st.stop()

    def safe_text(value, fallback="Đang cập nhật"):
        text = str(value).strip() if value is not None else ""
        return escape(text) if text else escape(fallback)

    def safe_url(value, fallback="https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"):
        text = str(value).strip() if value is not None else ""
        return escape(text if text else fallback, quote=True)
    
    def card_image_src(value):
        raw = str(value).strip() if value is not None else ""
        if not raw:
            return ""

        if raw.startswith(("http://", "https://", "data:")):
            return escape(raw, quote=True)

        return image_to_data_uri([
            raw,
            raw.replace(".jpg", ".png"),
            raw.replace(".jpg", ".jpeg"),
            raw.replace(".jpg", ".webp"),
            raw.replace(".png", ".jpg"),
            raw.replace(".png", ".jpeg"),
            raw.replace(".png", ".webp"),
            raw.replace(".jpeg", ".jpg"),
            raw.replace(".jpeg", ".png"),
            raw.replace(".jpeg", ".webp"),
            raw.replace(".webp", ".jpg"),
            raw.replace(".webp", ".png"),
            raw.replace(".webp", ".jpeg"),
        ])

    def normalize_text(value):
        return str(value or "").strip().lower()

    def safe_list(value):
        if isinstance(value, list):
            return [escape(str(item).strip()) for item in value if str(item).strip()]
        return []

    place_name_options = ["Tất cả địa điểm"] + sorted({item.get("name", "") for item in diemden_data if item.get("name")})
    area_options = ["Tất cả"] + sorted({item.get("area", "") for item in diemden_data if item.get("area")})
    category_options = ["Tất cả"] + sorted({item.get("category", "") for item in diemden_data if item.get("category")})
    season_options = ["Tất cả"] + sorted({item.get("season", "") for item in diemden_data if item.get("season")})

    total_places = len(diemden_data)
    total_areas = len(area_options) - 1
    total_categories = len(category_options) - 1

    hero_title = "Khám phá điểm đến nổi bật tại Lào Cai"
    hero_subtitle = "Tìm kiếm theo khu vực, loại hình và mùa đẹp để lựa chọn hành trình phù hợp với sở thích của bạn."

    diemden_hero_src = image_to_data_uri([
        "kho_anh/diem_den/chinh.jpg",
        "kho_anh/diem_den/chinh.png",
        "kho_anh/diem_den/chinh.jpeg",
        "kho_anh/diem_den/chinh.webp"
    ])

    st.markdown(f"""
    <style>
        .dd-hero-section{{
            position: relative;
            height: 500px;
            background:
                linear-gradient(180deg, rgba(7, 19, 37, 0.12) 0%, rgba(7, 19, 37, 0.22) 100%),
                url('{diemden_hero_src}');
            background-size: cover;
            background-position: center 58%;
            background-repeat: no-repeat;
            animation: ddHeroZoom 1.4s ease-out forwards;
            transform: scale(1.06);
            opacity: 0;
            overflow: hidden;
        }}

        @media (max-width: 768px) {{
            .dd-hero-section {{
                height: 210px !important;
                background-position: center center !important;
                border-radius: 0 !important;
            }}
        }}

        .dd-hero-section::after{{
            content: "";
            position: absolute;
            inset: 0;
            background: rgba(255,255,255,0.04);
            animation: ddHeroFade 1.2s ease-out forwards;
        }}

        @keyframes ddHeroZoom {{
            from {{
                transform: scale(1.06);
                opacity: 0;
            }}
            to {{
                transform: scale(1);
                opacity: 1;
            }}
        }}

        @keyframes ddHeroFade {{
            from {{
                opacity: 0.35;
            }}
            to {{
                opacity: 1;
            }}
        }}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <section class="dd-hero-section"></section>
    """, unsafe_allow_html=True)

    filtered_data = list(diemden_data)
    result_count = len(filtered_data)

    st.markdown("<div style='height:18px;'></div>", unsafe_allow_html=True)

    vi_sao_canhdep = image_to_data_uri([
        "kho_anh/diem_den/vi_sao/canhdep.webp",
        "kho_anh/diem_den/vi_sao/canhdep.png",
        "kho_anh/diem_den/vi_sao/canhdep.jpg",
        "kho_anh/diem_den/vi_sao/canhdep.jpeg"
    ])

    vi_sao_bandia = image_to_data_uri([
        "kho_anh/diem_den/vi_sao/bandia.png",
        "kho_anh/diem_den/vi_sao/bandia.webp",
        "kho_anh/diem_den/vi_sao/bandia.jpg",
        "kho_anh/diem_den/vi_sao/bandia.jpeg"
    ])

    vi_sao_vanhoa = image_to_data_uri([
        "kho_anh/diem_den/vi_sao/vanhoa.png",
        "kho_anh/diem_den/vi_sao/vanhoa.webp",
        "kho_anh/diem_den/vi_sao/vanhoa.jpg",
        "kho_anh/diem_den/vi_sao/vanhoa.jpeg"
    ])

    vi_sao_trainghiem = image_to_data_uri([
        "kho_anh/diem_den/vi_sao/trainghiem.png",
        "kho_anh/diem_den/vi_sao/trainghiem.webp",
        "kho_anh/diem_den/vi_sao/trainghiem.jpg",
        "kho_anh/diem_den/vi_sao/trainghiem.jpeg"
    ])

    components.html(f"""
    <style>
    body{{
        margin:0;
        background:transparent;
        font-family: Inter, Arial, sans-serif;
    }}

    .dd-benefit-wrap{{
        max-width: 1320px;
        margin: 26px auto 8px auto;
        padding: 0 6px;
        box-sizing: border-box;
        text-align: center;
    }}

    .dd-benefit-title{{
        font-size: 30px;
        font-weight: 900;
        color: #111827;
        text-transform: uppercase;
        margin-bottom: 8px;
        line-height: 1.3;
    }}

    .dd-benefit-line{{
        width: 86px;
        height: 4px;
        background: #f59e0b;
        border-radius: 999px;
        margin: 0 auto 16px auto;
    }}

    .dd-benefit-desc{{
        max-width: 760px;
        margin: 0 auto 28px auto;
        font-size: 16px;
        line-height: 1.8;
        color: #4b5563;
    }}

    .dd-benefit-grid{{
        display:grid;
        grid-template-columns: repeat(4, minmax(0, 1fr));
        gap:20px;
    }}

    .dd-benefit-card{{
        background:#ffffff;
        border:1px solid #e5edf6;
        border-radius:24px;
        padding:26px 20px;
        box-shadow:0 12px 28px rgba(15,23,42,0.06);
    }}

    .dd-benefit-icon{{
        width:72px;
        height:72px;
        margin:0 auto 16px auto;
        border-radius:50%;
        background:linear-gradient(135deg, #38bdf8 0%, #0ea5e9 100%);
        display:flex;
        align-items:center;
        justify-content:center;
        box-shadow:0 12px 24px rgba(14,165,233,0.22);
        overflow:hidden;
        padding:14px;
    }}

    .dd-benefit-icon img{{
        width:100%;
        height:100%;
        object-fit:contain;
        display:block;
    }}

    .dd-benefit-name{{
        font-size:18px;
        font-weight:800;
        color:#111827;
        margin-bottom:8px;
    }}

    .dd-benefit-text{{
        font-size:14px;
        line-height:1.75;
        color:#6b7280;
    }}
                    
    .dd-benefit-card{{
        background:#ffffff;
        border:1px solid #e5edf6;
        border-radius:24px;
        padding:26px 20px;
        box-shadow:0 12px 28px rgba(15,23,42,0.06);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }}

    .dd-benefit-card:hover{{
        transform: translateY(-6px);
        box-shadow:0 18px 34px rgba(15,23,42,0.12);
    }}

    .dd-benefit-name{{
        font-size:18px;
        font-weight:900;
        color:#111827;
        margin-bottom:10px;
        line-height:1.35;
        text-shadow:
            0 1px 0 rgba(255,255,255,0.85),
            0 4px 10px rgba(15,23,42,0.08);
        transition: transform 0.28s ease, color 0.28s ease, text-shadow 0.28s ease;
    }}

    .dd-benefit-text{{
        font-size:14px;
        line-height:1.8;
        color:#5b6472;
        transition: transform 0.28s ease, color 0.28s ease;
    }}

    .dd-benefit-card:hover .dd-benefit-name{{
        transform: translateY(-2px) scale(1.03);
        color:#0f4f9a;
        text-shadow: 0 6px 16px rgba(21,101,192,0.16);
    }}

    .dd-benefit-card:hover .dd-benefit-text{{
        transform: translateY(-1px);
        color:#374151;
    }}

    @media (max-width: 980px){{
        .dd-benefit-grid{{
            grid-template-columns: repeat(2, minmax(0, 1fr));
        }}
    }}

    @media (max-width: 768px){{
        .dd-benefit-wrap{{
            margin: 18px auto 6px auto;
            padding: 0 10px;
        }}

        .dd-benefit-title{{
            font-size: 20px;
            line-height: 1.35;
            margin-bottom: 8px;
        }}

        .dd-benefit-line{{
            width: 72px;
            height: 3px;
            margin: 0 auto 12px auto;
        }}

        .dd-benefit-desc{{
            font-size: 13px;
            line-height: 1.7;
            margin: 0 auto 18px auto;
            max-width: 100%;
        }}

        .dd-benefit-grid{{
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 10px;
        }}

        .dd-benefit-card{{
            border-radius: 16px;
            padding: 14px 10px;
        }}

        .dd-benefit-icon{{
            width: 50px;
            height: 50px;
            margin: 0 auto 10px auto;
            padding: 10px;
        }}

        .dd-benefit-name{{
            font-size: 14px;
            margin-bottom: 6px;
            line-height: 1.35;
        }}

        .dd-benefit-text{{
            font-size: 12px;
            line-height: 1.6;
        }}
    }}
                    
                    
    </style>

    <div class="dd-benefit-wrap">
        <div class="dd-benefit-title">VÌ SAO NÊN KHÁM PHÁ LÀO CAI CÙNG AI?</div>
        <div class="dd-benefit-line"></div>
        <div class="dd-benefit-desc">
            Khám phá Lào Cai theo cách hoàn toàn mới với trợ lý AI thông minh – giúp bạn tìm hiểu, lên lịch trình và trải nghiệm văn hóa vùng cao một cách sâu sắc, tiện lợi và cá nhân hóa.
        </div>

        <div class="dd-benefit-grid">
            <div class="dd-benefit-card">
                <div class="dd-benefit-icon"><img src="{vi_sao_canhdep}" alt="Cảnh đẹp"></div>
                <div class="dd-benefit-name">Cảnh đẹp nổi bật</div>
                <div class="dd-benefit-text">Khám phá Lào Cai qua góc nhìn trực quan với gợi ý thông minh từ AI – từ Sa Pa, Fansipan đến những điểm đến ít người biết.</div>
            </div>

            <div class="dd-benefit-card">
                <div class="dd-benefit-icon"><img src="{vi_sao_bandia}" alt="Bản địa"></div>
                <div class="dd-benefit-name">Bản sắc bản địa</div>
                <div class="dd-benefit-text">Tìm hiểu văn hóa, phong tục và đời sống của đồng bào dân tộc thông qua dữ liệu được chọn lọc và giải thích bởi AI.</div>
            </div>

            <div class="dd-benefit-card">
                <div class="dd-benefit-icon"><img src="{vi_sao_vanhoa}" alt="Văn hóa"></div>
                <div class="dd-benefit-name">Di tích - văn hóa</div>
                <div class="dd-benefit-text">Hệ thống AI cung cấp thông tin chính xác, dễ hiểu về các di tích, giúp bạn hiểu rõ giá trị lịch sử và văn hóa địa phương.</div>
            </div>

            <div class="dd-benefit-card">
                <div class="dd-benefit-icon"><img src="{vi_sao_trainghiem}" alt="Trải nghiệm"></div>
                <div class="dd-benefit-name">Trải nghiệm đa dạng</div>
                <div class="dd-benefit-text">Gợi ý lịch trình cá nhân hóa, hỗ trợ hỏi đáp nhanh chóng – giúp bạn khám phá Lào Cai thuận tiện và hiệu quả hơn bao giờ hết.</div>
            </div>
        </div>

        <script>
        function updateFrameHeight() {{
            const height = Math.max(
                document.body.scrollHeight,
                document.documentElement.scrollHeight
            );

            window.parent.postMessage({{
                type: "streamlit:setFrameHeight",
                height: height + 20
            }}, "*");
        }}

        window.addEventListener("load", () => {{
            setTimeout(updateFrameHeight, 120);
            setTimeout(updateFrameHeight, 400);
        }});

        window.addEventListener("resize", () => {{
            setTimeout(updateFrameHeight, 120);
        }});

        const resizeObserver = new ResizeObserver(() => {{
            setTimeout(updateFrameHeight, 80);
        }});

        resizeObserver.observe(document.body);
        </script>
        
    </div>
    """, height=500, scrolling=False)

    st.markdown(f"""
    <div style="
        max-width:1320px;
        margin: 8px auto 10px auto;
        padding: 0 12px;
        box-sizing:border-box;
    ">
        <div style="
            display:flex;
            align-items:flex-end;
            justify-content:space-between;
            gap:16px;
            flex-wrap:wrap;
        ">
            <div>
                <div style="
                    font-size:24px;
                    font-weight:900;
                    color:#111827;
                    line-height:1.2;
                    margin-bottom:6px;
                ">
                    Danh sách điểm đến
                </div>
                <div style="
                    font-size:14px;
                    color:#64748b;
                    line-height:1.7;
                ">
                    Tìm thấy <strong>{result_count}</strong> dành cho bạn.
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    cards_css = """
    <style>
    body{
        margin:0;
        background:transparent;
        font-family: Inter, Arial, sans-serif;
    }

    .dd-results-wrap{
        max-width: 1320px;
        margin: 12px auto 60px auto;
        padding: 0 6px;
        box-sizing: border-box;
    }

    .dd-grid{
        display:grid;
        grid-template-columns: 1fr;
        gap:24px;
    }

    .dd-card{
        display:grid;
        grid-template-columns: 320px 1fr;
        background:#ffffff;
        border:1px solid #e5edf6;
        border-radius:26px;
        overflow:hidden;
        box-shadow:0 18px 38px rgba(15,23,42,0.08);
        min-height:260px;
    }

    .dd-card-image{
        height:100%;
        min-height:260px;
        background-size:cover;
        background-position:center;
        position:relative;
        overflow:hidden;
    }

    .dd-card-image::after{
        content:"";
        position:absolute;
        inset:0;
        background:
            linear-gradient(
                to right,
                rgba(255,255,255,0) 0%,
                rgba(255,255,255,0.10) 45%,
                rgba(255,255,255,0.42) 68%,
                rgba(255,255,255,0.78) 84%,
                #ffffff 100%
            );
    }

    .dd-card-body{
        padding:24px 26px;
        display:flex;
        flex-direction:column;
        min-width:0;
        gap:18px;
    }

    .dd-card-header{
        display:flex;
        align-items:flex-start;
        justify-content:space-between;
        gap:16px;
    }

    .dd-card-header-main{
        min-width:0;
        flex:1;
    }

    .dd-card-title{
        font-size:28px;
        font-weight:900;
        color:#0f172a;
        line-height:1.2;
        margin:0;
    }

    .dd-card-top{
        display:flex;
        gap:8px;
        flex-wrap:wrap;
        justify-content:flex-end;
        flex-shrink:0;
        max-width:40%;
    }

    .dd-badge{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        padding:7px 12px;
        border-radius:999px;
        background:#ecf5ff;
        color:#1565c0;
        font-size:12px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.06em;
        border:1px solid #cfe5ff;
        white-space:nowrap;
    }

    .dd-badge-alt{
        background:#fff7ed;
        color:#c2410c;
        border:1px solid #fed7aa;
    }

    .dd-card-desc{
        font-size:16px;
        line-height:1.8;
        color:#475569;
        margin-top:-4px;
    }

    .dd-divider{
        height:1px;
        background:#e8eef5;
        margin:-2px 0 0 0;
    }

    .dd-info-grid{
        display:grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap:12px;
    }

    .dd-info-card{
        background:#f8fafc;
        border:1px solid #e6edf5;
        border-radius:18px;
        padding:14px 16px;
    }

    .dd-info-label{
        font-size:12px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.08em;
        color:#64748b;
        margin-bottom:6px;
    }

    .dd-info-value{
        font-size:15px;
        line-height:1.6;
        font-weight:700;
        color:#0f172a;
        word-break:break-word;
    }

    .dd-highlight-wrap{
        background:linear-gradient(180deg, #ffffff 0%, #fbfdff 100%);
        border:1px solid #edf2f7;
        border-radius:18px;
        padding:14px 16px;
    }

    .dd-highlight-title{
        font-size:12px;
        font-weight:800;
        text-transform:uppercase;
        letter-spacing:0.08em;
        color:#64748b;
        margin-bottom:10px;
    }

    .dd-chip-row{
        display:flex;
        gap:8px;
        flex-wrap:wrap;
    }

    .dd-chip{
        background:#f1f5f9;
        color:#334155;
        border-radius:999px;
        padding:8px 12px;
        font-size:13px;
        line-height:1.2;
        font-weight:700;
        border:1px solid #e2e8f0;
    }

    .dd-card-footer{
        display:flex;
        justify-content:flex-end;
        margin-top:auto;
    }

    .dd-detail-link{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        min-width:148px;
        padding:12px 18px;
        border-radius:14px;
        background:linear-gradient(135deg, #1565c0 0%, #2b7fd3 100%);
        color:#ffffff !important;
        text-decoration:none !important;
        font-size:14px;
        font-weight:800;
        letter-spacing:0.02em;
        box-shadow:0 12px 24px rgba(21,101,192,0.18);
        transition:transform 0.18s ease, box-shadow 0.18s ease;
    }

    .dd-detail-link:hover{
        transform:translateY(-1px);
        box-shadow:0 16px 28px rgba(21,101,192,0.24);
    }

    .dd-empty{
        background:linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
        border:1px dashed #cbd5e1;
        border-radius:26px;
        padding:42px 26px;
        text-align:center;
        color:#334155;
        box-shadow:0 14px 36px rgba(15,23,42,0.05);
    }

    .dd-empty-title{
        font-size:26px;
        font-weight:900;
        color:#0f172a;
        margin-bottom:10px;
    }

    .dd-empty-text{
        max-width:620px;
        margin:0 auto;
        font-size:16px;
        line-height:1.8;
        color:#475569;
    }

    @media (max-width: 980px){
        .dd-card{
            grid-template-columns: 260px 1fr;
        }

        .dd-card-header{
            flex-direction:column-reverse;
            align-items:flex-start;
        }

        .dd-card-top{
            justify-content:flex-start;
            max-width:100%;
        }

        .dd-card-title{
            font-size:24px;
        }

        .dd-card-desc{
            font-size:15px;
        }
    }

    @media (max-width: 760px){
        .dd-results-wrap{
            padding: 0;
        }

        .dd-grid{
            gap:18px;
        }

        .dd-card{
            grid-template-columns: 1fr;
            border-radius:22px;
        }

        .dd-card-image{
            min-height:220px;
        }

        .dd-card-body{
            padding:20px 18px;
            gap:16px;
        }

        .dd-card-title{
            font-size:22px;
        }

        .dd-info-grid{
            grid-template-columns: 1fr;
        }

        .dd-card-footer{
            justify-content:stretch;
        }

        .dd-detail-link{
            width:100%;
        }

        .dd-empty-title{
            font-size:22px;
        }

        .dd-empty-text{
            font-size:15px;
        }
    }
    </style>
    """

    cards_html = '<div class="dd-results-wrap">'

    if filtered_data:
        cards_html += '<div class="dd-grid">'
        for item in filtered_data:
            highlights = safe_list(item.get("highlights"))[:3]
            highlights_html = "".join([f'<div class="dd-chip">{h}</div>' for h in highlights])

            slug_value = safe_text(item.get("slug"), "dang-cap-nhat")
            slug_raw = str(item.get("slug", "")).strip()
            detail_href = build_app_url("diemden_detail", slug=quote(slug_raw))   
            image_src = card_image_src(item.get("image"))      

            cards_html += f"""
            <div class="dd-card">
                <div class="dd-card-image" style="background-image:url('{image_src}');"></div>

                <div class="dd-card-body">
                    <div class="dd-card-header">
                        <div class="dd-card-header-main">
                            <div class="dd-card-title">{safe_text(item.get("name"))}</div>
                        </div>

                        <div class="dd-card-top">
                            <div class="dd-badge">{safe_text(item.get("area"))}</div>
                            <div class="dd-badge dd-badge-alt">{safe_text(item.get("category"))}</div>
                        </div>
                    </div>

                    <div class="dd-card-desc">
                        {safe_text(item.get("short_desc"), "Điểm đến đang được cập nhật thông tin.")}
                    </div>

                    <div class="dd-divider"></div>

                    <div class="dd-info-grid">
                        <div class="dd-info-card">
                            <div class="dd-info-label">Mùa đẹp</div>
                            <div class="dd-info-value">{safe_text(item.get("season"), "Quanh năm")}</div>
                        </div>

                        <div class="dd-info-card">
                            <div class="dd-info-label">Mã địa điểm</div>
                            <div class="dd-info-value">{slug_value}</div>
                        </div>
                    </div>

                    <div class="dd-highlight-wrap">
                        <div class="dd-highlight-title">Điểm nổi bật</div>
                        <div class="dd-chip-row">
                            {highlights_html}
                        </div>
                    </div>

                    <div class="dd-card-footer">
                        <a class="dd-detail-link"
                        href="{detail_href}"
                        target="_blank"
                        rel="noopener noreferrer"
                        onclick="window.open('{detail_href}', '_blank', 'noopener,noreferrer'); return false;">
                            Xem chi tiết
                        </a>
                    </div>
                </div>
            </div>
            """

        cards_html += "</div>"
        rows = result_count if result_count else 1
        cards_height = max(3000, result_count * 440)
    else:
        cards_html += """
        <div class="dd-empty">
            <div class="dd-empty-title">Chưa tìm thấy điểm đến phù hợp</div>
            <div class="dd-empty-text">
                Bạn hãy thử đổi từ khóa tìm kiếm hoặc chuyển bộ lọc về <strong>Tất cả</strong>
                để xem toàn bộ điểm đến đang có trong hệ thống.
            </div>
        </div>
        """
        cards_height = 320

    cards_html += "</div>"
    components.html(cards_css + cards_html, height=cards_height, scrolling=False)

elif page == "diemden_detail":
    import os
    import json
    import base64
    import fitz
    from html import escape
    from string import Template
    from pathlib import Path
    import streamlit.components.v1 as components

    try:
        with open("diemden.json", "r", encoding="utf-8") as f:
            diemden_data = json.load(f)
    except Exception as e:
        st.error(f"Lỗi đọc file diemden.json: {e}")
        st.stop()

    def read_place_file_any(path_str):
        path = Path(path_str)
        if not path.exists():
            return ""

        suffix = path.suffix.lower()

        try:
            if suffix == ".txt":
                return path.read_text(encoding="utf-8").strip()

            elif suffix == ".docx":
                from docx import Document
                doc = Document(str(path))
                parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                return "\n".join(parts).strip()

            elif suffix == ".pdf":
                pdf = fitz.open(str(path))
                texts = []
                for page in pdf:
                    txt = page.get_text("text").strip()
                    if txt:
                        texts.append(txt)
                pdf.close()
                return "\n".join(texts).strip()
        except Exception:
            return ""

        return ""

    def pick_place_intro_text_file(place):
        folder = Path("text_diemden_gioithieu")
        slug_value = str(place.get("slug", "")).strip()
        if not slug_value:
            return ""

        file_name = slug_value.replace("-", "_").lower()

        candidates = [
            folder / f"{file_name}.txt",
            folder / f"{file_name}.docx",
            folder / f"{file_name}.pdf",
        ]

        for fp in candidates:
            if fp.exists():
                return read_place_file_any(fp)

        return ""

    def pick_place_culture_text_file(place):
        folder = Path("text_diemden_vanhoa")
        slug_value = str(place.get("slug", "")).strip()
        if not slug_value:
            return ""

        file_name = slug_value.replace("-", "_").lower()

        candidates = [
            folder / f"{file_name}.txt",
            folder / f"{file_name}.docx",
            folder / f"{file_name}.pdf",
        ]

        for fp in candidates:
            if fp.exists():
                return read_place_file_any(fp)

        return ""

    def pick_place_heritage_text_file(place):
        folder = Path("text_diemden_quanthe")
        slug_value = str(place.get("slug", "")).strip()
        if not slug_value:
            return ""

        file_name = slug_value.replace("-", "_").lower()

        candidates = [
            folder / f"{file_name}.txt",
            folder / f"{file_name}.docx",
            folder / f"{file_name}.pdf",
        ]

        for fp in candidates:
            if fp.exists():
                return read_place_file_any(fp)

        return ""

    def clean_text(value):
        if value is None:
            return ""
        return " ".join(str(value).replace("\r", "\n").split())

    def normalize_text(value):
        return clean_text(value).lower()

    def safe_text(value, fallback="Đang cập nhật"):
        text = clean_text(value)
        return escape(text if text else fallback)

    def safe_url(value, fallback="https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"):
        text = str(value).strip() if value is not None else ""
        return escape(text if text else fallback, quote=True)
    
    DEFAULT_DETAIL_IMAGE = "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"

    def file_to_data_uri(path):
        ext = os.path.splitext(path)[1].lower()
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
            ".gif": "image/gif",
            ".mp3": "audio/mpeg",
            ".wav": "audio/wav",
            ".ogg": "audio/ogg",
            ".m4a": "audio/mp4",
        }
        mime_type = mime_map.get(ext, "application/octet-stream")

        try:
            with open(path, "rb") as f:
                encoded = base64.b64encode(f.read()).decode("utf-8")
            return f"data:{mime_type};base64,{encoded}"
        except Exception:
            return ""

    def asset_to_src(value, fallback=DEFAULT_DETAIL_IMAGE):
        raw_value = str(value).strip() if value is not None else ""
        raw_fallback = str(fallback).strip() if fallback is not None else ""

        candidates = []
        if raw_value:
            candidates.append(raw_value)
        if raw_fallback:
            candidates.append(raw_fallback)

        for candidate in candidates:
            if candidate.startswith(("http://", "https://", "data:")):
                return escape(candidate, quote=True)

            if os.path.exists(candidate):
                data_uri = file_to_data_uri(candidate)
                if data_uri:
                    return data_uri

        return escape(DEFAULT_DETAIL_IMAGE, quote=True)

    def folder_image(base_folder, filename_no_ext):
        return image_to_data_uri([
            f"{base_folder}/{filename_no_ext}.png",
            f"{base_folder}/{filename_no_ext}.jpg",
            f"{base_folder}/{filename_no_ext}.jpeg",
            f"{base_folder}/{filename_no_ext}.webp"
        ])

    def slug_to_folder(slug: str) -> str:
        return str(slug or "").strip().replace("-", "_")

    def get_audio_src(place):
        def normalize_audio_name(value):
            text = str(value or "").strip().lower()
            text = text.replace("đ", "d").replace("Đ", "D")
            text = unicodedata.normalize("NFD", text)
            text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
            text = text.replace("-", "_")
            return text

        candidates = []

        # 1) Ưu tiên nếu JSON có khai báo sẵn đường dẫn
        for key in ["audio_file", "audio", "audio_path", "thuyet_minh", "thuyetminh"]:
            val = str(place.get(key, "")).strip()
            if val:
                candidates.append(val)

        # 2) Tự dò theo slug trong thư mục am_thanh
        slug_value = str(place.get("slug", "")).strip()
        if slug_value:
            slug_norm = normalize_audio_name(slug_value)
            for ext in [".wav", ".mp3", ".m4a", ".ogg"]:
                candidates.append(os.path.join("am_thanh", f"{slug_norm}{ext}"))

        # 3) Tự dò theo tên địa điểm nếu cần
        name_value = str(place.get("name", "")).strip()
        if name_value:
            name_norm = normalize_audio_name(name_value)
            for ext in [".wav", ".mp3", ".m4a", ".ogg"]:
                candidates.append(os.path.join("am_thanh", f"{name_norm}{ext}"))

        for candidate in candidates:
            if candidate.startswith(("http://", "https://", "data:")):
                return escape(candidate, quote=True)

            if os.path.exists(candidate):
                data_uri = file_to_data_uri(candidate)
                if data_uri:
                    return data_uri

        return ""

    def split_paragraphs(*values):
        results = []
        seen = set()

        for value in values:
            if value is None:
                continue
            text = str(value).replace("\r\n", "\n").replace("\r", "\n")
            parts = [p.strip() for p in text.split("\n") if p.strip()]
            if not parts and str(value).strip():
                parts = [str(value).strip()]

            for part in parts:
                normalized = " ".join(part.split()).strip()
                if normalized and normalized.lower() not in seen:
                    seen.add(normalized.lower())
                    results.append(normalized)
        return results

    def short_line(text, max_len=150):
        text = clean_text(text)
        if not text:
            return ""
        if len(text) <= max_len:
            return text
        trimmed = text[:max_len].rsplit(" ", 1)[0].strip()
        return (trimmed if trimmed else text[:max_len]).strip() + "..."

    def get_highlights(value):
        if isinstance(value, list):
            return [clean_text(x) for x in value if clean_text(x)]
        return []

    def get_images(place):
        images = []

        for key in ["hero_image", "banner_image", "image", "image_2", "image_3", "image_4", "image_5"]:
            val = place.get(key)
            if isinstance(val, str) and val.strip():
                images.append(val.strip())

        for key in ["gallery", "images", "detail_images"]:
            val = place.get(key)
            if isinstance(val, list):
                for item in val:
                    if isinstance(item, str) and item.strip():
                        images.append(item.strip())

        unique_images = []
        seen = set()
        for img in images:
            if img not in seen:
                seen.add(img)
                unique_images.append(img)

        if not unique_images:
            unique_images = [DEFAULT_DETAIL_IMAGE]

        while len(unique_images) < 8:
            unique_images.append(unique_images[-1])

        return [asset_to_src(x, DEFAULT_DETAIL_IMAGE) for x in unique_images[:8]]

    slug = str(st.query_params.get("slug", "")).strip()

    selected_place = None
    for item in diemden_data:
        if normalize_text(item.get("slug")) == normalize_text(slug):
            selected_place = item
            break

    if not selected_place:
        st.error("Không tìm thấy địa điểm.")
        st.stop()

    detail_folder_name = slug_to_folder(selected_place.get("slug", ""))
    detail_root = f"kho_anh/diem_den/tri_tiet/{detail_folder_name}"
    thuyet_minh_folder = f"{detail_root}/thuyet_minh"

    gioi_thieu_folder = f"{detail_root}/gioi_thieu"
    intro_image_big_custom = folder_image(gioi_thieu_folder, "1")
    intro_image_small_custom = folder_image(gioi_thieu_folder, "2")

    
    quan_the_di_tich_folder = f"{detail_root}/quan_the_di_tich"
    heritage_image_1 = folder_image(quan_the_di_tich_folder, "1")
    heritage_image_2 = folder_image(quan_the_di_tich_folder, "2")
    heritage_image_3 = folder_image(quan_the_di_tich_folder, "3")
    heritage_image_4 = folder_image(quan_the_di_tich_folder, "4")

    kham_pha_folder = f"{detail_root}/kham_pha"
    explore_image_1 = folder_image(kham_pha_folder, "1")
    explore_image_2 = folder_image(kham_pha_folder, "2")
    explore_image_3 = folder_image(kham_pha_folder, "3")
    explore_image_4 = folder_image(kham_pha_folder, "4")

    van_hoa_folder = f"{detail_root}/van_hoa"
    culture_custom_1 = folder_image(van_hoa_folder, "1")
    culture_custom_2 = folder_image(van_hoa_folder, "2")
    culture_custom_3 = folder_image(van_hoa_folder, "3")

    tin_tuc_hoat_dong_folder = f"{detail_root}/tin_tuc_hoat_dong"
    activity_image_1 = folder_image(tin_tuc_hoat_dong_folder, "1")
    activity_image_2 = folder_image(tin_tuc_hoat_dong_folder, "2")
    activity_image_3 = folder_image(tin_tuc_hoat_dong_folder, "3")
    activity_image_4 = folder_image(tin_tuc_hoat_dong_folder, "4")
    activity_image_5 = folder_image(tin_tuc_hoat_dong_folder, "5")
    activity_image_6 = folder_image(tin_tuc_hoat_dong_folder, "6")

    tin_tuc_le_hoi_folder = f"{detail_root}/tin_tuc_le_hoi"
    news_image_1 = folder_image(tin_tuc_le_hoi_folder, "1")
    news_image_2 = folder_image(tin_tuc_le_hoi_folder, "2")
    news_image_3 = folder_image(tin_tuc_le_hoi_folder, "3")
    news_image_4 = folder_image(tin_tuc_le_hoi_folder, "4")
    news_image_5 = folder_image(tin_tuc_le_hoi_folder, "5")

    gallery_1 = folder_image(thuyet_minh_folder, "1")
    gallery_2 = folder_image(thuyet_minh_folder, "2")
    gallery_3 = folder_image(thuyet_minh_folder, "3")
    gallery_4 = folder_image(thuyet_minh_folder, "4")
    gallery_5 = folder_image(thuyet_minh_folder, "5")
    gallery_6 = folder_image(thuyet_minh_folder, "6")
    gallery_7 = folder_image(thuyet_minh_folder, "7")
    gallery_8 = folder_image(thuyet_minh_folder, "8")

    name_raw = clean_text(selected_place.get("name")) or "Đang cập nhật"
    area_raw = clean_text(selected_place.get("area"))
    category_raw = clean_text(selected_place.get("category"))

    season_raw = clean_text(selected_place.get("season"))
    best_time_raw = clean_text(selected_place.get("best_time"))
    short_desc_raw = clean_text(selected_place.get("short_desc"))
    full_desc_raw = str(selected_place.get("full_desc") or "").strip()
    slug_raw = clean_text(selected_place.get("slug"))

    name = escape(name_raw)
    area = safe_text(area_raw)
    category = safe_text(category_raw)
    season = safe_text(season_raw)
    best_time = safe_text(best_time_raw)
    slug_value = safe_text(slug_raw)

    intro_file_text = pick_place_intro_text_file(selected_place)

    if intro_file_text:
        paragraphs = split_paragraphs(intro_file_text)
    else:
        paragraphs = []

    highlights = get_highlights(selected_place.get("highlights"))

    intro_1 = paragraphs[0] if len(paragraphs) > 0 else ""
    intro_2 = paragraphs[1] if len(paragraphs) > 1 else ""

    culture_file_text = pick_place_culture_text_file(selected_place)

    if culture_file_text:
        culture_paragraphs = split_paragraphs(culture_file_text)
    else:
        culture_paragraphs = []

    culture_1 = culture_paragraphs[0] if len(culture_paragraphs) > 0 else ""
    culture_2 = culture_paragraphs[1] if len(culture_paragraphs) > 1 else ""

    heritage_file_text = pick_place_heritage_text_file(selected_place)

    if heritage_file_text:
        heritage_paragraphs = split_paragraphs(heritage_file_text)
    else:
        heritage_paragraphs = []

    standout_text = heritage_paragraphs[0] if len(heritage_paragraphs) > 0 else ""
    visit_text = heritage_paragraphs[1] if len(heritage_paragraphs) > 1 else ""

    hero_kicker = " • ".join([x for x in [area_raw, category_raw] if x]) or "Khám phá điểm đến"
    hero_subtitle = short_line(short_desc_raw or intro_1, 160)

    images = get_images(selected_place)
    team_image_src = asset_to_src("assets/anime_team.png", images[0])
    audio_src = get_audio_src(selected_place)

    narration_image_1 = gallery_1
    narration_image_2 = gallery_2
    narration_image_3 = gallery_3
    narration_image_4 = gallery_4
    narration_image_5 = gallery_5
    narration_image_6 = gallery_6
    narration_image_7 = gallery_7
    narration_image_8 = gallery_8

    if audio_src:
        narration_audio_html = f"""
        <div class="ddt-audio-wrap">
            <audio controls preload="metadata" src="{audio_src}"></audio>
            <div class="ddt-audio-note">Phần thuyết minh dành cho: {escape(name_raw)}</div>
        </div>
        """
    else:
        narration_audio_html = """
        <div class="ddt-audio-empty">
            Chưa tìm thấy file thuyết minh cho địa điểm này.
        </div>
        """
    hero_image = images[0]
    intro_image_big = intro_image_big_custom or images[0]
    intro_image_small = intro_image_small_custom or images[1]

    card_image_1 = explore_image_1 or images[2]
    card_image_2 = explore_image_2 or images[3]
    card_image_3 = explore_image_3 or images[4]

    culture_image_1 = culture_custom_1 or images[1]
    culture_image_2 = culture_custom_2 or images[2]
    culture_image_3 = culture_custom_3 or images[3]
    
    standout_image_1 = heritage_image_1 or images[0]
    standout_image_2 = heritage_image_2 or images[1]
    standout_image_3 = heritage_image_3 or images[2]
    standout_image_4 = heritage_image_4 or images[3]

    if highlights:
        highlight_html = "".join(
            f'<span class="ddt-chip">{escape(item)}</span>' for item in highlights[:6]
        )
    else:
        highlight_html = '<span class="ddt-chip">Đang cập nhật</span>'

    news_titles = highlights[:5] if highlights else [
        f"Không gian nổi bật tại {name_raw}",
        f"Du khách khám phá {name_raw}",
        f"Hoạt động trải nghiệm quanh {name_raw}",
        f"Góc nhìn văn hóa tại {name_raw}",
        f"Hình ảnh ấn tượng về {name_raw}",
    ]

    while len(news_titles) < 5:
        news_titles.append(news_titles[-1])

    news_images = [
        news_image_1 or images[0],
        news_image_2 or images[1],
        news_image_3 or images[2],
        news_image_4 or images[3],
        news_image_5 or images[4],
    ]

    news_cards_html = "".join(
        f'''
        <div class="ddt-news-card">
            <div class="ddt-news-image" style="background-image:url('{news_images[i]}');"></div>
        </div>
        '''
        for i in range(5)
    )

    activity_title_seed = highlights[:6] if highlights else [
        f"Hoạt động nổi bật tại {name_raw}",
        f"Thông báo về {name_raw}",
        f"Cập nhật trải nghiệm tại {name_raw}",
        f"Sự kiện quanh {name_raw}",
        f"Góc nhìn mới về {name_raw}",
        f"Tin tức tham quan {name_raw}",
    ]

    while len(activity_title_seed) < 6:
        activity_title_seed.append(activity_title_seed[-1])

    activity_desc_seed = [
        short_line(culture_1, 120),
        short_line(culture_2, 120),
        short_line(intro_1, 120),
        short_line(intro_2, 120),
        short_line(standout_text, 120),
        short_line(visit_text, 120),
    ]

    activity_images = [
        activity_image_1 or images[0],
        activity_image_2 or images[1],
        activity_image_3 or images[2],
        activity_image_4 or images[3],
        activity_image_5 or images[4],
        activity_image_6 or images[5],
    ]

    activity_date_value = clean_text(selected_place.get("updated_at")) or "Đang cập nhật"

    activity_cards_html = "".join(
        f'''
        <div class="ddt-activity-card">
            <div class="ddt-activity-image" style="background-image:url('{activity_images[i]}');"></div>
        </div>
        '''
        for i in range(6)
    )
    
    page_height = 5200 + max(0, len(full_desc_raw) // 800) * 120

    html_template = Template("""
    <style>
        html, body {
            margin: 0;
            padding: 0;
            background: #f3f3f3;
            font-family: "Segoe UI", Arial, sans-serif;
            color: #111827;
        }

        * {
            box-sizing: border-box;
        }

        .ddt-page {
            background: #f3f3f3;
            padding-bottom: 70px;
        }

        .ddt-wrap {
            max-width: 1320px;
            margin: 0 auto;
            padding: 0 24px;
        }
                             
        .ddt-narration-section {
            padding: 42px 0 8px 0;
        }

        .ddt-narration-box {
            background: transparent;
            border: none;
            border-radius: 0;
            padding: 28px 0;
            display: grid;
            grid-template-columns: 430px 1fr 430px;
            gap: 36px;
            align-items: center;
            box-shadow: none;
        }
                                    
        .ddt-narration-section .ddt-wrap {
            max-width: 1500px;
            padding: 0 18px;
        }

        .ddt-team-block {
            position: relative;
            min-height: 300px;
            display: flex;
            align-items: flex-end;
            justify-content: center;
        }

        .ddt-team-hello {
            position: absolute;
            top: 6px;
            left: 0;
            font-size: 46px;
            font-weight: 900;
            color: #7ed957;
            transform: rotate(-10deg);
            text-shadow: 0 4px 10px rgba(0,0,0,0.10);
            text-transform: uppercase;
        }

        .ddt-team-image {
            width: 100%;
            max-width: 420px;
            object-fit: contain;
            display: block;
            filter: drop-shadow(0 10px 18px rgba(15,23,42,0.10));
            -webkit-mask-image: linear-gradient(to bottom, rgba(0,0,0,1) 0%, rgba(0,0,0,1) 72%, rgba(0,0,0,0) 100%);
            mask-image: linear-gradient(to bottom, rgba(0,0,0,1) 0%, rgba(0,0,0,1) 72%, rgba(0,0,0,0) 100%);
        }
                             
        .ddt-narration-center {
            display: flex;
            flex-direction: column;
            align-items: flex-start;
            justify-content: center;
        }

        .ddt-narration-kicker {
            font-size: 13px;
            font-weight: 800;
            letter-spacing: 0.16em;
            text-transform: uppercase;
            color: #475569;
            margin-bottom: 8px;
        }

        .ddt-narration-title {
            font-size: 60px;
            line-height: 0.95;
            margin: 0;
            color: #2f46c7;
            font-weight: 900;
            text-transform: uppercase;
        }

        .ddt-narration-subtitle {
            font-size: 22px;
            line-height: 1.4;
            font-weight: 800;
            color: #111827;
            text-transform: uppercase;
            margin-top: 5px;
            margin-left: 107px;
        }
        .ddt-narration-place {
            font-size: 16px;
            line-height: 1.7;
            color: #64748b;
            margin-top: 8px;
        }

        .ddt-audio-wrap {
            width: 100%;
            margin-top: 20px;
            background: #ffffff;
            border: 1px solid #eadcff;
            border-radius: 18px;
            padding: 14px 16px 12px 16px;
            box-shadow: 0 10px 20px rgba(168,85,247,0.08);
        }

        .ddt-audio-wrap audio {
            width: 100%;
            height: 44px;
        }

        .ddt-audio-note {
            margin-top: 8px;
            font-size: 13px;
            line-height: 1.6;
            color: #7c3aed;
            font-weight: 700;
        }

        .ddt-audio-empty {
            width: 100%;
            margin-top: 20px;
            padding: 18px 16px;
            border-radius: 18px;
            border: 1px dashed #cbd5e1;
            background: #ffffff;
            color: #64748b;
            font-size: 15px;
            line-height: 1.7;
            font-weight: 700;
        }

        .ddt-narration-collage {
            display: grid;
            grid-template-columns: 56px 84px 124px 84px 56px;
            grid-template-rows: 52px 80px 88px 80px 52px;
            gap: 10px;
            justify-content: center;
            align-content: center;
            overflow: visible;
        }

        .ddt-narration-item {
            background-size: cover;
            background-position: center;
            border: 3px solid #ffffff;
            border-radius: 10px;
            box-shadow: 0 10px 20px rgba(15,23,42,0.10);
            transition: transform 0.28s ease, box-shadow 0.28s ease, z-index 0.28s ease;
            position: relative;
            z-index: 1;
            cursor: pointer;
        }

        .ddt-narration-item:hover {
            transform: scale(1.18);
            box-shadow: 0 22px 40px rgba(15,23,42,0.22);
            z-index: 20;
        }

        .ddt-ni-1 { grid-column: 1 / 2; grid-row: 2 / 4; }
        .ddt-ni-2 { grid-column: 2 / 3; grid-row: 1 / 3; }
        .ddt-ni-3 { grid-column: 3 / 4; grid-row: 1 / 3; }
        .ddt-ni-4 { grid-column: 4 / 5; grid-row: 1 / 3; }
        .ddt-ni-5 { grid-column: 2 / 3; grid-row: 3 / 5; }
        .ddt-ni-6 { grid-column: 3 / 5; grid-row: 3 / 4; }
        .ddt-ni-7 { grid-column: 3 / 4; grid-row: 4 / 6; }
        .ddt-ni-8 { grid-column: 5 / 6; grid-row: 2 / 4; }

        .ddt-hero {
            position: relative;
            width: 100%;
            min-height: 620px;
            border-radius: 0;
            overflow: hidden;
            background-image: url('$hero_image');
            background-size: cover;
            background-position: center;
            box-shadow: none;
        }

        .ddt-hero::before {
            content: "";
            position: absolute;
            inset: 0;
            background: linear-gradient(180deg, rgba(0,0,0,0.18) 0%, rgba(0,0,0,0.34) 65%, rgba(0,0,0,0.42) 100%);
        }

        .ddt-hero-content {
            position: relative;
            z-index: 2;
            min-height: 560px;
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            text-align: center;
            padding: 32px;
            color: #ffffff;
        }

        .ddt-hero-kicker {
            font-size: 15px;
            font-weight: 800;
            letter-spacing: 0.16em;
            text-transform: uppercase;
            margin-bottom: 16px;
        }

        .ddt-script {
            font-family: "Segoe Script", "Brush Script MT", cursive;
        }

        .ddt-hero-title {
            font-size: 72px;
            line-height: 1.05;
            font-weight: 800;
            margin: 0 0 14px 0;
            text-shadow: 0 4px 18px rgba(0,0,0,0.32);
        }

        .ddt-hero-subtitle {
            max-width: 760px;
            font-size: 19px;
            line-height: 1.7;
            font-weight: 700;
            text-transform: uppercase;
            text-shadow: 0 3px 14px rgba(0,0,0,0.28);
        }

        .ddt-section {
            position: relative;
            padding: 78px 0 10px 0;
        }

        .ddt-section::before {
            content: "";
            position: absolute;
            width: 430px;
            height: 430px;
            border: 1px solid rgba(15,23,42,0.04);
            border-radius: 50%;
            left: 50%;
            top: 50%;
            transform: translate(-50%, -50%);
            box-shadow:
                0 0 0 18px rgba(15,23,42,0.015),
                0 0 0 38px rgba(15,23,42,0.010);
            pointer-events: none;
        }

        .ddt-two-col {
            position: relative;
            z-index: 1;
            display: grid;
            grid-template-columns: 1.05fr 0.95fr;
            gap: 56px;
            align-items: center;
        }

        .ddt-two-col.reverse {
            grid-template-columns: 0.95fr 1.05fr;
        }

        .ddt-section-kicker {
            font-size: 15px;
            font-weight: 800;
            letter-spacing: 0.18em;
            text-transform: uppercase;
            color: #334155;
            margin-bottom: 8px;
        }

        .ddt-section-title {
            font-size: 58px;
            line-height: 1.04;
            margin: 0 0 18px 0;
            color: #930000;
            font-weight: 800;
        }

        .ddt-text {
            font-size: 17px;
            line-height: 1.85;
            color: #374151;
            margin-bottom: 18px;
        }

        .ddt-button {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-width: 98px;
            padding: 11px 24px;
            border-radius: 999px;
            background: #f7b52a;
            color: #111827;
            text-decoration: none;
            font-size: 15px;
            font-weight: 800;
            box-shadow: 0 8px 18px rgba(247,181,42,0.24);
        }

        .ddt-meta-grid {
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 12px;
            margin: 24px 0 22px 0;
        }

        .ddt-meta-card {
            background: rgba(255,255,255,0.78);
            border: 1px solid #e8ebf0;
            border-radius: 18px;
            padding: 14px 16px;
            box-shadow: 0 10px 22px rgba(15,23,42,0.04);
        }

        .ddt-meta-label {
            font-size: 12px;
            font-weight: 800;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #64748b;
            margin-bottom: 6px;
        }

        .ddt-meta-value {
            font-size: 16px;
            line-height: 1.6;
            font-weight: 700;
            color: #111827;
            word-break: break-word;
        }

        .ddt-chip-row {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-top: 8px;
        }

        .ddt-chip {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            padding: 8px 13px;
            border-radius: 999px;
            background: #ffffff;
            border: 1px solid #e5e7eb;
            color: #334155;
            font-size: 13px;
            font-weight: 700;
            box-shadow: 0 6px 16px rgba(15,23,42,0.04);
        }

        .ddt-collage-left {
            position: relative;
            min-height: 520px;
        }

        .ddt-collage-left .main {
            width: 74%;
            height: 500px;
            background-size: cover;
            background-position: center;
            box-shadow: 0 16px 36px rgba(15,23,42,0.10);
        }

        .ddt-collage-left .small {
            position: absolute;
            right: 0;
            top: 110px;
            width: 40%;
            height: 290px;
            background-size: cover;
            background-position: center;
            box-shadow: 0 16px 36px rgba(15,23,42,0.12);
        }

        .ddt-card-section {
            padding: 56px 0 14px 0;
        }

        .ddt-center-heading {
            text-align: center;
            margin-bottom: 34px;
        }

        .ddt-card-grid {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 30px;
        }

        .ddt-card {
            text-align: center;
        }

        .ddt-card-image {
            width: 100%;
            height: 310px;
            border-radius: 4px;
            background-size: cover;
            background-position: center;
            box-shadow: 0 14px 30px rgba(15,23,42,0.08);
        }

        .ddt-card-title {
            margin: 18px 0 10px 0;
            font-size: 18px;
            line-height: 1.4;
            font-weight: 900;
            color: #0f172a;
            text-transform: uppercase;
        }

        .ddt-card-desc {
            font-size: 15px;
            line-height: 1.7;
            color: #475569;
            max-width: 88%;
            margin: 0 auto 16px auto;
        }

        .ddt-mosaic-right {
            position: relative;
            min-height: 540px;
            display: grid;
            grid-template-columns: 1fr 1fr;
            grid-template-rows: 330px 140px;
            gap: 18px;
        }

        .ddt-mosaic-right .img-a {
            grid-column: 1 / 3;
            grid-row: 1 / 2;
        }

        .ddt-mosaic-right .img-b {
            grid-column: 1 / 2;
            grid-row: 2 / 3;
        }

        .ddt-mosaic-right .img-c {
            grid-column: 2 / 3;
            grid-row: 2 / 3;
        }

        .ddt-mosaic-right .img-a,
        .ddt-mosaic-right .img-b,
        .ddt-mosaic-right .img-c {
            background-size: cover;
            background-position: center;
            border-radius: 4px;
            box-shadow: 0 14px 30px rgba(15,23,42,0.08);
        }

        .ddt-mosaic-left {
            position: relative;
            min-height: 540px;
            display: grid;
            grid-template-columns: 1fr 1fr;
            grid-template-rows: 260px 190px 150px;
            gap: 18px;
        }

        .ddt-mosaic-left .img-a {
            grid-column: 1 / 2;
            grid-row: 1 / 3;
        }

        .ddt-mosaic-left .img-b {
            grid-column: 2 / 3;
            grid-row: 1 / 2;
        }

        .ddt-mosaic-left .img-c {
            grid-column: 1 / 2;
            grid-row: 3 / 4;
        }

        .ddt-mosaic-left .img-d {
            grid-column: 2 / 3;
            grid-row: 2 / 4;
        }

        .ddt-mosaic-left .img-a,
        .ddt-mosaic-left .img-b,
        .ddt-mosaic-left .img-c,
        .ddt-mosaic-left .img-d {
            background-size: cover;
            background-position: center;
            border-radius: 4px;
            box-shadow: 0 14px 30px rgba(15,23,42,0.08);
        }
                             
        .ddt-news-section{
            position: relative;
            margin-top: 42px;
            background: #a50000;
            padding: 48px 0 52px 0;
            overflow: hidden;
        }

        .ddt-news-section::before{
            content:"";
            position:absolute;
            width:520px;
            height:520px;
            left:50%;
            top:50%;
            transform:translate(-50%, -50%);
            border-radius:50%;
            border:1px solid rgba(255,255,255,0.08);
            box-shadow:
                0 0 0 20px rgba(255,255,255,0.035),
                0 0 0 42px rgba(255,255,255,0.02);
            pointer-events:none;
        }

        .ddt-news-head{
            position:relative;
            z-index:2;
            text-align:center;
            margin-bottom:22px;
        }

        .ddt-news-kicker{
            font-size:15px;
            font-weight:800;
            letter-spacing:0.16em;
            text-transform:uppercase;
            color:rgba(255,255,255,0.86);
            margin-bottom:8px;
        }

        .ddt-news-title{
            font-size:42px;
            line-height:1.05;
            margin:0;
            color:#ffffff;
            font-weight:800;
        }

        #tin-tuc-le-hoi .ddt-wrap{
            max-width:100%;
            padding:0;
        }

        .ddt-news-slider{
            position:relative;
            z-index:2;
            overflow:hidden;
        }

        .ddt-news-viewport{
            overflow:hidden;
            padding:0 24px;
        }

        .ddt-news-track{
            display:flex;
            gap:16px;
            transition:transform 0.5s ease;
            will-change:transform;
        }

        .ddt-news-card{
            flex:0 0 30%;
            min-width:0;
        }

        .ddt-news-image{
            width:100%;
            height:260px;
            background-size:cover;
            background-position:center;
            box-shadow:0 18px 36px rgba(0,0,0,0.14);
        }

        .ddt-news-controls{
            position:relative;
            z-index:2;
            display:flex;
            justify-content:center;
            gap:14px;
            margin-top:18px;
        }
        .ddt-news-btn{
            width:42px;
            height:42px;
            border:none;
            border-radius:50%;
            background:#ffffff;
            color:#7c2d12;
            font-size:20px;
            font-weight:800;
            cursor:pointer;
            box-shadow:0 8px 18px rgba(0,0,0,0.14);
        }

        .ddt-news-btn:disabled{
            opacity:0.45;
            cursor:not-allowed;
        }
                             
        .ddt-activity-section{
            padding:72px 0 50px 0;
            background:#f3f3f3;
        }

        .ddt-activity-head{
            text-align:center;
            margin-bottom:34px;
        }

        .ddt-activity-kicker{
            font-size:15px;
            font-weight:800;
            letter-spacing:0.16em;
            text-transform:uppercase;
            color:#334155;
            margin-bottom:10px;
        }

        .ddt-activity-title{
            font-size:56px;
            line-height:1.05;
            margin:0;
            color:#930000;
            font-weight:800;
        }

        .ddt-activity-grid{
            display:grid;
            grid-template-columns:repeat(3, minmax(0, 1fr));
            gap:32px 26px;
        }

        .ddt-activity-card{
            min-width:0;
        }

        .ddt-activity-image{
            width:100%;
            height:215px;
            background-size:cover;
            background-position:center;
            margin-bottom:14px;
        }

        .ddt-activity-title-card{
            font-size:17px;
            line-height:1.45;
            font-weight:900;
            color:#111827;
            margin-bottom:10px;
            text-transform:uppercase;
        }

        @media (max-width: 1100px) {
            .ddt-two-col,
            .ddt-two-col.reverse {
                grid-template-columns: 1fr;
                gap: 34px;
            }

            .ddt-card-grid {
                grid-template-columns: 1fr;
                gap: 24px;
            }

            .ddt-card-image {
                height: 280px;
            }

            .ddt-hero-title {
                font-size: 56px;
            }

            .ddt-section-title {
                font-size: 48px;
            }
                
            .ddt-news-title {
                font-size: 46px;
            }

            .ddt-news-card {
                flex: 0 0 44%;
            }

            .ddt-news-image {
                height: 230px;
            }
                            
            .ddt-activity-title{
                font-size:46px;
            }

            .ddt-activity-grid{
                grid-template-columns:repeat(2, minmax(0, 1fr));
            }

            .ddt-activity-image{
                height:210px;
            }
                             
            .ddt-narration-box {
                grid-template-columns: 1fr;
                gap: 22px;
            }

            .ddt-narration-center {
                align-items: center;
                text-align: center;
            }

            .ddt-team-block {
                min-height: 260px;
            }
        }

        @media (max-width: 768px) {
            .ddt-wrap {
                padding: 0 14px;
            }

            .ddt-hero {
                min-height: 420px;
                border-radius: 22px;
            }

            .ddt-hero-content {
                min-height: 420px;
                padding: 24px 18px;
            }

            .ddt-hero-title {
                font-size: 40px;
            }

            .ddt-hero-subtitle {
                font-size: 15px;
            }

            .ddt-section {
                padding: 52px 0 4px 0;
            }

            .ddt-section-title {
                font-size: 38px;
            }

            .ddt-text {
                font-size: 15px;
            }

            .ddt-meta-grid {
                grid-template-columns: 1fr;
            }

            .ddt-collage-left {
                min-height: auto;
                display: flex;
                flex-direction: column;
                gap: 16px;
            }

            .ddt-collage-left .main,
            .ddt-collage-left .small {
                position: relative;
                width: 100%;
                height: 240px;
                top: auto;
                right: auto;
            }

            .ddt-mosaic-right,
            .ddt-mosaic-left {
                min-height: auto;
                grid-template-columns: 1fr;
                grid-template-rows: none;
            }

            .ddt-mosaic-right .img-a,
            .ddt-mosaic-right .img-b,
            .ddt-mosaic-right .img-c,
            .ddt-mosaic-left .img-a,
            .ddt-mosaic-left .img-b,
            .ddt-mosaic-left .img-c,
            .ddt-mosaic-left .img-d {
                grid-column: auto;
                grid-row: auto;
                height: 220px;
            }
                             
            .ddt-news-section {
                padding: 56px 0 64px 0;
            }

            .ddt-news-title {
                font-size: 36px;
            }

            .ddt-news-card {
                flex: 0 0 86%;
            }

            .ddt-news-image {
                height: 190px;
            }

            .ddt-news-caption {
                font-size: 15px;
            }
                             
            .ddt-activity-section{
                padding:56px 0 36px 0;
            }

            .ddt-activity-title{
                font-size:36px;
            }

            .ddt-activity-grid{
                grid-template-columns:1fr;
                gap:24px;
            }

            .ddt-activity-image{
                height:220px;
            }
                             
            .ddt-narration-box {
                padding: 20px 16px;
                border-radius: 20px;
            }

            .ddt-team-hello {
                font-size: 34px;
            }

            .ddt-team-image {
                max-width: 240px;
            }

            .ddt-narration-title {
                font-size: 42px;
            }

            .ddt-narration-subtitle {
                font-size: 17px;
            }

            .ddt-narration-collage {
                grid-template-columns: repeat(3, 1fr);
                grid-template-rows: none;
            }

            .ddt-ni-1, .ddt-ni-2, .ddt-ni-3, .ddt-ni-4,
            .ddt-ni-5, .ddt-ni-6, .ddt-ni-7, .ddt-ni-8 {
                grid-column: auto;
                grid-row: auto;
                min-height: 90px;
            }
        }
    </style>

    <div class="ddt-page">
        <section class="ddt-hero">
            <div class="ddt-hero-content">
                <h1 class="ddt-hero-title ddt-script">$name</h1>
                <div class="ddt-hero-subtitle">$hero_subtitle</div>
            </div>
        </section>

        <section class="ddt-narration-section" id="thuyet-minh">
            <div class="ddt-wrap">
                <div class="ddt-narration-box">

                    <div class="ddt-team-block">
                        <img class="ddt-team-image" src="$team_image_src" alt="Đội thuyết minh" />
                    </div>

                    <div class="ddt-narration-center">
                        <div class="ddt-narration-kicker"></div>
                        <h2 class="ddt-narration-title">Thuyết minh</h2>
                        <div class="ddt-narration-subtitle">Nghe thuyết minh tại đây</div>
                        $narration_audio_html
                    </div>

                    <div class="ddt-narration-collage">
                        <div class="ddt-narration-item ddt-ni-1" style="background-image:url('$narration_image_1');"></div>
                        <div class="ddt-narration-item ddt-ni-2" style="background-image:url('$narration_image_2');"></div>
                        <div class="ddt-narration-item ddt-ni-3" style="background-image:url('$narration_image_3');"></div>
                        <div class="ddt-narration-item ddt-ni-4" style="background-image:url('$narration_image_4');"></div>
                        <div class="ddt-narration-item ddt-ni-5" style="background-image:url('$narration_image_5');"></div>
                        <div class="ddt-narration-item ddt-ni-6" style="background-image:url('$narration_image_6');"></div>
                        <div class="ddt-narration-item ddt-ni-7" style="background-image:url('$narration_image_7');"></div>
                        <div class="ddt-narration-item ddt-ni-8" style="background-image:url('$narration_image_8');"></div>
                    </div>

                </div>
            </div>
        </section>
                             
        <section class="ddt-section" id="gioi-thieu">
            <div class="ddt-wrap">
                <div class="ddt-two-col">
                    <div class="ddt-collage-left">
                        <div class="main" style="background-image:url('$intro_image_big');"></div>
                        <div class="small" style="background-image:url('$intro_image_small');"></div>
                    </div>

                    <div>
                        <div class="ddt-section-kicker">Về điểm đến</div>
                        <h2 class="ddt-section-title ddt-script">Giới thiệu</h2>

                        <div class="ddt-text">$intro_1</div>
                        <div class="ddt-text">$intro_2</div>

                        <div class="ddt-meta-grid">
                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Khu vực</div>
                                <div class="ddt-meta-value">$area</div>
                            </div>

                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Loại hình</div>
                                <div class="ddt-meta-value">$category</div>
                            </div>

                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Mùa đẹp</div>
                                <div class="ddt-meta-value">$season</div>
                            </div>

                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Mã địa điểm</div>
                                <div class="ddt-meta-value">$slug_value</div>
                            </div>
                        </div>

                        <div class="ddt-chip-row">
                            $highlight_html
                        </div>
                    </div>
                </div>
            </div>
        </section>

        <section class="ddt-card-section" id="kham-pha">
            <div class="ddt-wrap">
                <div class="ddt-center-heading">
                    <div class="ddt-section-kicker">Đầu tư</div>
                    <h2 class="ddt-section-title ddt-script">Khám phá</h2>
                </div>

                <div class="ddt-card-grid">
                    <div class="ddt-card">
                        <div class="ddt-card-image" style="background-image:url('$card_image_1');"></div>
                        <div class="ddt-card-title">Văn hóa, lịch sử</div>
                        <div class="ddt-card-desc">Tìm hiểu chiều sâu văn hóa và những nét nổi bật gắn với $name.</div>
                    </div>

                    <div class="ddt-card">
                        <div class="ddt-card-image" style="background-image:url('$card_image_2');"></div>
                        <div class="ddt-card-title">Điểm nổi bật</div>
                        <div class="ddt-card-desc">Khám phá các chi tiết giúp địa điểm này tạo dấu ấn trong hành trình tham quan.</div>
                    </div>

                    <div class="ddt-card">
                        <div class="ddt-card-image" style="background-image:url('$card_image_3');"></div>
                        <div class="ddt-card-title">Thông tin tham quan</div>
                        <div class="ddt-card-desc">Xem nhanh khu vực, thời điểm phù hợp và gợi ý trải nghiệm trước khi ghé thăm.</div>
                    </div>
                </div>
            </div>
        </section>

        <section class="ddt-section" id="van-hoa-lich-su">
            <div class="ddt-wrap">
                <div class="ddt-two-col reverse">
                    <div>
                        <div class="ddt-section-kicker">Khám phá</div>
                        <h2 class="ddt-section-title ddt-script">Văn hóa, lịch sử</h2>

                        <div class="ddt-text">$culture_1</div>
                        <div class="ddt-text">$culture_2</div>

                    </div>

                    <div class="ddt-mosaic-right">
                        <div class="img-a" style="background-image:url('$culture_image_1');"></div>
                        <div class="img-b" style="background-image:url('$culture_image_2');"></div>
                        <div class="img-c" style="background-image:url('$culture_image_3');"></div>
                    </div>
                </div>
            </div>
        </section>

        <section class="ddt-section" id="diem-noi-bat">
            <div class="ddt-wrap">
                <div class="ddt-two-col">
                    <div class="ddt-mosaic-left">
                        <div class="img-a" style="background-image:url('$standout_image_1');"></div>
                        <div class="img-b" style="background-image:url('$standout_image_2');"></div>
                        <div class="img-c" style="background-image:url('$standout_image_3');"></div>
                        <div class="img-d" style="background-image:url('$standout_image_4');"></div>
                    </div>

                    <div id="tham-quan">
                        <div class="ddt-section-kicker">Tham quan</div>
                        <h2 class="ddt-section-title ddt-script">Quần thể di tích</h2>

                        <div class="ddt-text">$standout_text</div>
                        <div class="ddt-text">$visit_text</div>

                        <div class="ddt-meta-grid">
                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Thời điểm phù hợp</div>
                                <div class="ddt-meta-value">$best_time</div>
                            </div>

                            <div class="ddt-meta-card">
                                <div class="ddt-meta-label">Mùa tham quan</div>
                                <div class="ddt-meta-value">$season</div>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </section>
                        
        <section class="ddt-news-section" id="tin-tuc-le-hoi">
            <div class="ddt-wrap">
                <div class="ddt-news-head">
                    <div class="ddt-news-kicker">Thông tin</div>
                    <h2 class="ddt-news-title ddt-script">Sắc màu văn hóa Lào Cai</h2>
                </div>

                <div class="ddt-news-slider">
                    <div class="ddt-news-viewport">
                        <div class="ddt-news-track" id="ddtNewsTrack">
                            $news_cards_html
                        </div>
                    </div>

                    <div class="ddt-news-controls">
                        <button class="ddt-news-btn" id="ddtNewsPrev">‹</button>
                        <button class="ddt-news-btn" id="ddtNewsNext">›</button>
                    </div>
                </div>
            </div>
        </section>

        <script>
        (function () {
            const track = document.getElementById("ddtNewsTrack");
            const prevBtn = document.getElementById("ddtNewsPrev");
            const nextBtn = document.getElementById("ddtNewsNext");

            if (!track || !prevBtn || !nextBtn) return;

            let currentIndex = 0;

            function getVisibleCount() {
                if (window.innerWidth <= 768) return 1;
                if (window.innerWidth <= 1100) return 2;
                return 3;
            }

            function updateSlider() {
                const cards = track.querySelectorAll(".ddt-news-card");
                if (!cards.length) return;

                const gap = 18;
                const cardWidth = cards[0].offsetWidth + gap;
                const maxIndex = Math.max(0, cards.length - getVisibleCount());

                if (currentIndex > maxIndex) currentIndex = maxIndex;

                const peekOffset = 110;
                track.style.transform = 'translateX(-' + ((currentIndex * cardWidth) + peekOffset) + 'px)';
                prevBtn.disabled = currentIndex === 0;
                nextBtn.disabled = currentIndex >= maxIndex;
            }

            prevBtn.addEventListener("click", function () {
                if (currentIndex > 0) {
                    currentIndex--;
                    updateSlider();
                }
            });

            nextBtn.addEventListener("click", function () {
                const cards = track.querySelectorAll(".ddt-news-card");
                const maxIndex = Math.max(0, cards.length - getVisibleCount());
                if (currentIndex < maxIndex) {
                    currentIndex++;
                    updateSlider();
                }
            });

            window.addEventListener("resize", updateSlider);
            setTimeout(updateSlider, 120);
        })();
        </script>

        </div>
                             
        <section class="ddt-activity-section" id="tin-tuc-hoat-dong">
            <div class="ddt-wrap">
                <div class="ddt-activity-head">
                    <div class="ddt-activity-kicker">Thông tin</div>
                    <h2 class="ddt-activity-title ddt-script">Hoạt động nổi bật</h2>
                </div>

                <div class="ddt-activity-grid">
                    $activity_cards_html
                </div>
            </div>
        </section>
    </div>
    """)

    html_output = html_template.substitute(
        hero_image=hero_image,
        hero_kicker=escape(hero_kicker),
        name=name,
        hero_subtitle=escape(hero_subtitle),
        intro_image_big=intro_image_big,
        intro_image_small=intro_image_small,
        intro_1=escape(intro_1),
        intro_2=escape(intro_2),
        area=area,
        category=category,
        season=season,
        slug_value=slug_value,
        highlight_html=highlight_html,
        card_image_1=card_image_1,
        card_image_2=card_image_2,
        card_image_3=card_image_3,
        culture_1=escape(culture_1),
        culture_2=escape(culture_2),
        culture_image_1=culture_image_1,
        culture_image_2=culture_image_2,
        culture_image_3=culture_image_3,
        standout_image_1=standout_image_1,
        standout_image_2=standout_image_2,
        standout_image_3=standout_image_3,
        standout_image_4=standout_image_4,
        standout_text=escape(standout_text),
        visit_text=escape(visit_text),
        best_time=best_time,
        news_cards_html=news_cards_html,
        activity_cards_html=activity_cards_html,

        team_image_src=team_image_src,
        narration_audio_html=narration_audio_html,
        narration_image_1=narration_image_1,
        narration_image_2=narration_image_2,
        narration_image_3=narration_image_3,
        narration_image_4=narration_image_4,
        narration_image_5=narration_image_5,
        narration_image_6=narration_image_6,
        narration_image_7=narration_image_7,
        narration_image_8=narration_image_8,
    )

    components.html(html_output, height=page_height, scrolling=False)

elif page == "chatbot":
    import base64
    from html import escape
    import streamlit.components.v1 as components

    def _image_to_data_uri(candidates):
        mime_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp"
        }

        for path in candidates:
            if path and os.path.exists(path):
                ext = os.path.splitext(path)[1].lower()
                mime = mime_map.get(ext, "image/png")
                try:
                    with open(path, "rb") as f:
                        encoded = base64.b64encode(f.read()).decode("utf-8")
                    return f"data:{mime};base64,{encoded}"
                except Exception:
                    pass
        return ""

    logo_src = _image_to_data_uri([
        "assets/anime_teamlogo.png",
        "anime_teamlogo.png"
    ])

    default_demo_messages = [
        {
            "role": "user",
            "content": "xin chào"
        },
        {
            "role": "assistant",
            "content": "Xin chào! Mình đây. Bạn cần mình hỗ trợ gì ngay bây giờ?"
        },

        {
            "role": "user",
            "content": "giới thiệu giúp về đền thượng lào cai"
        },
        {
            "role": "assistant",
            "content": "Mình tìm nhanh thông tin đáng tin cậy về Đền Thượng để giới thiệu cho bạn gọn, dễ nói, và đúng địa danh Lào Cai.\n\nĐền Thượng Lào Cai là một điểm du lịch tâm linh và di tích văn hóa nổi bật của thành phố Lào Cai. Nơi đây gắn với việc thờ Hưng Đạo Đại Vương Trần Quốc Tuấn, mang ý nghĩa tưởng nhớ công lao giữ nước và bảo vệ vùng biên cương. Lễ hội Đền Thượng cũng là một sinh hoạt văn hóa lớn đầu xuân của địa phương.\n\nBạn có thể giới thiệu ngắn gọn như sau:\n\nĐền Thượng nằm tại phường Lào Cai, thành phố Lào Cai, là một địa điểm tâm linh nổi tiếng và giàu giá trị lịch sử của vùng đất biên cương. Ngôi đền thờ Hưng Đạo Đại Vương Trần Quốc Tuấn, thể hiện truyền thống uống nước nhớ nguồn và lòng biết ơn đối với vị anh hùng dân tộc có công với đất nước. Hằng năm, nơi đây diễn ra Lễ hội Đền Thượng vào dịp tháng Giêng, thu hút đông đảo người dân và du khách đến dâng hương, tham quan và hòa mình vào các hoạt động văn hóa truyền thống đặc sắc. Đền Thượng không chỉ là nơi sinh hoạt tín ngưỡng mà còn góp phần làm nổi bật bản sắc văn hóa của Lào Cai."
        },

        {
            "role": "user",
            "content": "ngắn gọn giúp mình tạo lịch trình 3 ngày 3 đêm tại sapa"
        },
        {
            "role": "assistant",
            "content": "Mình ghép cho bạn một lịch trình ngắn, dễ demo, dựa trên các điểm nổi bật và thông tin hiện có ở Sa Pa.\n\nĐây là lịch trình ngắn gọn 3 ngày 3 đêm ở Sa Pa:\n\nĐêm 1:\nDi chuyển đến Sa Pa, nhận phòng khách sạn, đi dạo nhà thờ đá Sa Pa, ăn đồ nướng và nghỉ ngơi. Sa Pa là điểm du lịch nổi bật của Lào Cai, nổi tiếng với khí hậu mát, cảnh núi và văn hóa bản địa.\n\nNgày 1:\nTham quan bản Cát Cát buổi sáng, chiều check-in Moana Sa Pa hoặc khu trung tâm thị xã, tối dạo chợ đêm. Cát Cát là một trong những điểm tham quan phổ biến nhất ở Sa Pa.\n\nĐêm 2:\nNghỉ tại Sa Pa, thưởng thức đặc sản địa phương.\n\nNgày 2:\nĐi Fansipan bằng cáp treo vào buổi sáng, chiều nghỉ ngơi hoặc cafe ngắm núi. Fansipan là điểm đến nổi bật của Sa Pa và thường được xem là biểu tượng du lịch của khu vực.\n\nĐêm 3:\nTự do khám phá thị xã Sa Pa, mua quà, chuẩn bị về.\n\nNgày 3:\nTham quan Lao Chải – Tả Van để ngắm ruộng bậc thang, tìm hiểu đời sống người dân bản địa, sau đó lên xe trở về. Đây là tuyến trải nghiệm rất phổ biến khi đi Sa Pa.\n\nBạn cứ ngắn để cho vào web hoặc thuyết trình:\n\n“Đêm 1 đến Sa Pa nghỉ ngơi, ngày 1 đi bản Cát Cát, ngày 2 chinh phục Fansipan, ngày 3 tham quan Lao Chải – Tả Van rồi trở về.”"
        },

        {
            "role": "user",
            "content": "du lịch lào cai lên ăn món nào"
        },
        {
            "role": "assistant",
            "content": "Mình kiểm tra nhanh các món đặc trưng Lào Cai để gợi ý đúng kiểu “đi Lào Cai nên ăn gì”, ưu tiên món nổi tiếng và dễ giới thiệu.\n\nLên Lào Cai, nhất là Sa Pa, bạn nên thử mấy món này:\n\n1. Thắng cố\nMón đặc trưng vùng cao, rất nổi tiếng ở Sa Pa và thường được nhắc như một món “phải thử” khi đến đây.\n\n2. Lẩu cá hồi, cá tầm\nĐây là nhóm món nổi bật của Sa Pa, hợp thời tiết lạnh và được báo địa phương giới thiệu nhiều cho du khách.\n\n3. Thịt lợn cắp nách\nMột đặc sản quen thuộc của vùng cao Lào Cai, thường có trong danh sách món nên ăn khi du lịch Sa Pa.\n\n4. Thịt trâu gác bếp\nMón rất hợp mua ăn tại chỗ hoặc làm quà, cũng là đặc sản Tây Bắc được nhắc nhiều ở Sa Pa.\n\n5. Gà đen\nLà món đặc sản địa phương, thường xuất hiện trong các bữa ăn dành cho khách du lịch ở Sa Pa."
        }
    ]
    

    current_user_id = st.session_state["auth_user"]["id"]

    if (
        "gemini_loaded_user_id" not in st.session_state
        or st.session_state.gemini_loaded_user_id != current_user_id
    ):
        ensure_default_chat_history(current_user_id, default_demo_messages)
        st.session_state.gemini_messages = load_chat_history(current_user_id)
        st.session_state.gemini_loaded_user_id = current_user_id

    def send_chat_message(user_text: str):
        clean_user_text = str(user_text or "").strip()
        if not clean_user_text:
            return

        save_chat_message(current_user_id, "user", clean_user_text)

        st.session_state.gemini_messages.append({
            "role": "user",
            "content": clean_user_text
        })

        answer = ask_gemini(
            user_message=clean_user_text,
            payload=CHATBOT_PAYLOAD,
            current_page=page,
            messages=st.session_state.gemini_messages[:-1]
        )

        save_chat_message(current_user_id, "assistant", answer)

        st.session_state.gemini_messages.append({
            "role": "assistant",
            "content": answer
        })

    suggestion_questions = [
        "Sa Pa có gì đẹp?",
        "Mùa nào đi Fansipan đẹp?",
        "Gợi ý lịch trình đi Bắc Hà",
        "Chi phí đi Sa Pa 2 ngày",
        "Đền Bảo Hà có gì nổi bật?",
        "Thời tiết Lào Cai hôm nay"
    ]

    st.markdown("""
    <style>
    /* khung ngoài chatbot */
    .chatbot-page-title{
        max-width: 1420px;
        margin: 16px auto 8px auto;
        padding: 0 14px;
        box-sizing: border-box;
    }

    /* 2 khối chính */
    div[data-testid="stVerticalBlockBorderWrapper"]{
        border: 2px solid #bfeeee !important;
        border-radius: 28px !important;
        background: #ffffff !important;
        box-shadow: 0 14px 28px rgba(15,23,42,0.05) !important;
        padding: 14px 16px 16px 16px !important;
        min-height: 760px !important;
    }

    .chat-divider{
        width: 2px;
        height: 760px;
        margin: 0 auto;
        border-radius: 999px;
        background: linear-gradient(
            to bottom,
            rgba(143,231,231,0.00) 0%,
            rgba(143,231,231,0.35) 12%,
            rgba(143,231,231,0.50) 50%,
            rgba(143,231,231,0.35) 88%,
            rgba(143,231,231,0.00) 100%
        );
        filter: blur(0.6px);
        opacity: 0.65;
        box-shadow: 0 0 10px rgba(143,231,231,0.18);
    }

    .chat-left-logo-wrap{
        display: flex;
        justify-content: center;
        align-items: center;
        min-height: 290px;
        margin-bottom: 18px;
    }

    .chat-left-logo{
        width: 100%;
        max-width: 310px;
        object-fit: contain;
        display: block;
        filter: drop-shadow(0 10px 18px rgba(15,23,42,0.10));
        -webkit-mask-image: linear-gradient(
            to bottom,
            rgba(0,0,0,1) 0%,
            rgba(0,0,0,1) 72%,
            rgba(0,0,0,0.88) 82%,
            rgba(0,0,0,0.55) 90%,
            rgba(0,0,0,0.18) 96%,
            rgba(0,0,0,0) 100%
        );
        mask-image: linear-gradient(
            to bottom,
            rgba(0,0,0,1) 0%,
            rgba(0,0,0,1) 72%,
            rgba(0,0,0,0.88) 82%,
            rgba(0,0,0,0.55) 90%,
            rgba(0,0,0,0.18) 96%,
            rgba(0,0,0,0) 100%
        );
    }

    .chat-left-title{
        text-align: center;
        font-size: 16px;
        font-weight: 800;
        color: #334155;
        margin: 2px 0 16px 0;
        line-height: 1.6;
    }

    .chat-right-head{
        padding: 2px 6px 8px 6px;
        margin-bottom: 6px;
    }

    .chat-right-title{
        font-size: 34px;
        font-weight: 900;
        color: #111827;
        line-height: 1.1;
        margin-bottom: 6px;
    }

    .chat-right-subtitle{
        font-size: 15px;
        line-height: 1.8;
        color: #64748b;
    }

    .chat-bottom-note{
        margin-top: 8px;
        font-size: 13px;
        line-height: 1.6;
        color: #64748b;
        padding-left: 4px;
    }

    /* nút gợi ý bên trái */
        div[data-testid="stButton"] > button{
        width: 100%;
        min-height: 35px !important;
        border-radius: 16px !important;
        border: 1.5px solid #bfeeee !important;
        background: #ffffff !important;
        color: #4b6472 !important;
        font-size: 14px !important;
        font-weight: 600 !important;
        text-align: center !important;
        justify-content: center !important;
        padding: 0 14px !important;
        box-shadow: none !important;
        margin-bottom: 0px !important;
    }

    div[data-testid="stButton"] > button:hover{
        background: #f7ffff !important;
        border-color: #7ee4e4 !important;
        color: #334155 !important;
    }

    /* ô nhập */
    div[data-testid="stTextInput"] input{
        min-height: 36px !important;
        border-radius: 999px !important;
        border: 2px solid #111827 !important;
        background: #ffffff !important;
        color: #111827 !important;
        font-size: 17px !important;
        padding-left: 20px !important;
        padding-right: 20px !important;
        box-shadow: none !important;
    }

    div[data-testid="stTextInput"] input::placeholder{
        color: #94a3b8 !important;
    }

    @media (max-width: 992px){
        div[data-testid="stVerticalBlockBorderWrapper"]{
            min-height: auto !important;
        }

        .chat-divider{
            display: none;
        }

        .chat-right-title{
            font-size: 28px;
        }
                
    .chat-input-wrap{
    width: 100%;
    display: flex;
    justify-content: center;
    margin-top: 14px;
    margin-bottom: 6px;
}

    .chat-input-wrap > div{
        width: 92%;
        max-width: 760px;
    }

    .chat-input-wrap div[data-testid="stForm"]{
        margin: 0 !important;
        padding: 0 !important;
        background: transparent !important;
        box-shadow: none !important;
        border: none !important;
    }

    .chat-input-wrap div[data-testid="stTextInput"]{
        margin-bottom: 0 !important;
    }

    .chat-input-wrap{
        width: 100%;
        margin-top: 14px;
    }

    .chat-input-wrap div[data-testid="stForm"]{
        margin: 0 !important;
        padding: 0 !important;
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }

    .chat-input-wrap div[data-testid="stForm"] > div{
        width: 100%;
    }

    .chat-input-wrap div[data-testid="stHorizontalBlock"]{
        gap: 0 !important;
        align-items: center !important;
    }

    .chat-input-wrap div[data-testid="stTextInput"]{
        margin-bottom: 0 !important;
    }

    .chat-input-wrap div[data-testid="stTextInput"] > div{
        position: relative !important;
    }

    .chat-input-wrap div[data-testid="stTextInput"] input{
        height: 68px !important;
        min-height: 68px !important;
        border-radius: 999px !important;
        border: 1.8px solid #b9d9ff !important;
        background: #f3f6fb !important;
        color: #111827 !important;
        font-size: 18px !important;
        padding-left: 34px !important;
        padding-right: 88px !important;
        box-shadow: none !important;
    }

    .chat-input-wrap div[data-testid="stTextInput"] input::placeholder{
        color: #7b8ba1 !important;
    }

    .chat-input-wrap div[data-testid="stFormSubmitButton"]{
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        margin-left: -64px !important;
        z-index: 5 !important;
        width: 56px !important;
        min-width: 56px !important;
    }

    .chat-input-wrap div[data-testid="stFormSubmitButton"] button{
        width: 40px !important;
        height: 40px !important;
        min-width: 40px !important;
        min-height: 40px !important;
        border-radius: 50% !important;
        border: none !important;
        background: linear-gradient(180deg, #58b8ff 0%, #339af0 100%) !important;
        color: white !important;
        font-size: 20px !important;
        font-weight: 900 !important;
        padding: 0 !important;
        box-shadow: 0 6px 14px rgba(51,154,240,0.28) !important;
    }

    .chat-input-wrap div[data-testid="stFormSubmitButton"] button:hover{
        background: linear-gradient(180deg, #4aaeff 0%, #228be6 100%) !important;
    }

    .chat-input-wrap div[data-testid="stForm"] [data-testid="stCaptionContainer"],
    .chat-input-wrap div[data-testid="stForm"] p{
        display: none !important;
    }

    }
    </style>
    """, unsafe_allow_html=True)

    outer_left, outer_mid, outer_right = st.columns([0.02, 0.96, 0.02])

    with outer_mid:
        left_col, divider_col, right_col = st.columns([0.35, 0.02, 0.63], gap="small")

        with left_col:
            with st.container(border=False):
                if logo_src:
                    st.markdown(f"""
                    <div class="chat-left-logo-wrap">
                        <img class="chat-left-logo" src="{logo_src}" alt="anime_teamlogo">
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div class="chat-left-logo-wrap">
                        <div class="chat-left-title">Chưa tìm thấy ảnh anime_teamlogo.png</div>
                    </div>
                    """, unsafe_allow_html=True)

                for idx, question in enumerate(suggestion_questions):
                    if st.button(question, key=f"chat_suggest_{idx}", use_container_width=True):
                        send_chat_message(question)
                        st.rerun()

        with divider_col:
            st.markdown('<div class="chat-divider"></div>', unsafe_allow_html=True)

        with right_col:
            with st.container(border=False):
                st.markdown("""
                <div class="chat-right-head">
                    <div class="chat-right-title">Xin chào</div>
                    <div class="chat-right-subtitle">
                        Hỏi bất kì điều gì về điểm đến, lịch trình và trải nghiệm du lịch tại Lào Cai.
                    </div>
                </div>
                """, unsafe_allow_html=True)

                message_blocks = []
                for msg in st.session_state.gemini_messages[-40:]:
                    role = "user" if msg.get("role") == "user" else "assistant"
                    content = escape(str(msg.get("content", ""))).replace("\n", "<br>")
                    message_blocks.append(
                        f'<div class="msg-row {role}"><div class="msg-bubble {role}">{content}</div></div>'
                    )

                messages_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        html, body {{
                            margin: 0;
                            padding: 0;
                            background: transparent;
                            font-family: "Segoe UI", Arial, sans-serif;
                        }}

                        .msg-wrap {{
                            height: 560px;
                            overflow-y: auto;
                            padding: 8px 4px 10px 4px;
                            box-sizing: border-box;
                            background: #f7f7f7;
                            border-radius: 22px;
                        }}

                        .msg-row {{
                            display: flex;
                            margin-bottom: 16px;
                        }}

                        .msg-row.user {{
                            justify-content: flex-end;
                        }}

                        .msg-row.assistant {{
                            justify-content: flex-start;
                        }}

                        .msg-bubble {{
                            max-width: 76%;
                            padding: 14px 18px;
                            font-size: 16px;
                            line-height: 1.75;
                            color: #111827;
                            word-break: break-word;
                            box-sizing: border-box;
                            box-shadow: 0 6px 14px rgba(15,23,42,0.04);
                        }}

                        .msg-bubble.user {{
                            background: #dddddd;
                            border-radius: 18px 18px 8px 18px;
                        }}

                        .msg-bubble.assistant {{
                            background: #e9e9e9;
                            border-radius: 18px 18px 18px 8px;
                        }}
                    </style>
                </head>
                <body>
                    <div class="msg-wrap" id="msgWrap">
                        {''.join(message_blocks)}
                    </div>

                    <script>
                        const box = document.getElementById("msgWrap");
                        if (box) {{
                            box.scrollTop = box.scrollHeight;
                        }}
                    </script>
                </body>
                </html>
                """

                components.html(messages_html, height=575, scrolling=False)

                st.markdown('<div class="chat-input-wrap">', unsafe_allow_html=True)

                with st.form("gemini_chat_main_form", clear_on_submit=True, border=False):
                    input_col, send_col = st.columns([0.94, 0.06], gap="small")

                    with input_col:
                        user_text = st.text_input(
                            "Nhập câu hỏi",
                            placeholder="Hỏi bất cứ điều gì...",
                            label_visibility="collapsed"
                        )

                    with send_col:
                        submitted = st.form_submit_button("↑", use_container_width=True)

                st.markdown('</div>', unsafe_allow_html=True)

                if submitted and user_text.strip():
                    send_chat_message(user_text.strip())
                    st.rerun()

elif page == "gioithieu":
    import json
    from html import escape
    import streamlit.components.v1 as components

    def load_json_list(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data if isinstance(data, list) else []
        except Exception:
            return []

    DEFAULT_IMG = "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1600&auto=format&fit=crop"

    def safe_url(value, fallback=DEFAULT_IMG):
        text = str(value).strip() if value is not None else ""
        return escape(text if text else fallback, quote=True)

    def get_image(data, idx, fallback=DEFAULT_IMG):
        try:
            value = str(data[idx].get("image", "")).strip()
            return safe_url(value if value else fallback)
        except Exception:
            return safe_url(fallback)

    diemden_data = load_json_list("diemden.json")
    lichtrinh_data = load_json_list("lichtrinh.json")

    hero_image = get_image(diemden_data, 0)
    overview_image = image_to_data_uri([
        "assets/hanhtrinh.png",
        "assets/hanhtrinh.jpg",
        "assets/hanhtrinh.jpeg",
        "assets/hanhtrinh.webp"
    ])
    collage_img_1 = get_image(diemden_data, 2)
    collage_img_2 = get_image(diemden_data, 3)
    service_img_1 = get_image(diemden_data, 4)
    service_img_2 = get_image(diemden_data, 5)
    service_img_3 = get_image(diemden_data, 6)
    service_img_4 = get_image(diemden_data, 7)

    khonggianvanhoa_icon = image_to_data_uri([
        "assets/khonggianvanhoa.png",
        "assets/khonggianvanhoa.jpg",
        "assets/khonggianvanhoa.jpeg",
        "assets/khonggianvanhoa.webp"
    ])

    dulieudiaphuong_icon = image_to_data_uri([
        "assets/dulieudiaphuong.png",
        "assets/dulieudiaphuong.jpg",
        "assets/dulieudiaphuong.jpeg",
        "assets/dulieudiaphuong.webp"
    ])

    trolyai_icon = image_to_data_uri([
        "assets/troliai.png",
        "assets/troliai.jpg",
        "assets/troliai.jpeg",
        "assets/troliai.webp"
    ])

    dinhhuongbenvung_icon = image_to_data_uri([
        "assets/dinhhuongbenvung.png",
        "assets/dinhhuongbenvung.jpg",
        "assets/dinhhuongbenvung.jpeg",
        "assets/dinhhuongbenvung.webp"
    ])

    gioithieu_hero_image = image_to_data_uri([
        "assets/anime_teamtrangchu2.png"
    ])

    gt_story_img_1 = "https://cdn.vntrip.vn/cam-nang/wp-content/uploads/2017/03/%E1%BA%A3nh-1-5.jpg"
    gt_story_img_2 = "https://dulichtaybac.vn/wp-content/uploads/2025/11/LC-10.jpg"

    total_places = len(diemden_data)
    total_routes = len(lichtrinh_data)
    total_areas = len({str(item.get("area", "")).strip() for item in diemden_data if str(item.get("area", "")).strip()})
    total_categories = len({str(item.get("category", "")).strip() for item in diemden_data if str(item.get("category", "")).strip()})

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8" />
        <meta name="viewport" content="width=device-width, initial-scale=1.0" />
        <style>
            * {{
                box-sizing: border-box;
            }}

            html, body {{
                margin: 0;
                padding: 0;
                font-family: Inter, Arial, sans-serif;
                background: #f4f1ec;
                color: #243127;
            }}

            .gt-page {{
                width: 100%;
                padding: 28px 18px 52px 18px;
                background: linear-gradient(180deg, #f4f1ec 0%, #f7f4ef 100%);
            }}

            .gt-wrap {{
                max-width: 1360px;
                margin: 0 auto;
            }}

            /* ===== PHẦN 1 ===== */
            .gt-hero {{
                display: grid;
                grid-template-columns: 1.02fr 1.18fr;
                gap: 24px;
                align-items: center;
                margin-bottom: 22px;
            }}

            .gt-left {{
                padding: 18px 8px 18px 6px;
            }}

            .gt-kicker {{
                display: inline-flex;
                align-items: center;
                gap: 10px;
                font-size: 13px;
                font-weight: 700;
                letter-spacing: 0.12em;
                text-transform: uppercase;
                color: #8c7b66;
                margin-bottom: 18px;
            }}

            .gt-kicker::before {{
                content: "";
                width: 26px;
                height: 2px;
                background: #b59e82;
                border-radius: 999px;
            }}

            .gt-title {{
                font-size: 56px;
                line-height: 1.03;
                font-weight: 500;
                color: #4a5a4a;
                margin: 0 0 18px 0;
                font-family: Georgia, "Times New Roman", serif;
            }}

            .gt-sub {{
                font-size: 18px;
                line-height: 1.8;
                color: #68756a;
                margin-bottom: 22px;
                max-width: 560px;
            }}

            .gt-actions {{
                display: flex;
                gap: 14px;
                flex-wrap: wrap;
            }}

            .gt-btn {{
                min-width: 178px;
                height: 48px;
                border-radius: 12px;
                border: 1.5px solid #50624f;
                background: transparent;
                color: #50624f;
                font-size: 15px;
                font-weight: 700;
                cursor: pointer;
                display: inline-flex;
                align-items: center;
                justify-content: center;
                gap: 10px;
                transition: 0.2s ease;
            }}

            .gt-btn.primary {{
                background: #50624f;
                color: #ffffff;
                box-shadow: 0 10px 22px rgba(80, 98, 79, 0.18);
            }}

            .gt-btn:hover {{
                transform: translateY(-1px);
            }}

            .gt-image {{
                position: relative;
                min-height: 440px;
                border-radius: 0 0 0 120px;
                overflow: hidden;
                background-image: url('{gioithieu_hero_image}');
                background-size: cover;
                background-position: center;
                box-shadow: 0 18px 38px rgba(0,0,0,0.08);
            }}

            .gt-feature-bar {{
                display: grid;
                grid-template-columns: repeat(4, 1fr);
                gap: 12px;
                background: #f2efea;
                border: 1px solid #e7e0d6;
                border-radius: 18px;
                padding: 14px;
                margin-bottom: 34px;
                box-shadow: 0 10px 24px rgba(35, 49, 39, 0.04);
            }}

            .gt-feature {{
                display: flex;
                gap: 14px;
                align-items: flex-start;
                padding: 10px 10px;
            }}

            .gt-feature-icon{{
                width: 46px;
                height: 46px;
                min-width: 42px;
                border-radius: 12px;
                background: #f7f4ee;
                border: 1px solid #eadfce;
                display: flex;
                align-items: center;
                justify-content: center;
                overflow: hidden;
                padding: 8px;
            }}

            .gt-feature-icon img{{
                width: 100%;
                height: 100%;
                object-fit: contain;
                display: block;
            }}

            .gt-feature-title {{
                font-size: 16px;
                font-weight: 800;
                color: #4b5a4b;
                margin-bottom: 6px;
            }}

            .gt-feature-text {{
                font-size: 13px;
                line-height: 1.65;
                color: #7a8478;
            }}

            .gt-bottom {{
                display: grid;
                grid-template-columns: 0.82fr 1.18fr;
                gap: 26px;
                align-items: center;
                margin-bottom: 54px;
            }}

            .gt-bottom-left {{
                padding: 12px 4px;
            }}

            .gt-bottom-title {{
                font-size: 57px;
                line-height: 1.08;
                font-weight: 500;
                color: #4a5a4a;
                margin: 0 0 16px 0;
                font-family: Georgia, "Times New Roman", serif;
            }}

            .gt-bottom-text {{
                font-size: 16px;
                line-height: 1.9;
                color: #6f7a6e;
                max-width: 500px;
                margin-bottom: 24px;
            }}

            .gt-small-btn {{
                display: inline-flex;
                align-items: center;
                justify-content: center;
                min-width: 140px;
                height: 46px;
                border-radius: 12px;
                background: #50624f;
                color: white;
                font-size: 14px;
                font-weight: 800;
                border: none;
                cursor: pointer;
                box-shadow: 0 10px 22px rgba(80, 98, 79, 0.16);
            }}

            .gt-bottom-right {{
                position: relative;
                min-height: 390px;
                border-radius: 18px;
                overflow: hidden;
                background-image: url('{overview_image}');
                background-size: cover;
                background-position: center;
                box-shadow: 0 18px 38px rgba(0,0,0,0.08);
            }}

            .gt-bottom-right::before {{
                content: "";
                position: absolute;
                inset: 0;
                background: linear-gradient(180deg, rgba(255,255,255,0.10) 0%, rgba(0,0,0,0.18) 100%);
            }}

            .gt-stats{{
                position: absolute;
                left: 24px;
                right: 24px;
                bottom: 16px;
                display: grid;
                grid-template-columns: repeat(4, 1fr);
                gap: 10px;
                background: rgba(255, 255, 255, 0.16);
                border: 1px solid rgba(255, 255, 255, 0.22);
                border-radius: 18px;
                padding: 14px;
                backdrop-filter: blur(18px);
                -webkit-backdrop-filter: blur(18px);
                box-shadow: 0 8px 24px rgba(0,0,0,0.10);
            }}

            .gt-stat{{
                background: rgba(255, 255, 255, 0.08);
                border: 1px solid rgba(255, 255, 255, 0.16);
                border-radius: 14px;
                padding: 14px 10px;
                text-align: center;
                box-shadow: inset 0 1px 0 rgba(255,255,255,0.10);
                transition: transform 0.22s ease, box-shadow 0.22s ease, background 0.22s ease;
                cursor: pointer;
                position: relative;
            }}

            .gt-stat:hover{{
                transform: scale(1.08) translateY(-4px);
                background: rgba(255, 255, 255, 0.14);
                box-shadow:
                    0 14px 30px rgba(0,0,0,0.18),
                    inset 0 1px 0 rgba(255,255,255,0.16);
                z-index: 2;
            }}

            .gt-stat {{
                text-align: center;
                padding: 6px 4px;
            }}

            .gt-stat-value{{
                font-size: 24px;
                font-weight: 800;
                color: #ffffff;
                line-height: 1.1;
            }}

            .gt-stat-label{{
                font-size: 14px;
                font-weight: 700;
                color: #ffffff;
                margin-top: 4px;
            }}

            .gt-stat-sub{{
                font-size: 12px;
                line-height: 1.5;
                color: rgba(255,255,255,0.92);
                margin-top: 2px;
            }}

            /* ===== PHẦN 2 - GIAO DIỆN BẠN VỪA GỬI ===== */
            .gt-story-section {{
                background: #faf8f4;
                border-radius: 22px;
                padding: 42px 34px 40px 34px;
                margin-bottom: 0;
            }}

            .gt-story-grid {{
                display: grid;
                grid-template-columns: 0.88fr 1.12fr;
                gap: 34px;
                align-items: center;
            }}

            .gt-collage-wrap {{
                position: relative;
                min-height: 320px;
                display: flex;
                align-items: center;
                justify-content: center;
            }}

            .gt-collage-circle {{
                position: absolute;
                width: 240px;
                height: 240px;
                border-radius: 50%;
                background: #ece6dd;
                left: 0;
                top: 34px;
            }}

            .gt-collage-card-1 {{
                position: relative;
                width: 250px;
                height: 160px;
                background-image: url("{gt_story_img_1}");
                background-size: cover;
                background-position: center;
                box-shadow: 0 18px 34px rgba(0,0,0,0.14);
                z-index: 2;
                margin-left: 34px;
            }}

            .gt-collage-card-2 {{
                position: absolute;
                width: 170px;
                height: 115px;
                background-image: url("{gt_story_img_2}");
                background-size: cover;
                background-position: center;
                left: 78px;
                bottom: 36px;
                box-shadow: 0 16px 30px rgba(0,0,0,0.14);
                z-index: 3;
            }}

            .gt-badge-float {{
                position: absolute;
                left: 18px;
                top: 124px;
                width: 52px;
                height: 52px;
                border-radius: 50%;
                background: #d8c84b;
                display: flex;
                align-items: center;
                justify-content: center;
                color: #ffffff;
                font-size: 22px;
                font-weight: 900;
                z-index: 4;
                box-shadow: 0 10px 22px rgba(216,200,75,0.30);
            }}

            .gt-story-kicker {{
                font-size: 13px;
                font-weight: 700;
                letter-spacing: 0.16em;
                text-transform: uppercase;
                color: #9b8d7c;
                margin-bottom: 14px;
            }}

            .gt-story-title {{
                font-size: 42px;
                line-height: 1.1;
                font-weight: 500;
                color: #4a5a4a;
                margin: 0 0 16px 0;
                font-family: Georgia, "Times New Roman", serif;
                max-width: 520px;
            }}

            .gt-story-desc {{
                font-size: 15px;
                line-height: 1.9;
                color: #748071;
                max-width: 600px;
                margin-bottom: 22px;
            }}

            .gt-mini-features {{
                display: grid;
                grid-template-columns: repeat(2, 1fr);
                gap: 16px 18px;
                margin-bottom: 22px;
            }}

            .gt-mini-item {{
                display: flex;
                gap: 12px;
                align-items: flex-start;
            }}

            .gt-mini-icon {{
                width: 34px;
                height: 34px;
                border-radius: 10px;
                background: #f2eee8;
                border: 1px solid #e6ddd1;
                display: flex;
                align-items: center;
                justify-content: center;
                font-size: 17px;
                color: #7e8d6f;
                flex-shrink: 0;
            }}

            .gt-mini-title {{
                font-size: 14px;
                font-weight: 800;
                color: #4e5c4d;
                margin-bottom: 4px;
            }}

            .gt-mini-text {{
                font-size: 12px;
                line-height: 1.65;
                color: #7c857a;
            }}

            .gt-story-btn {{
                display: inline-flex;
                align-items: center;
                justify-content: center;
                min-width: 138px;
                height: 42px;
                border-radius: 10px;
                background: #d8c84b;
                color: #2f382d;
                font-size: 14px;
                font-weight: 800;
                border: none;
                cursor: pointer;
                box-shadow: 0 10px 22px rgba(216,200,75,0.22);
            }}

            /* ===== SERVICES ===== */
            .gt-services {{
                background: #ebe6de;
                margin-top: 0;
                padding: 52px 28px 40px 28px;
                border-radius: 0 0 22px 22px;
            }}

            .gt-services-head {{
                text-align: center;
                margin-bottom: 30px;
            }}

            .gt-services-kicker {{
                font-size: 12px;
                font-weight: 700;
                letter-spacing: 0.18em;
                text-transform: uppercase;
                color: #968775;
                margin-bottom: 10px;
            }}

            .gt-services-title {{
                font-size: 34px;
                line-height: 1.15;
                font-weight: 500;
                color: #4c5b4c;
                margin: 0;
                font-family: Georgia, "Times New Roman", serif;
            }}

            .gt-service-grid {{
                display: grid;
                grid-template-columns: repeat(4, 1fr);
                gap: 16px;
            }}

            .gt-service-card {{
                background: #f7f4ee;
                padding: 18px 14px 14px 14px;
                min-height: 290px;
                display: flex;
                flex-direction: column;
                justify-content: space-between;
                border: 1px solid #e3ddd4;
                box-shadow: 0 8px 18px rgba(0,0,0,0.04);
            }}

            .gt-service-number {{
                font-size: 12px;
                font-weight: 800;
                color: #b0a08f;
                letter-spacing: 0.10em;
                text-transform: uppercase;
                margin-bottom: 14px;
            }}

            .gt-service-name {{
                font-size: 18px;
                line-height: 1.3;
                font-weight: 800;
                color: #4f5d4f;
                margin-bottom: 10px;
                text-transform: uppercase;
            }}

            .gt-service-text {{
                font-size: 13px;
                line-height: 1.75;
                color: #7a8478;
                margin-bottom: 14px;
                min-height: 88px;
            }}

            .gt-service-image {{
                width: 100%;
                height: 95px;
                background-size: cover;
                background-position: center;
                margin-top: auto;
            }}

            @media (max-width: 1100px) {{
                .gt-hero,
                .gt-bottom,
                .gt-story-grid {{
                    grid-template-columns: 1fr;
                }}

                .gt-title {{
                    font-size: 48px;
                }}

                .gt-bottom-title {{
                    font-size: 42px;
                }}

                .gt-feature-bar {{
                    grid-template-columns: repeat(2, 1fr);
                }}

                .gt-service-grid {{
                    grid-template-columns: repeat(2, 1fr);
                }}
            }}

            @media (max-width: 720px) {{
                .gt-page {{
                    padding: 18px 10px 28px 10px;
                }}

                .gt-title {{
                    font-size: 38px;
                }}

                .gt-bottom-title,
                .gt-story-title {{
                    font-size: 34px;
                }}

                .gt-sub,
                .gt-bottom-text,
                .gt-story-desc {{
                    font-size: 15px;
                }}

                .gt-image{{
                    min-height: 420px;
                    border-radius: 0 0 0 110px;
                    background-image: url('assets/anime_teamtrangchu2.png');
                    background-size: contain;
                    background-repeat: no-repeat;
                    background-position: center center;
                    background-color: transparent;
                }}

                .gt-feature-bar,
                .gt-stats,
                .gt-mini-features,
                .gt-service-grid {{
                    grid-template-columns: 1fr;
                }}

                .gt-bottom-right {{
                    min-height: 420px;
                }}

                .gt-story-section {{
                    padding: 28px 16px 28px 16px;
                }}

                .gt-collage-wrap {{
                    min-height: 270px;
                }}

                .gt-collage-circle {{
                    width: 180px;
                    height: 180px;
                    left: 6px;
                    top: 34px;
                }}

                .gt-collage-card-1 {{
                    width: 210px;
                    height: 135px;
                    margin-left: 28px;
                }}

                .gt-collage-card-2 {{
                    width: 138px;
                    height: 92px;
                    left: 76px;
                    bottom: 42px;
                }}

                .gt-badge-float {{
                    width: 44px;
                    height: 44px;
                    left: 10px;
                    top: 112px;
                    font-size: 18px;
                }}

                .gt-services {{
                    padding: 34px 16px 24px 16px;
                }}
            }}
            
        </style>
    </head>
    <body>
        <div class="gt-page">
            <div class="gt-wrap">

                <section class="gt-hero">
                    <div class="gt-left">
                        <div class="gt-kicker">Dự án</div>
                        <h1 class="gt-title">LAO CAI HERITAGE AI</h1>
                        <div class="gt-sub">
                            Lao Cai Heritage AI là nền tảng giới thiệu văn hóa, lịch sử và các điểm đến tiêu biểu của Lào Cai theo hướng hiện đại,
                            trực quan và ứng dụng trí tuệ nhân tạo. Dự án giúp người dùng dễ dàng tra cứu thông tin, khám phá giá trị di sản địa phương và có trải nghiệm tìm hiểu sinh động, thuận tiện hơn.
                        </div>

                    </div>

                    <div class="gt-image"></div>
                </section>

                <section class="gt-feature-bar">
                    <div class="gt-feature">
                        <div class="gt-feature-icon">
                            <img src="{khonggianvanhoa_icon}" alt="Không gian văn hóa">
                        </div>
                        <div>
                            <div class="gt-feature-title">Không gian văn hóa</div>
                            <div class="gt-feature-text">
                                Tôn vinh vẻ đẹp di sản, bản sắc địa phương và giá trị truyền thống của Lào Cai.
                            </div>
                        </div>
                    </div>

                    <div class="gt-feature">
                        <div class="gt-feature-icon">
                            <img src="{dulieudiaphuong_icon}" alt="Dữ liệu địa phương">
                        </div>
                        <div>
                            <div class="gt-feature-title">Dữ liệu địa phương</div>
                            <div class="gt-feature-text">
                                Kết nối thông tin về điểm đến, khu vực, mùa đẹp và nội dung giới thiệu rõ ràng.
                            </div>
                        </div>
                    </div>

                    <div class="gt-feature">
                        <div class="gt-feature-icon">
                            <img src="{trolyai_icon}" alt="Trợ lý AI thông minh">
                        </div>
                        <div>
                            <div class="gt-feature-title">Trợ lý AI thông minh</div>
                            <div class="gt-feature-text">
                                Hỗ trợ hỏi đáp, tư vấn lịch trình và giải đáp nhanh các câu hỏi cơ bản cho du khách.
                            </div>
                        </div>
                    </div>

                    <div class="gt-feature">
                        <div class="gt-feature-icon">
                            <img src="{dinhhuongbenvung_icon}" alt="Định hướng bền vững">
                        </div>
                        <div>
                            <div class="gt-feature-title">Định hướng bền vững</div>
                            <div class="gt-feature-text">
                                Ứng dụng công nghệ để quảng bá văn hóa địa phương theo hướng hiện đại và lâu dài.
                            </div>
                        </div>
                    </div>
                </section>

                <section class="gt-bottom">
                    <div class="gt-bottom-left">
                        <div class="gt-kicker">Giới thiệu dự án</div>
                        <h2 class="gt-bottom-title">Mở ra hành trình khám phá mới</h2>
                        <div class="gt-bottom-text">
                            Lao Cai Heritage AI được phát triển với mục tiêu hỗ trợ giới thiệu danh lam thắng cảnh, ại và dễ tiếp cận. 
                            Dự án hướng tới việc kết hợp hiệu quả giữa dữ liệu số và trí tuệ nhân tạo nhằm nâng cao trải nghiệm tra cứu thông tin, tham khảo lịch trình và khám phá Lào Cai một cách thuận tiện, sinh động hơn.
                            di tích lịch sử và những nét đẹp văn hóa đặc trưng của địa phương bằng hình thức trực quan, hiện đ
                        </div>

                    </div>

                    <div class="gt-bottom-right">
                        <div class="gt-stats">
                            <div class="gt-stat">
                                <div class="gt-stat-value">{total_places}+</div>
                                <div class="gt-stat-label">Điểm đến</div>
                                <div class="gt-stat-sub">Nội dung trong hệ thống</div>
                            </div>

                            <div class="gt-stat">
                                <div class="gt-stat-value">{total_routes}+</div>
                                <div class="gt-stat-label">Lịch trình</div>
                                <div class="gt-stat-sub">Gợi ý tham khảo</div>
                            </div>

                            <div class="gt-stat">
                                <div class="gt-stat-value">{total_areas}+</div>
                                <div class="gt-stat-label">Khu vực</div>
                                <div class="gt-stat-sub">Phân loại dữ liệu</div>
                            </div>

                            <div class="gt-stat">
                                <div class="gt-stat-value">{total_categories}+</div>
                                <div class="gt-stat-label">Loại hình</div>
                                <div class="gt-stat-sub">Trải nghiệm khám phá</div>
                            </div>
                        </div>
                    </div>
                </section>

                <section class="gt-story-section">
                    <div class="gt-story-grid">
                        <div class="gt-collage-wrap">
                            <div class="gt-collage-circle"></div>
                            <div class="gt-badge-float">✦</div>
                            <div class="gt-collage-card-1"></div>
                            <div class="gt-collage-card-2"></div>
                        </div>

                        <div>
                            <div class="gt-story-kicker">Giá trị cốt lõi</div>
                            <h2 class="gt-story-title">Giải pháp quảng bá văn hóa và du lịch địa phương</h2>
                            <div class="gt-story-desc">
                                Dự án không chỉ tập trung vào việc giới thiệu thông tin mà còn hướng đến việc tạo ra trải nghiệm tra cứu hiện đại, 
                                thuận tiện và gần gũi với nhiều nhóm người dùng như học sinh, du khách và cộng đồng quan tâm đến văn hóa địa phương. 
                                Hệ thống được xây dựng với giao diện trực quan, nội dung mạch lạc và định hướng phát triển bền vững trong tương lai.
                            </div>

                            <div class="gt-mini-features">
                                <div class="gt-mini-item">
                                    <div class="gt-mini-icon">🧭</div>
                                    <div>
                                        <div class="gt-mini-title">Định vị nội dung rõ ràng</div>
                                        <div class="gt-mini-text">Phân tách điểm đến, lịch trình và thuyết minh thành từng nhóm dễ tra cứu.</div>
                                    </div>
                                </div>

                                <div class="gt-mini-item">
                                    <div class="gt-mini-icon">📚</div>
                                    <div>
                                        <div class="gt-mini-title">Tăng giá trị học tập</div>
                                        <div class="gt-mini-text">Giúp người dùng tìm hiểu lịch sử, văn hóa và thông tin địa phương thuận tiện hơn.</div>
                                    </div>
                                </div>

                                <div class="gt-mini-item">
                                    <div class="gt-mini-icon">⚙️</div>
                                    <div>
                                        <div class="gt-mini-title">Tích hợp AI thực tế</div>
                                        <div class="gt-mini-text">Chatbot hỗ trợ hỏi đáp cơ bản, gợi ý lịch trình và trả lời theo dữ liệu hệ thống.</div>
                                    </div>
                                </div>

                                <div class="gt-mini-item">
                                    <div class="gt-mini-icon">🚀</div>
                                    <div>
                                        <div class="gt-mini-title">Dễ mở rộng về sau</div>
                                        <div class="gt-mini-text">Có thể tiếp tục bổ sung phản hồi, dữ liệu mới và các tính năng nâng cao sau này.</div>
                                    </div>
                                </div>
                            </div>

                        </div>
                    </div>
                </section>

                <section class="gt-services">
                    <div class="gt-services-head">
                        <div class="gt-services-kicker">Tính năng</div>
                        <h2 class="gt-services-title">Các chức năng nổi bật</h2>
                    </div>

                    <div class="gt-service-grid">
                        <div class="gt-service-card">
                            <div>
                                <div class="gt-service-number">01</div>
                                <div class="gt-service-name">Giới thiệu điểm đến</div>
                                <div class="gt-service-text">
                                    Cung cấp thông tin về các điểm đến nổi bật của Lào Cai thông qua cách trình bày trực quan, rõ ràng và dễ theo dõi. Người dùng có thể nhanh chóng tìm hiểu đặc điểm, giá trị văn hóa, lịch sử và những nét nổi bật của từng địa danh.
                                </div>
                            </div>
                            <div class="gt-service-image" style="background-image:url('{service_img_1}');"></div>
                        </div>

                        <div class="gt-service-card">
                            <div>
                                <div class="gt-service-number">02</div>
                                <div class="gt-service-name">Lịch trình thông minh</div>
                                <div class="gt-service-text">
                                    Đề xuất các hành trình tham khảo dựa trên khu vực, loại hình trải nghiệm và mức chi phí cơ bản. Tính năng này giúp người dùng thuận tiện hơn trong việc lựa chọn lịch trình phù hợp với thời gian, nhu cầu và mục đích khám phá thực tế.
                                </div>
                            </div>
                            <div class="gt-service-image" style="background-image:url('{service_img_2}');"></div>
                        </div>

                        <div class="gt-service-card">
                            <div>
                                <div class="gt-service-number">03</div>
                                <div class="gt-service-name">Chatbot AI</div>
                                <div class="gt-service-text">
                                    Hỗ trợ giải đáp nhanh các câu hỏi liên quan đến địa điểm, mùa đẹp, lịch trình gợi ý và những thông tin cần thiết trong quá trình tham khảo. Tính năng này góp phần tăng khả năng tương tác và giúp người dùng tiếp cận thông tin một cách linh hoạt hơn.
                                </div>
                            </div>
                            <div class="gt-service-image" style="background-image:url('{service_img_3}');"></div>
                        </div>

                        <div class="gt-service-card">
                            <div>
                                <div class="gt-service-number">04</div>
                                <div class="gt-service-name">Thuyết minh địa điểm</div>
                                <div class="gt-service-text">
                                    Mang đến trải nghiệm khám phá sinh động hơn thông qua nội dung thuyết minh được tích hợp ngay tại trang chi tiết. Người dùng không chỉ đọc thông tin mà còn có thể tiếp cận nội dung theo cách trực quan, hấp dẫn và gần gũi hơn.
                                </div>
                            </div>
                            <div class="gt-service-image" style="background-image:url('{service_img_4}');"></div>
                        </div>
                    </div>
                </section>

            </div>
        </div>
    </body>
    </html>
    """

    components.html(html, height=2220, scrolling=False)

elif page == "caidat":
    current_user = get_user_by_id(st.session_state["auth_user"]["id"])

    if not current_user:
        st.error("Không tìm thấy thông tin tài khoản.")
        st.stop()

    st.markdown("""
    <style>

    .settings-title{
        font-size: 26px;
        font-weight: 900;
        color: #111827;
        margin-bottom: 18px;
    }

    .settings-card{
        width: 100%;
        background:#ffffff;
        border:1px solid #e5edf6;
        border-radius:20px;
        padding:24px;
        box-shadow:0 14px 32px rgba(15,23,42,0.06);
        box-sizing: border-box;
    }
                
    div[data-testid="stTabs"] + div{
        margin-top: 0 !important;
        padding-top: 0 !important;
    }

    .settings-section-title{
        font-size:16px;
        font-weight:900;
        color:#111827;
        margin-bottom:14px;
    }

    .settings-box{
        width: 100%;
        background:#f8fafc;
        border:1px solid #e2e8f0;
        border-radius:14px;
        padding:16px 18px;
        margin-bottom:16px;
        box-sizing: border-box;
    }

    .settings-label{
        font-size:13px;
        font-weight:800;
        color:#64748b;
        margin-bottom:8px;
        text-transform:uppercase;
        letter-spacing:.04em;
    }

    .settings-value{
        font-size:17px;
        font-weight:700;
        color:#0f172a;
        line-height:1.6;
        word-break:break-word;
    }

    .settings-note{
        font-size:13px;
        color:#64748b;
        line-height:1.7;
        margin-top:6px;
    }

    .settings-action-row{
        display:flex;
        justify-content:center;
        margin-top: 0px;
    }

    .settings-logout-btn{
        display:inline-flex;
        align-items:center;
        justify-content:center;
        min-width:160px;
        height:46px;
        padding:0 18px;
        border-radius:12px;
        background:linear-gradient(135deg, #dc2626 0%, #ef4444 100%);
        color:#ffffff !important;
        text-decoration:none !important;
        font-size:15px;
        font-weight:800;
        box-shadow:0 10px 22px rgba(220,38,38,0.18);
        transition:transform 0.18s ease, box-shadow 0.18s ease;
    }

    .settings-logout-btn:hover{
        transform:translateY(-1px);
        box-shadow:0 14px 26px rgba(220,38,38,0.24);
        color:#ffffff !important;
    }

    div[data-testid="stTabs"]{
        margin-bottom: 0 !important;
    }

    div[data-testid="stTabs"] > div{
        max-width: 980px;
        margin: 0 auto;
        background: transparent !important;
        box-shadow: none !important;
        border: none !important;
    }
                
    div[data-testid="stTabs"] [role="tablist"]{
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
    }

    div[data-testid="stTabs"] [data-baseweb="tab-border"]{
        display: none !important;
    }
                

    div[data-testid="stForm"]{
        width: 100%;
        background:#f8fafc !important;
        border:1px solid #e2e8f0 !important;
        border-radius:14px !important;
        padding:16px 18px !important;
        box-shadow:none !important;
        margin-top:0 !important;
    }

    div[data-testid="stTextInput"]{
        width: 100%;
    }

    div[data-testid="stTextInput"] input{
        min-height: 48px !important;
        border-radius: 12px !important;
        border: 1px solid #dbe4ef !important;
        background: #ffffff !important;
        box-shadow: none !important;
    }

    div[data-testid="stFormSubmitButton"]{
        margin-top: 10px;
    }

    div[data-testid="stFormSubmitButton"] button{
        border-radius:12px !important;
        min-height:46px !important;
        font-weight:800 !important;
        padding: 0 18px !important;
    }

    @media (max-width: 768px){
        .settings-outer{
            padding: 84px 12px 28px 12px;
        }

        .settings-title{
            font-size: 22px;
            margin-bottom: 14px;
        }

        .settings-card{
            border-radius: 16px;
            padding: 16px;
        }

        .settings-box{
            border-radius: 12px;
            padding: 14px;
            margin-bottom: 12px;
        }

        .settings-value{
            font-size: 15px;
        }

        .settings-action-row{
            justify-content: stretch;
        }

        .settings-logout-btn{
            width: 100%;
        }
    }
    </style>
    """, unsafe_allow_html=True)

    tab1, tab2 = st.tabs(["Thông tin tài khoản", "Mật khẩu & Bảo mật"])

    with tab1:
        st.markdown('<div class="settings-section-title">Thông tin tài khoản</div>', unsafe_allow_html=True)

        st.markdown(f'''
        <div class="settings-box">
            <div class="settings-label">Tên đăng nhập</div>
            <div class="settings-value">{escape(current_user["username"])}</div>
        </div>
        ''', unsafe_allow_html=True)

        st.markdown(f'''
        <div class="settings-box">
            <div class="settings-label">Mã người dùng</div>
            <div class="settings-value">USER-{current_user["id"]}</div>
        </div>
        ''', unsafe_allow_html=True)

        st.markdown(f'''
        <div class="settings-box">
            <div class="settings-label">Ngày tạo tài khoản</div>
            <div class="settings-value">{escape(str(current_user["created_at"]))}</div>
        </div>
        ''', unsafe_allow_html=True)

        st.markdown(f'''
        <div class="settings-action-row">
            <a class="settings-logout-btn" href="{build_app_url('home', logout=1)}" target="_self">
                Đăng xuất
            </a>
        </div>
        ''', unsafe_allow_html=True)

    with tab2:
        st.markdown('<div class="settings-card">', unsafe_allow_html=True)
        st.markdown('<div class="settings-section-title">Mật khẩu & Bảo mật</div>', unsafe_allow_html=True)

        st.markdown(f'''
        <div class="settings-box">
            <div class="settings-label">Tên đăng nhập</div>
            <div class="settings-value">{escape(current_user["username"])}</div>
        </div>
        ''', unsafe_allow_html=True)

        st.markdown('''
        <div class="settings-box">
            <div class="settings-label">Mật khẩu hiện tại</div>
            <div class="settings-value">••••••••</div>
        </div>
        ''', unsafe_allow_html=True)

        with st.form("change_password_form", clear_on_submit=True):
            old_password = st.text_input("Mật khẩu hiện tại", type="password")
            new_password = st.text_input("Mật khẩu mới", type="password")
            confirm_password = st.text_input("Nhập lại mật khẩu mới", type="password")
            submit_change_password = st.form_submit_button("Lưu thay đổi", use_container_width=False)

        st.markdown('</div>', unsafe_allow_html=True)

        if submit_change_password:
            if new_password != confirm_password:
                st.error("Mật khẩu mới nhập lại chưa khớp.")
            else:
                ok, msg = update_user_password(
                    st.session_state["auth_user"]["id"],
                    old_password,
                    new_password
                )
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)

        st.markdown('</div>', unsafe_allow_html=True)

elif page == "congdong":
    from html import escape

    def safe_text(value, fallback="Đang cập nhật"):
        text = str(value).strip() if value is not None else ""
        return escape(text) if text else escape(fallback)

    def safe_url(value):
        text = str(value).strip() if value is not None else ""
        if text.startswith(("http://", "https://", "data:")):
            return escape(text, quote=True)
        return ""

    community_posts = load_community_posts(limit=50)

    def build_avatar_label(username: str) -> str:
        name = str(username or "").strip()
        return (name[:1] or "U").upper()
    
    def normalize_post_image_src(value: str) -> str:
        raw = str(value or "").strip()
        if not raw:
            return ""

        raw = raw.replace("&amp;", "&")

        if raw.startswith("data:"):
            return raw

        if raw.startswith(("http://", "https://")):
            # Google Drive: đổi link share sang link xem ảnh trực tiếp
            if "drive.google.com/file/d/" in raw:
                try:
                    file_id = raw.split("/file/d/")[1].split("/")[0]
                    return f"https://drive.google.com/uc?export=view&id={file_id}"
                except Exception:
                    return raw

            # GitHub blob -> raw
            if "github.com" in raw and "/blob/" in raw:
                raw = raw.replace("github.com/", "raw.githubusercontent.com/")
                raw = raw.replace("/blob/", "/")
                return raw

            return raw

        # nếu là đường dẫn local trong project
        local_candidates = [
            raw,
            raw.replace(".jpg", ".png"),
            raw.replace(".jpg", ".jpeg"),
            raw.replace(".jpg", ".webp"),
            raw.replace(".png", ".jpg"),
            raw.replace(".png", ".jpeg"),
            raw.replace(".png", ".webp"),
            raw.replace(".jpeg", ".jpg"),
            raw.replace(".jpeg", ".png"),
            raw.replace(".jpeg", ".webp"),
            raw.replace(".webp", ".jpg"),
            raw.replace(".webp", ".png"),
            raw.replace(".webp", ".jpeg"),
        ]

        local_image = image_to_data_uri(local_candidates)
        return local_image if local_image else ""

    hero_bg = hero_home_src if hero_home_src else ""

    st.markdown(f"""
    <style>
    .cd-page {{
        max-width: 1320px;
        margin: 0 auto 42px auto;
        padding: 0 12px 24px 12px;
        box-sizing: border-box;
    }}

    .cd-hero {{
        position: relative;
        margin: 18px auto 22px auto;
        min-height: 320px;
        border-radius: 30px;
        overflow: hidden;
        background:
            linear-gradient(180deg, rgba(7,19,37,0.30) 0%, rgba(7,19,37,0.60) 100%),
            url('{hero_bg}');
        background-size: cover;
        background-position: center;
        display: flex;
        align-items: center;
        padding: 34px 34px;
        box-shadow: 0 18px 38px rgba(15,23,42,0.10);
    }}

    .cd-hero-inner {{
        max-width: 760px;
        color: #ffffff;
        position: relative;
        z-index: 2;
    }}

    .cd-kicker {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        padding: 8px 14px;
        border-radius: 999px;
        background: rgba(255,255,255,0.16);
        border: 1px solid rgba(255,255,255,0.28);
        font-size: 13px;
        font-weight: 800;
        letter-spacing: 0.06em;
        text-transform: uppercase;
        margin-bottom: 14px;
    }}

    .cd-title {{
        font-size: 40px;
        font-weight: 900;
        line-height: 1.15;
        margin: 0 0 12px 0;
    }}

    .cd-sub {{
        font-size: 17px;
        line-height: 1.85;
        color: rgba(255,255,255,0.94);
        max-width: 700px;
    }}

    .cd-grid {{
        display: grid;
        grid-template-columns: 420px 1fr;
        gap: 22px;
        align-items: start;
    }}

    .cd-compose {{
        position: sticky;
        top: 92px;
        background: rgba(255,255,255,0.92);
        border: 1px solid #e5edf6;
        border-radius: 26px;
        padding: 18px;
        box-shadow: 0 16px 32px rgba(15,23,42,0.08);
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
    }}

    .cd-compose-title {{
        font-size: 22px;
        font-weight: 900;
        color: #0f172a;
        margin-bottom: 6px;
    }}

    .cd-compose-sub {{
        font-size: 14px;
        line-height: 1.7;
        color: #64748b;
        margin-bottom: 14px;
    }}

    .cd-feed {{
        display: flex;
        flex-direction: column;
        gap: 18px;
    }}

    .cd-post {{
        background: #ffffff;
        border: 1px solid #e5edf6;
        border-radius: 24px;
        overflow: hidden;
        box-shadow: 0 14px 28px rgba(15,23,42,0.06);
    }}

    .cd-post-head {{
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 14px;
        padding: 18px 18px 12px 18px;
    }}

    .cd-user {{
        display: flex;
        align-items: center;
        gap: 12px;
        min-width: 0;
    }}

    .cd-avatar {{
        width: 46px;
        height: 46px;
        border-radius: 50%;
        background: linear-gradient(135deg, #1565c0 0%, #43b0f1 100%);
        color: #ffffff;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 18px;
        font-weight: 900;
        flex-shrink: 0;
    }}

    .cd-user-main {{
        min-width: 0;
    }}

    .cd-username {{
        font-size: 16px;
        font-weight: 900;
        color: #0f172a;
        line-height: 1.3;
    }}

    .cd-time {{
        font-size: 13px;
        color: #64748b;
        line-height: 1.5;
        margin-top: 2px;
    }}

    .cd-badges {{
        display: flex;
        gap: 8px;
        flex-wrap: wrap;
        justify-content: flex-end;
    }}

    .cd-badge {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        padding: 7px 11px;
        border-radius: 999px;
        background: #ecf5ff;
        color: #1565c0;
        border: 1px solid #cfe5ff;
        font-size: 12px;
        font-weight: 800;
        white-space: nowrap;
    }}

    .cd-badge-alt {{
        background: #fff7ed;
        color: #c2410c;
        border: 1px solid #fed7aa;
    }}

    .cd-content {{
        padding: 0 18px 16px 18px;
        font-size: 16px;
        line-height: 1.9;
        color: #334155;
        white-space: pre-wrap;
    }}

    .cd-image {{
        width: 100%;
        height: 330px;
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        border-top: 1px solid #edf2f7;
    }}

    .cd-empty {{
        background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
        border: 1px dashed #cbd5e1;
        border-radius: 24px;
        padding: 36px 22px;
        text-align: center;
        color: #475569;
        font-size: 15px;
        line-height: 1.8;
    }}

    div[data-testid="stForm"] {{
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
        padding: 0 !important;
    }}

    div[data-testid="stTextArea"] textarea {{
        min-height: 140px !important;
        border-radius: 18px !important;
        border: 1px solid #dbe4ef !important;
        background: #f8fbff !important;
        font-size: 15px !important;
    }}

    div[data-testid="stTextInput"] input,
    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {{
        min-height: 48px !important;
        border-radius: 16px !important;
        border: 1px solid #dbe4ef !important;
        background: #f8fbff !important;
        font-size: 14px !important;
    }}

    div[data-testid="stFormSubmitButton"] button {{
        width: 100% !important;
        min-height: 48px !important;
        border: none !important;
        border-radius: 16px !important;
        background: linear-gradient(135deg, #1565c0 0%, #2b7fd3 100%) !important;
        color: #ffffff !important;
        font-size: 15px !important;
        font-weight: 800 !important;
        box-shadow: 0 14px 24px rgba(21,101,192,0.20) !important;
    }}

    @media (max-width: 980px) {{
        .cd-grid {{
            grid-template-columns: 1fr;
        }}

        .cd-compose {{
            position: static;
            top: auto;
        }}
    }}

    @media (max-width: 768px) {{
        .cd-page {{
            padding: 0 10px 20px 10px;
        }}

        .cd-hero {{
            min-height: 250px;
            padding: 22px 18px;
            border-radius: 22px;
        }}

        .cd-title {{
            font-size: 28px;
        }}

        .cd-sub {{
            font-size: 14px;
            line-height: 1.75;
        }}

        .cd-post-head {{
            flex-direction: column;
            align-items: flex-start;
        }}

        .cd-badges {{
            justify-content: flex-start;
        }}

        .cd-image {{
            height: 240px;
        }}
    }}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="cd-page">
        <div class="cd-hero">
            <div class="cd-hero-inner">
                <div class="cd-kicker">Không gian chia sẻ</div>
                <div class="cd-title">Cộng đồng Lao Cai Heritage AI</div>
                <div class="cd-sub">
                    Nơi mọi người cùng giới thiệu điểm đến, lễ hội, làng nghề, chia sẻ hình ảnh thực tế
                    và lan tỏa vẻ đẹp văn hóa – du lịch Lào Cai theo cách gần gũi, hiện đại.
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <style>
    .cd-native-post{
        background:#ffffff;
        border:1px solid #e5edf6;
        border-radius:24px;
        padding:16px 16px 14px 16px;
        box-shadow:0 14px 28px rgba(15,23,42,0.06);
        margin-bottom:18px;
    }

    .cd-native-userrow{
        display:flex;
        align-items:center;
        gap:12px;
        margin-bottom:10px;
    }

    .cd-native-avatar{
        width:46px;
        height:46px;
        border-radius:50%;
        background:linear-gradient(135deg, #1565c0 0%, #43b0f1 100%);
        color:#ffffff;
        display:flex;
        align-items:center;
        justify-content:center;
        font-size:18px;
        font-weight:900;
        flex-shrink:0;
    }

    .cd-native-name{
        font-size:16px;
        font-weight:900;
        color:#0f172a;
        line-height:1.2;
    }

    .cd-native-time{
        font-size:13px;
        color:#64748b;
        margin-top:2px;
    }

    .cd-native-pill{
        display:inline-block;
        padding:6px 10px;
        border-radius:999px;
        font-size:12px;
        font-weight:800;
        margin-right:8px;
        margin-bottom:8px;
        border:1px solid #dbe4ef;
        background:#f8fbff;
        color:#1565c0;
    }

    .cd-native-pill-alt{
        background:#fff7ed;
        color:#c2410c;
        border:1px solid #fed7aa;
    }

    .cd-native-content{
        font-size:15px;
        line-height:1.85;
        color:#334155;
        white-space:pre-wrap;
        word-break:break-word;
        margin-top:2px;
    }

    .cd-native-image-wrap img{
        border-radius:18px !important;
        border:1px solid #edf2f7;
    }
    </style>
    """, unsafe_allow_html=True)

    page_left, page_center, page_right = st.columns([0.06, 0.88, 0.06])

    with page_center:
        left_col, right_col = st.columns([0.34, 0.66], gap="large")

        with left_col:
            st.markdown('<div class="cd-compose">', unsafe_allow_html=True)
            st.markdown('<div class="cd-compose-title">Đăng chia sẻ mới</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="cd-compose-sub">Viết cảm nhận, giới thiệu địa điểm và thêm link ảnh để bài đăng sinh động hơn.</div>',
                unsafe_allow_html=True
            )

            with st.form("community_post_form", clear_on_submit=True):
                post_content = st.text_area(
                    "Bạn muốn chia sẻ gì?",
                    placeholder="Ví dụ: Mình vừa ghé lễ hội Gầu Tào ở Bắc Hà và thấy không khí rất đặc sắc...",
                    key="community_post_content"
                )

                c1, c2 = st.columns(2)
                with c1:
                    post_category = st.selectbox(
                        "Loại nội dung",
                        ["Trải nghiệm", "Điểm đến", "Lễ hội", "Làng nghề", "Di tích"],
                        key="community_post_category"
                    )
                with c2:
                    post_area = st.text_input(
                        "Khu vực",
                        placeholder="Ví dụ: Sa Pa, Bắc Hà...",
                        key="community_post_area"
                    )

                post_image_url = st.text_input(
                    "Link ảnh",
                    placeholder="Dán link ảnh tại đây",
                    key="community_post_image_url"
                )

                submit_post = st.form_submit_button("Đăng bài", use_container_width=True)

            if submit_post:
                ok, msg = create_community_post(
                    user_id=st.session_state["auth_user"]["id"],
                    username=st.session_state["auth_user"]["username"],
                    content=post_content,
                    image_url=post_image_url,
                    category=post_category,
                    area=post_area
                )
                if ok:
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)

            st.markdown('</div>', unsafe_allow_html=True)

        with right_col:
            if community_posts:
                for post in community_posts:
                    username = str(post.get("username", "")).strip() or "Người dùng"
                    avatar_letter = build_avatar_label(username)
                    content = str(post.get("content", "")).strip()
                    category = str(post.get("category", "")).strip() or "Chia sẻ"
                    area = str(post.get("area", "")).strip() or "Lào Cai"
                    created_at = format_post_time(str(post.get("created_at", "")).strip())
                    image_url = str(post.get("image_url", "")).strip()

                    st.markdown('<div class="cd-native-post">', unsafe_allow_html=True)

                    # Hàng đầu: avatar + tên + thời gian
                    st.markdown(f"""
                    <div class="cd-native-userrow">
                        <div class="cd-native-avatar">{escape(avatar_letter)}</div>
                        <div>
                            <div class="cd-native-name">{escape(username)}</div>
                            <div class="cd-native-time">{escape(created_at)}</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                    # Badge
                    st.markdown(f"""
                    <div>
                        <span class="cd-native-pill">{escape(category)}</span>
                        <span class="cd-native-pill cd-native-pill-alt">{escape(area)}</span>
                    </div>
                    """, unsafe_allow_html=True)

                    # Nội dung
                    st.markdown(
                        f'<div class="cd-native-content">{escape(content)}</div>',
                        unsafe_allow_html=True
                    )

                    # Ảnh
                    raw_image_url = str(post.get("image_url", "")).strip()
                image_src = normalize_post_image_src(raw_image_url)

                if image_src:
                    st.markdown('<div class="cd-native-image-wrap">', unsafe_allow_html=True)
                    st.image(image_src, use_container_width=True)
                    st.markdown('</div>', unsafe_allow_html=True)

                    st.markdown('</div>', unsafe_allow_html=True)

            else:
                st.info("Chưa có bài đăng nào trong cộng đồng. Hãy là người đầu tiên chia sẻ một địa điểm, một lễ hội hoặc một trải nghiệm đẹp về Lào Cai.")

elif page == "lehoi_langnghe":
    import json
    from html import escape
    from pathlib import Path
    import fitz
    import re
    import unicodedata

    def clean_text(value):
        if value is None:
            return ""
        return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()

    def normalize_file_name(value):
        text = clean_text(value).lower()
        text = text.replace("đ", "d").replace("Đ", "D")
        text = unicodedata.normalize("NFD", text)
        text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
        text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
        return text

    def read_text_file_any(path_obj):
        path_obj = Path(path_obj)
        if not path_obj.exists():
            return ""

        suffix = path_obj.suffix.lower()

        try:
            if suffix == ".txt":
                return path_obj.read_text(encoding="utf-8").strip()

            elif suffix == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(path_obj))
                    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    return "\n".join(parts).strip()
                except Exception:
                    return ""

            elif suffix == ".pdf":
                try:
                    pdf = fitz.open(str(path_obj))
                    texts = []
                    for page in pdf:
                        txt = page.get_text("text").strip()
                        if txt:
                            texts.append(txt)
                    pdf.close()
                    return "\n".join(texts).strip()
                except Exception:
                    return ""

        except Exception:
            return ""

        return ""

    def pick_intro_text_file(item):
        folder = Path("text_lehoi_langnghe")
        base_name = normalize_file_name(item.get("name", ""))

        candidates = [
            folder / f"{base_name}.txt",
            folder / f"{base_name}.docx",
            folder / f"{base_name}.pdf",
        ]

        for fp in candidates:
            if fp.exists():
                return read_text_file_any(fp)

        return ""

    def build_intro_preview(item):
        intro_text = pick_intro_text_file(item)

        if not intro_text:
            intro_text = clean_text(item.get("full_desc")) or clean_text(item.get("short_desc"))

        intro_text = " ".join(intro_text.split())

        if len(intro_text) > 260:
            intro_text = intro_text[:260].rsplit(" ", 1)[0] + "..."

        return escape(intro_text if intro_text else "Thông tin giới thiệu đang được cập nhật.")

    try:
        with open("lehoi_langnghe.json", "r", encoding="utf-8") as f:
            lehoi_data = json.load(f)
    except Exception as e:
        st.error(f"Lỗi đọc file lehoi_langnghe.json: {e}")
        st.stop()

    def safe_text(value, fallback="Đang cập nhật"):
        text = str(value).strip() if value is not None else ""
        return escape(text) if text else escape(fallback)

    def normalize_text(value):
        return str(value or "").strip().lower()

    def card_image_src(value):
        raw = str(value).strip() if value is not None else ""
        if not raw:
            return ""
        if raw.startswith(("http://", "https://", "data:")):
            return raw
        return image_to_data_uri([
            raw,
            raw.replace(".jpg", ".png"),
            raw.replace(".jpg", ".jpeg"),
            raw.replace(".jpg", ".webp"),
            raw.replace(".png", ".jpg"),
            raw.replace(".png", ".jpeg"),
            raw.replace(".png", ".webp"),
        ]) or ""

    st.markdown("""
    <style>
    .lln-hero{
        max-width:1320px;
        margin:20px auto 16px auto;
        padding:48px 28px;
        border-radius:28px;
        background: linear-gradient(135deg, #1565c0 0%, #2e7d32 100%);
        color:white;
        text-align:center;
        box-shadow:0 18px 38px rgba(15,23,42,0.10);
    }
    .lln-title{
        font-size:38px;
        font-weight:900;
        margin-bottom:10px;
    }
    .lln-sub{
        font-size:17px;
        line-height:1.8;
        max-width:850px;
        margin:0 auto;
        opacity:0.96;
    }
    .lln-grid{
        max-width:1320px;
        margin:18px auto 50px auto;
        display:grid;
        grid-template-columns:repeat(3, minmax(0, 1fr));
        gap:24px;
    }
    .lln-card{
        background:#ffffff;
        border-radius:24px;
        overflow:hidden;
        box-shadow:0 12px 28px rgba(15,23,42,0.08);
        border:1px solid #e5edf6;
    }
    .lln-img{
        height:220px;
        background-size:cover;
        background-position:center;
        background-repeat:no-repeat;
    }
    .lln-body{
        padding:18px;
    }elif page == "lehoi_langnghe":
        display:inline-block;
        padding:6px 12px;
        border-radius:999px;
        background:#ecf5ff;
        color:#1565c0;
        font-size:12px;
        font-weight:800;
        margin-bottom:10px;
    }
    .lln-name{
        font-size:22px;
        font-weight:900;
        color:#111827;
        margin-bottom:10px;
        line-height:1.3;
    }
    .lln-area{
        font-size:14px;
        color:#64748b;
        margin-bottom:10px;
    }
    .lln-desc{
        font-size:15px;
        line-height:1.7;
        color:#475569;
    }
    
                
    .lln-filter-wrap{
        max-width:1320px;
        margin: 6px auto 10px auto;
        padding: 0 8px;
        box-sizing:border-box;
    }

    .lln-filter-title{
        text-align:center;
        font-size:15px;
        font-weight:800;
        color:#334155;
        margin-bottom:10px;
        letter-spacing:0.01em;
    }

    div[data-testid="stSelectbox"][aria-label="Lọc nội dung"],
    div[data-testid="stSelectbox"]:has(label p:contains("Lọc nội dung")){
        width:100%;
    }

    div[data-testid="stSelectbox"] label{
        font-size:15px !important;
        font-weight:800 !important;
        color:#334155 !important;
        margin-bottom:8px !important;
    }

    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div{
        min-height:56px !important;
        border-radius:18px !important;
        border:1.6px solid #dbe4ef !important;
        background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%) !important;
        box-shadow: 0 10px 24px rgba(15,23,42,0.06) !important;
        padding-left:16px !important;
        padding-right:16px !important;
        font-size:18px !important;
        font-weight:700 !important;
        color:#0f172a !important;
        transition: all 0.2s ease !important;
    }

    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div:hover{
        border-color:#93c5fd !important;
        box-shadow: 0 14px 28px rgba(21,101,192,0.10) !important;
    }

    div[data-testid="stSelectbox"] div[data-baseweb="select"] > div:focus-within{
        border-color:#1565c0 !important;
        box-shadow: 0 0 0 4px rgba(21,101,192,0.12) !important;
    }

    @media (max-width: 640px){
        .lln-filter-wrap{
            padding:0 10px;
            margin: 4px auto 10px auto;
        }

        div[data-testid="stSelectbox"] div[data-baseweb="select"] > div{
            min-height:52px !important;
            border-radius:16px !important;
            font-size:16px !important;
        }
    }
                
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="lln-hero">
        <div class="lln-title">Lễ hội - Làng nghề Lào Cai</div>
        <div class="lln-sub">
            Khám phá các lễ hội truyền thống và làng nghề đặc sắc của Lào Cai,
            nơi lưu giữ những giá trị văn hóa bản địa độc đáo của vùng cao.
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="lln-filter-wrap">', unsafe_allow_html=True)

    left_filter, center_filter, right_filter = st.columns([1.2, 2.2, 1.2])

    with center_filter:
        selected_type = st.selectbox(
            "",
            ["Tất cả", "Lễ hội", "Làng nghề"],
            key="lln_type_filter"
        )

    st.markdown('</div>', unsafe_allow_html=True)
    
    filtered_data = lehoi_data
    if selected_type == "Lễ hội":
        filtered_data = [x for x in lehoi_data if normalize_text(x.get("type")) == "le_hoi"]
    elif selected_type == "Làng nghề":
        filtered_data = [x for x in lehoi_data if normalize_text(x.get("type")) == "lang_nghe"]

    html_cards = '<div class="lln-grid">'

    for item in filtered_data:
        item_type = normalize_text(item.get("type"))
        badge_text = "Lễ hội" if item_type == "le_hoi" else "Làng nghề"

        image_src = card_image_src(item.get("image"))
        if not image_src:
            image_src = "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee?q=80&w=1200&auto=format&fit=crop"

        intro_preview = build_intro_preview(item)

        html_cards += f"""
        <div class="lln-card">
            <div class="lln-img" style="background-image:url('{image_src}');"></div>

            <div class="lln-body">
                <div class="lln-badge">{badge_text}</div>
                <div class="lln-name">{safe_text(item.get("name"))}</div>
                <div class="lln-area">
                    Khu vực: {safe_text(item.get("area"))} • Thời điểm: {safe_text(item.get("season"))}
                </div>
                <div class="lln-desc">{safe_text(item.get("short_desc"))}</div>
            </div>

            <div class="lln-hover">
                <div class="lln-hover-box">
                    <div class="lln-hover-badge">{badge_text}</div>
                    <div class="lln-hover-name">{safe_text(item.get("name"))}</div>
                    <div class="lln-hover-intro">{intro_preview}</div>
                </div>
            </div>
        </div>
        """

    html_cards += "</div>"

    card_count = len(filtered_data)
    if card_count == 0:
        html_cards = """
        <div style="
            max-width:1320px;
            margin:18px auto 50px auto;
            background:#ffffff;
            border:1px dashed #cbd5e1;
            border-radius:24px;
            padding:40px 24px;
            text-align:center;
            color:#475569;
            font-size:16px;
            line-height:1.8;
        ">
            Chưa có dữ liệu phù hợp để hiển thị.
        </div>
        """
        card_height = 220
    else:
        rows = (card_count + 2) // 3
        if card_count == 1:
            rows = 1
        card_height = max(420, rows * 470)

    components.html(
        f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                html, body {{
                    margin: 0;
                    padding: 0;
                    background: transparent;
                    font-family: Inter, Arial, sans-serif;
                }}

                .lln-grid {{
                    max-width:1320px;
                    margin:18px auto 50px auto;
                    display:grid;
                    grid-template-columns:repeat(3, minmax(0, 1fr));
                    gap:24px;
                    padding:0 8px;
                    box-sizing:border-box;
                }}

                <style>
                html, body {{
                    margin: 0;
                    padding: 0;
                    background: transparent;
                    font-family: Inter, Arial, sans-serif;
                }}

                .lln-grid {{
                    max-width:1320px;
                    margin:18px auto 50px auto;
                    display:grid;
                    grid-template-columns:repeat(3, minmax(0, 1fr));
                    gap:24px;
                    padding:0 8px;
                    box-sizing:border-box;
                }}

                .lln-card {{
                    position: relative;
                    min-height: 495px;
                    border-radius: 26px;
                    overflow: hidden;
                    background: #ffffff;
                    border: 1px solid #e5edf6;
                    box-shadow: 0 12px 28px rgba(15,23,42,0.08);
                    transition: transform 0.28s ease, box-shadow 0.28s ease;
                }}

                .lln-card:hover {{
                    transform: translateY(-6px);
                    box-shadow: 0 18px 36px rgba(15,23,42,0.14);
                }}

                .lln-img {{
                    height: 250px;
                    background-size: cover;
                    background-position: center;
                    background-repeat: no-repeat;
                    transition: transform 0.55s ease;
                }}

                .lln-card:hover .lln-img {{
                    transform: scale(1.06);
                }}

                .lln-body {{
                    position: relative;
                    padding: 18px 20px 20px 20px;
                    background: #ffffff;
                    z-index: 2;
                    transition: opacity 0.28s ease, transform 0.28s ease;
                }}

                .lln-card:hover .lln-body {{
                    opacity: 0.08;
                    transform: translateY(6px);
                }}

                .lln-badge {{
                    display:inline-block;
                    padding:6px 12px;
                    border-radius:999px;
                    background:#ecf5ff;
                    color:#1565c0;
                    font-size:12px;
                    font-weight:800;
                    margin-bottom:10px;
                }}

                .lln-name {{
                    font-size:22px;
                    font-weight:900;
                    color:#111827;
                    margin-bottom:10px;
                    line-height:1.3;
                }}

                .lln-area {{
                    font-size:14px;
                    color:#64748b;
                    margin-bottom:12px;
                    line-height:1.7;
                }}

                .lln-desc {{
                    font-size:15px;
                    line-height:1.75;
                    color:#475569;
                }}

                .lln-hover {{
                    position: absolute;
                    inset: 0;
                    z-index: 3;
                    display: flex;
                    align-items: flex-end;
                    padding: 22px;
                    background: linear-gradient(
                        180deg,
                        rgba(15,23,42,0.04) 0%,
                        rgba(15,23,42,0.12) 28%,
                        rgba(15,23,42,0.78) 72%,
                        rgba(15,23,42,0.92) 100%
                    );
                    opacity: 0;
                    pointer-events: none;
                    transition: opacity 0.35s ease;
                }}

                .lln-card:hover .lln-hover {{
                    opacity: 1;
                }}

                .lln-hover-box {{
                    width: 100%;
                    transform: translateY(28px);
                    transition: transform 0.38s ease;
                }}

                .lln-card:hover .lln-hover-box {{
                    transform: translateY(0);
                }}

                .lln-hover-badge {{
                    display:inline-block;
                    padding:7px 12px;
                    border-radius:999px;
                    background: rgba(255,255,255,0.16);
                    border: 1px solid rgba(255,255,255,0.24);
                    color:#ffffff;
                    font-size:12px;
                    font-weight:800;
                    margin-bottom:10px;
                }}

                .lln-hover-name {{
                    font-size:24px;
                    font-weight:900;
                    color:#ffffff;
                    line-height:1.3;
                    margin-bottom:10px;
                    text-shadow: 0 4px 14px rgba(0,0,0,0.24);
                }}

                .lln-hover-intro {{
                    font-size:15px;
                    line-height:1.8;
                    color: rgba(255,255,255,0.94);
                    display: -webkit-box;
                    -webkit-line-clamp: 6;
                    -webkit-box-orient: vertical;
                    overflow: hidden;
                    text-shadow: 0 2px 8px rgba(0,0,0,0.25);
                }}

                @media (max-width: 640px) {{
                    .lln-grid {{
                        grid-template-columns:1fr;
                        padding:0 10px;
                    }}

                    .lln-card {{
                        min-height: 460px;
                    }}

                    .lln-img {{
                        height: 220px;
                    }}

                    .lln-name,
                    .lln-hover-name {{
                        font-size: 20px;
                    }}

                    .lln-desc,
                    .lln-hover-intro {{
                        font-size: 14px;
                    }}
                }}
                </style>
                
            </style>
        </head>
        <body>
            {html_cards}
        </body>
        </html>
        """,
        height=card_height,
        scrolling=False
    )